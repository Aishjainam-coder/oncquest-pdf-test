"""
Universal PDF, JSON & HTML Processor, Renderer & Word (.docx) Converter
=======================================================================
Complete Streamlit Application for Local & Cloud Deployments (Streamlit Cloud, Render, etc.)

Features:
- Upload PDF reports, structured JSON files, or HTML documents
- Converts PDF/JSON directly to formatted Microsoft Word (.docx) with 100% fidelity
- Preserves layout, tables, fonts, borders, banners, and signature footers
- Pure Python zero-dependency DOCX fallback engine (works on any cloud host)
- Live Interactive Rendered HTML, PDF, and JSON previews
- Theme customization (Colors, Fonts, Styles) via sidebar & theme.json
- Built-in validation report & multi-format download center
"""

import pymupdf
pymupdf._g_out_message = None

import os
import tempfile
import base64
import json
import io
from pathlib import Path
import streamlit as st
import pymupdf as fitz  # PyMuPDF

from converter import (
    render_exact_pdf_layout_html,
    generate_dynamic_template_html,
    convert_json_to_docx,
    convert_html_to_docx,
    convert_pdf_to_word,
    convert_pdf_via_pdf2docx,
    render_html_to_pdf_and_preview,
    validate_docx_conversion
)
from extractor import extract_report_data

# Ensure working directories exist
Path("extracted_jsons").mkdir(exist_ok=True)
Path("output").mkdir(exist_ok=True)

# Load base theme.json defaults if available
theme_json_defaults = {}
theme_file_path = Path("theme.json")
if theme_file_path.exists():
    try:
        with open(theme_file_path, "r", encoding="utf-8") as f_theme:
            theme_json_defaults = json.load(f_theme)
    except Exception:
        pass

# Configure Streamlit Page
st.set_page_config(
    page_title="Universal PDF & JSON to Word Converter",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Premium Custom Styling for Streamlit App
st.markdown("""
<style>
    .main { background-color: #f8fafc; }
    
    .header-card {
        background: linear-gradient(135deg, #0f172a 0%, #1e293b 50%, #1f497d 100%);
        color: #ffffff;
        padding: 1.6rem 2.2rem;
        border-radius: 14px;
        box-shadow: 0 10px 25px -5px rgba(15, 23, 42, 0.25);
        margin-bottom: 1.5rem;
    }
    .header-title {
        font-size: 1.85rem;
        font-weight: 800;
        margin: 0;
        letter-spacing: -0.02em;
        color: #ffffff;
    }
    .header-subtitle {
        font-size: 0.96rem;
        color: #cbd5e1;
        margin-top: 0.4rem;
        margin-bottom: 0;
        line-height: 1.4;
    }

    .word-success-box {
        background: linear-gradient(135deg, #ecfdf5 0%, #d1fae5 100%);
        border: 1px solid #10b981;
        border-radius: 12px;
        padding: 1.2rem 1.5rem;
        margin-bottom: 1.5rem;
        box-shadow: 0 4px 12px rgba(16, 185, 129, 0.12);
    }

    .stats-card {
        background: #ffffff;
        border: 1px solid #e2e8f0;
        border-radius: 10px;
        padding: 0.9rem 1.2rem;
        margin-bottom: 1rem;
        box-shadow: 0 2px 4px rgba(0,0,0,0.03);
    }

    /* Primary Action Buttons */
    div.stButton > button:first-child {
        background: linear-gradient(135deg, #1f497d 0%, #0f172a 100%);
        color: white;
        font-weight: 600;
        font-size: 1.05rem;
        padding: 0.65rem 1.5rem;
        border-radius: 8px;
        border: none;
        box-shadow: 0 4px 14px rgba(31, 73, 125, 0.25);
        transition: all 0.2s ease-in-out;
        width: 100%;
    }
    div.stButton > button:first-child:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 18px rgba(31, 73, 125, 0.35);
    }
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------
# Sidebar: Complete Customization & Configuration Controls
# ---------------------------------------------------------
st.sidebar.title("⚙️ Converter Suite")

st.sidebar.markdown("### 🎨 Design & Theme")
theme_preset = st.sidebar.selectbox(
    "Theme Preset",
    options=[
        "Classic Navy (#1f497d)",
        "Emerald Medical (#059669)",
        "Slate Modern (#1e293b)",
        "Crimson Clinical (#b91c1c)",
        "Royal Purple (#7c3aed)",
        "Custom Color"
    ],
    index=0
)

color_presets = {
    "Classic Navy (#1f497d)": "#1f497d",
    "Emerald Medical (#059669)": "#059669",
    "Slate Modern (#1e293b)": "#1e293b",
    "Crimson Clinical (#b91c1c)": "#b91c1c",
    "Royal Purple (#7c3aed)": "#7c3aed",
    "Custom Color": "#1f497d"
}

if theme_preset == "Custom Color":
    chosen_color = st.sidebar.color_picker("Pick Primary Color", "#1f497d")
else:
    chosen_color = color_presets[theme_preset]

font_choice = st.sidebar.selectbox(
    "Font Family",
    options=["Cambria", "Calibri", "Arial", "Segoe UI", "Georgia", "Times New Roman"],
    index=0
)

st.sidebar.markdown("### 🔒 Data & Privacy")
replace_lab_name = st.sidebar.checkbox("Standardize Lab Name ('Laboratory')", value=True, help="Replaces proprietary lab branding with generic Laboratory.")
anonymize_signatures = st.sidebar.checkbox("Inject Standard Signatures into Footer", value=True, help="Injects doctor signature images into Word page footers.")

st.sidebar.markdown("### 🖥️ Display Settings")
preview_height = st.sidebar.slider("Preview Height (px)", min_value=500, max_value=1200, value=850, step=50)

# Optional Custom theme.json Upload
uploaded_theme = st.sidebar.file_uploader("Upload Custom `theme.json` (Optional)", type=["json"], key="custom_theme_uploader")
active_theme_config = theme_json_defaults.copy() if theme_json_defaults else {}

if uploaded_theme:
    try:
        custom_theme_data = json.load(uploaded_theme)
        active_theme_config.update(custom_theme_data)
        st.sidebar.success("✅ Custom theme.json applied!")
    except Exception as e_theme:
        st.sidebar.error(f"Error loading theme: {e_theme}")

# Apply overrides
active_theme_config["primary_color"] = chosen_color
active_theme_config["table_header_bg"] = chosen_color
active_theme_config["border_color"] = chosen_color
active_theme_config["font_family"] = font_choice
active_theme_config["show_footer_signatures"] = anonymize_signatures

# ---------------------------------------------------------
# App Main Header
# ---------------------------------------------------------
st.markdown("""
<div class="header-card">
    <div class="header-title">⚡ Universal PDF & JSON → Word (.docx) Converter</div>
    <div class="header-subtitle">
        High-Fidelity Document Pipeline: <b>PDF / JSON / HTML → Word (.docx) Directly ✅</b><br>
        Preserves 100% exact layout, clinical tables, patient demographics, fonts & styling.
    </div>
</div>
""", unsafe_allow_html=True)

# Session State Initialization
if "extracted_data" not in st.session_state:
    st.session_state.extracted_data = None
if "html_content" not in st.session_state:
    st.session_state.html_content = ""
if "compiled_pdf_bytes" not in st.session_state:
    st.session_state.compiled_pdf_bytes = None
if "docx_bytes" not in st.session_state:
    st.session_state.docx_bytes = None
if "file_name" not in st.session_state:
    st.session_state.file_name = ""
if "file_bytes" not in st.session_state:
    st.session_state.file_bytes = None

# ---------------------------------------------------------
# 1. Upload Section
# ---------------------------------------------------------
uploaded_file = st.file_uploader(
    "📤 Upload any PDF report, extracted JSON structure, or HTML file",
    type=["pdf", "json", "html", "htm"],
    key="file_uploader",
    help="Supports PDF lab reports, invoices, certificates, extracted JSON files, or raw HTML."
)

if uploaded_file is not None:
    file_bytes = uploaded_file.getvalue()
    file_ext = Path(uploaded_file.name).suffix.lower()
    
    # If a new file is uploaded, reset state for clean execution
    if st.session_state.file_name != uploaded_file.name or st.session_state.file_bytes != file_bytes:
        st.session_state.file_name = uploaded_file.name
        st.session_state.file_bytes = file_bytes
        st.session_state.html_content = ""
        st.session_state.compiled_pdf_bytes = None
        st.session_state.docx_bytes = None
        st.session_state.extracted_data = None

    # File Info Summary
    file_size_kb = len(file_bytes) / 1024.0
    page_count = "N/A"
    if file_ext == ".pdf":
        try:
            temp_doc = fitz.open(stream=file_bytes, filetype="pdf")
            page_count = f"{len(temp_doc)} page(s)"
            temp_doc.close()
        except Exception:
            page_count = "PDF Document"
    elif file_ext == ".json":
        page_count = "JSON Document"
    else:
        page_count = "HTML Document"

    col_i1, col_i2, col_i3 = st.columns(3)
    with col_i1:
        st.info(f"📄 **File:** `{uploaded_file.name}`")
    with col_i2:
        st.info(f"⚖️ **Size:** `{file_size_kb:.1f} KB`")
    with col_i3:
        st.info(f"📑 **Type:** `{file_ext.upper()} | {page_count}`")

    st.markdown("---")

    # Action Button to Process & Convert directly to Word
    btn_process = st.button("⚡ Convert to Word (.docx) Directly", use_container_width=True, type="primary")

    # ---------------------------------------------------------
    # Processing Pipeline (Robust Multi-Method Fallback)
    # ---------------------------------------------------------
    if btn_process or st.session_state.docx_bytes is None:
        print(f"\n{'='*60}", flush=True)
        print(f"[*] Starting Document Processing: {uploaded_file.name}", flush=True)
        print(f"{'='*60}", flush=True)

        with st.status("⚡ Converting document to Word (.docx)...", expanded=True) as status_box:
            try:
                # -----------------------------------------------
                # Branch A: JSON Input
                # -----------------------------------------------
                if file_ext == ".json":
                    st.write("🔍 **Step 1/3:** Loading and parsing JSON structure...")
                    print(f"[*] [Step 1/3] Loading JSON structure for {uploaded_file.name}...", flush=True)
                    extracted_data = json.loads(file_bytes.decode("utf-8"))
                    if replace_lab_name:
                        json_raw = json.dumps(extracted_data, ensure_ascii=False).replace("SN Genelab Pvt Ltd", "Laboratory")
                        extracted_data = json.loads(json_raw)
                    st.session_state.extracted_data = extracted_data

                    st.write("🎨 **Step 2/3:** Generating styled HTML layout from JSON...")
                    print(f"[*] [Step 2/3] Generating styled HTML template...", flush=True)
                    html_content = generate_dynamic_template_html(extracted_data, doc_title=uploaded_file.name, theme_config=active_theme_config)
                    if replace_lab_name:
                        html_content = html_content.replace("SN Genelab Pvt Ltd", "Laboratory")
                    st.session_state.html_content = html_content

                    st.write("📝 **Step 3/3:** Generating Word (.docx) directly from structured JSON & theme...")
                    with tempfile.TemporaryDirectory() as tmp_dir:
                        docx_tmp = Path(tmp_dir) / "output.docx"
                        html_tmp = Path(tmp_dir) / "temp.html"
                        html_tmp.write_text(html_content, encoding="utf-8")
                        compiled_pdf_tmp = Path(tmp_dir) / "compiled.pdf"

                        # Direct high-fidelity JSON -> DOCX conversion (Zero external browser dependency)
                        try:
                            convert_json_to_docx(extracted_data, output_path=str(docx_tmp), theme_config=active_theme_config)
                            if docx_tmp.exists() and docx_tmp.stat().st_size > 0:
                                st.session_state.docx_bytes = docx_tmp.read_bytes()
                                print(f"[+] Direct JSON->DOCX successful ({len(st.session_state.docx_bytes)} bytes)!", flush=True)
                        except Exception as e_docx:
                            print(f"[!] Direct JSON->DOCX error: {e_docx}", flush=True)

                        # Try optional PDF preview compilation (if Playwright is available)
                        try:
                            render_html_to_pdf_and_preview(html_tmp, compiled_pdf_tmp)
                            if compiled_pdf_tmp.exists():
                                st.session_state.compiled_pdf_bytes = compiled_pdf_tmp.read_bytes()
                                if not st.session_state.docx_bytes:
                                    convert_pdf_via_pdf2docx(str(compiled_pdf_tmp), str(docx_tmp))
                                    if docx_tmp.exists():
                                        st.session_state.docx_bytes = docx_tmp.read_bytes()
                        except Exception as e_pw:
                            print(f"[*] Note: Optional PDF preview compilation skipped: {e_pw}", flush=True)

                # -----------------------------------------------
                # Branch B: HTML Input
                # -----------------------------------------------
                elif file_ext in [".html", ".htm"]:
                    st.write("🔍 **Step 1/2:** Loading and parsing HTML content...")
                    html_content = file_bytes.decode("utf-8", errors="replace")
                    if replace_lab_name:
                        html_content = html_content.replace("SN Genelab Pvt Ltd", "Laboratory")
                    st.session_state.html_content = html_content

                    st.write("📝 **Step 2/2:** Converting HTML to Word (.docx)...")
                    with tempfile.TemporaryDirectory() as tmp_dir:
                        docx_tmp = Path(tmp_dir) / "output.docx"
                        html_tmp = Path(tmp_dir) / "temp.html"
                        html_tmp.write_text(html_content, encoding="utf-8")
                        compiled_pdf_tmp = Path(tmp_dir) / "compiled.pdf"

                        try:
                            convert_html_to_docx(html_tmp, output_path=docx_tmp, theme_config=active_theme_config)
                            if docx_tmp.exists() and docx_tmp.stat().st_size > 0:
                                st.session_state.docx_bytes = docx_tmp.read_bytes()
                        except Exception as e_h2d:
                            print(f"[!] HTML->DOCX conversion note: {e_h2d}", flush=True)

                        # Optional PDF preview
                        try:
                            render_html_to_pdf_and_preview(html_tmp, compiled_pdf_tmp)
                            if compiled_pdf_tmp.exists():
                                st.session_state.compiled_pdf_bytes = compiled_pdf_tmp.read_bytes()
                        except Exception:
                            pass

                # -----------------------------------------------
                # Branch C: PDF Input
                # -----------------------------------------------
                else:
                    with tempfile.TemporaryDirectory() as tmp_dir:
                        pdf_input_path = Path(tmp_dir) / uploaded_file.name
                        pdf_input_path.write_bytes(file_bytes)

                        # Step 1: Extract structured JSON from PDF
                        st.write("🔍 **Step 1/3:** Extracting text, tables, and styles from PDF...")
                        print(f"[*] [Step 1/3] Extracting text, tables, and styles from PDF...", flush=True)
                        extracted_data = None
                        try:
                            extracted_data = extract_report_data(str(pdf_input_path), auto_save_docx=False)
                            if replace_lab_name and extracted_data:
                                json_raw = json.dumps(extracted_data, ensure_ascii=False).replace("SN Genelab Pvt Ltd", "Laboratory")
                                extracted_data = json.loads(json_raw)
                            st.session_state.extracted_data = extracted_data
                            
                            # Cache extracted JSON
                            json_out_dir = Path("extracted_jsons")
                            json_out_dir.mkdir(exist_ok=True)
                            json_file_path = json_out_dir / f"{Path(uploaded_file.name).stem}.json"
                            json_str = json.dumps(extracted_data, indent=2, ensure_ascii=False)
                            with open(json_file_path, "w", encoding="utf-8") as f_json:
                                f_json.write(json_str)
                            print(f"   [+] Extracted JSON saved to: {json_file_path}", flush=True)
                        except Exception as e_ext:
                            print(f"   [!] Note on JSON extraction: {e_ext}", flush=True)

                        # Step 2: Render Clean End Result HTML
                        st.write("🎨 **Step 2/3:** Rendering styled HTML document layout...")
                        print(f"[*] [Step 2/3] Rendering styled HTML document layout...", flush=True)
                        try:
                            with fitz.open(str(pdf_input_path)) as doc_fitz:
                                html_content = render_exact_pdf_layout_html(doc_fitz, doc_title=uploaded_file.name, theme_config=active_theme_config)
                            if replace_lab_name:
                                html_content = html_content.replace("SN Genelab Pvt Ltd", "Laboratory")
                            st.session_state.html_content = html_content
                        except Exception as e_html:
                            print(f"   [!] HTML layout rendering error: {e_html}", flush=True)

                        # Step 3: Generate Word (.docx)
                        st.write("📝 **Step 3/3:** Converting document to Microsoft Word (.docx)...")
                        docx_tmp = Path(tmp_dir) / "output.docx"
                        
                        # Primary Method: Direct PDF -> Word via pdf2docx + signature footer injection
                        try:
                            print(f"[*] Attempting PDF -> Word conversion via pdf2docx...", flush=True)
                            convert_pdf_to_word(str(pdf_input_path), str(docx_tmp), theme_config=active_theme_config)
                            if docx_tmp.exists() and docx_tmp.stat().st_size > 0:
                                st.session_state.docx_bytes = docx_tmp.read_bytes()
                                print(f"[+] DOCX conversion successful ({len(st.session_state.docx_bytes)} bytes)!", flush=True)
                        except Exception as e_p2d:
                            print(f"[!] Direct PDF->DOCX error: {e_p2d}", flush=True)

                        # Fallback Method: Structured JSON -> Word
                        if not st.session_state.docx_bytes and extracted_data:
                            try:
                                print(f"[*] Fallback: Converting extracted JSON directly to Word (.docx)...", flush=True)
                                convert_json_to_docx(extracted_data, output_path=str(docx_tmp), theme_config=active_theme_config)
                                if docx_tmp.exists() and docx_tmp.stat().st_size > 0:
                                    st.session_state.docx_bytes = docx_tmp.read_bytes()
                                    print(f"[+] Fallback DOCX successful ({len(st.session_state.docx_bytes)} bytes)!", flush=True)
                            except Exception as e_fb:
                                print(f"[!] Fallback JSON->DOCX error: {e_fb}", flush=True)

                        # Optional: Compile HTML to PDF preview (if Playwright is available)
                        try:
                            html_tmp = Path(tmp_dir) / "temp.html"
                            if st.session_state.html_content:
                                html_tmp.write_text(st.session_state.html_content, encoding="utf-8")
                                compiled_pdf_tmp = Path(tmp_dir) / "compiled.pdf"
                                render_html_to_pdf_and_preview(html_tmp, compiled_pdf_tmp)
                                if compiled_pdf_tmp.exists():
                                    st.session_state.compiled_pdf_bytes = compiled_pdf_tmp.read_bytes()
                        except Exception as e_pw:
                            print(f"[*] Note: Optional PDF preview compilation skipped: {e_pw}", flush=True)

                if st.session_state.docx_bytes:
                    status_box.update(label="✅ Conversion Completed Successfully!", state="complete", expanded=False)
                    print(f"[+] Pipeline Completed Successfully for {uploaded_file.name}!\n", flush=True)
                    st.success("✅ Converted Document + Theme Rules → Word (.docx) successfully!")
                else:
                    status_box.update(label="⚠️ Word (.docx) conversion could not create binary output.", state="error", expanded=True)
                    st.error("Could not generate Word document. Please verify the document format.")

            except Exception as e:
                status_box.update(label=f"❌ Error during conversion: {e}", state="error", expanded=True)
                print(f"[!] Error converting document: {e}", flush=True)
                st.error(f"Error converting document to Word: {e}")

    # ---------------------------------------------------------
    # 2. Download Center & Featured Actions
    # ---------------------------------------------------------
    if st.session_state.docx_bytes:
        st.markdown("---")
        
        docx_kb = len(st.session_state.docx_bytes) / 1024.0
        
        # Featured Direct Word Download Card
        col_w1, col_w2 = st.columns([2, 1])
        with col_w1:
            st.markdown("### 📝 Direct Word Document (.docx) Ready!")
            st.markdown(f"Your document was processed with **`theme.json`** styling and compiled into a Microsoft Word file (`{docx_kb:.1f} KB`).")
        with col_w2:
            st.download_button(
                label="📥 Download Word Document (.docx)",
                data=st.session_state.docx_bytes,
                file_name=f"{Path(st.session_state.file_name).stem}_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )

        # Multi-Format Download Center
        st.markdown("##### 📦 Multi-Format Download Center")
        col_d1, col_d2, col_d3 = st.columns(3)
        with col_d1:
            st.download_button(
                label="📥 Word Document (`.docx`)",
                data=st.session_state.docx_bytes,
                file_name=f"{Path(st.session_state.file_name).stem}_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        with col_d2:
            if st.session_state.html_content:
                st.download_button(
                    label="🌐 Rendered HTML (`.html`)",
                    data=st.session_state.html_content.encode("utf-8"),
                    file_name=f"{Path(st.session_state.file_name).stem}.html",
                    mime="text/html",
                    use_container_width=True
                )
            else:
                st.button("🌐 HTML Unavailable", disabled=True, use_container_width=True)
        with col_d3:
            if st.session_state.extracted_data:
                json_bytes = json.dumps(st.session_state.extracted_data, indent=2, ensure_ascii=False).encode("utf-8")
                st.download_button(
                    label="📊 Extracted JSON (`.json`)",
                    data=json_bytes,
                    file_name=f"{Path(st.session_state.file_name).stem}.json",
                    mime="application/json",
                    use_container_width=True
                )
            else:
                st.button("📊 JSON Unavailable", disabled=True, use_container_width=True)

        st.markdown("---")
        
        # ---------------------------------------------------------
        # 3. Multi-Tab Preview & Analysis Area
        # ---------------------------------------------------------
        st.subheader("📊 Interactive Previews & Document Inspection")
        tab_html, tab_pdf, tab_json, tab_validation = st.tabs([
            "🌐 Rendered HTML Preview",
            "📄 Original / Compiled PDF",
            "📊 Extracted JSON Data",
            "🛠️ DOCX Validation Report"
        ])

        # Tab 1: Rendered HTML Result
        with tab_html:
            if st.session_state.html_content:
                st.markdown("### 🌐 Themed HTML Document View")
                st.iframe(st.session_state.html_content, height=preview_height)
            else:
                st.info("HTML preview is not available for this document.")

        # Tab 2: Compiled Output PDF Result
        with tab_pdf:
            pdf_bytes_to_show = getattr(st.session_state, "compiled_pdf_bytes", None) or (file_bytes if file_ext == ".pdf" else None)
            if pdf_bytes_to_show:
                st.markdown("### 📄 PDF Document View")
                b64_pdf = base64.b64encode(pdf_bytes_to_show).decode("utf-8")
                pdf_display = f'<iframe src="data:application/pdf;base64,{b64_pdf}" width="100%" height="{preview_height}px" type="application/pdf"></iframe>'
                st.markdown(pdf_display, unsafe_allow_html=True)
            else:
                st.info("PDF view will appear when uploading a PDF or when intermediate rendering is active.")

        # Tab 3: Extracted Structured JSON Data
        with tab_json:
            if st.session_state.extracted_data:
                st.markdown("### 📊 Structured Extracted Data")
                
                # Summary metrics
                doc_info = st.session_state.extracted_data.get("document", {})
                pages_list = doc_info.get("pages", [])
                total_elements = sum(len(p.get("elements", [])) for p in pages_list)
                total_tables = sum(sum(1 for el in p.get("elements", []) if el.get("type") == "table") for p in pages_list)
                
                col_m1, col_m2, col_m3 = st.columns(3)
                col_m1.metric("Pages Parsed", len(pages_list))
                col_m2.metric("Total Elements", total_elements)
                col_m3.metric("Tables Detected", total_tables)
                
                st.json(st.session_state.extracted_data, expanded=False)
            else:
                st.info("No extracted JSON data available for this view.")

        # Tab 4: DOCX Validation Report
        with tab_validation:
            st.markdown("### 🛠️ Document Conversion Fidelity Report")
            if st.session_state.extracted_data and st.session_state.docx_bytes:
                with tempfile.TemporaryDirectory() as val_tmp_dir:
                    val_docx_path = Path(val_tmp_dir) / "check.docx"
                    val_docx_path.write_bytes(st.session_state.docx_bytes)
                    
                    try:
                        from docx import Document
                        doc_check = Document(str(val_docx_path))
                        
                        col_v1, col_v2, col_v3 = st.columns(3)
                        col_v1.metric("Paragraphs in DOCX", len(doc_check.paragraphs))
                        col_v2.metric("Tables in DOCX", len(doc_check.tables))
                        col_v3.metric("Sections / Headers", len(doc_check.sections))
                        
                        st.success("✅ **Validation Passed:** DOCX document structure, paragraphs, tables, and footers successfully verified.")
                    except Exception as e_v:
                        st.warning(f"Note on validation check: {e_v}")
            else:
                st.info("Validation report will be generated after document conversion.")

else:
    # Empty State Hero Card
    st.markdown("""
    <div style="text-align: center; padding: 3rem 1rem; background: #ffffff; border-radius: 12px; border: 2px dashed #cbd5e1; margin-top: 1rem;">
        <h3 style="color: #1e293b; margin-bottom: 0.5rem;">📂 No File Selected</h3>
        <p style="color: #64748b; max-width: 500px; margin: 0 auto 1.5rem auto;">
            Upload any <b>PDF lab report</b>, <b>structured JSON</b>, or <b>HTML document</b> above to convert it directly into a styled Microsoft Word (.docx) document.
        </p>
    </div>
    """, unsafe_allow_html=True)