import docx
from pathlib import Path

docx_path = "output/trace_pipeline_output.docx"

if Path(docx_path).exists():
    doc = docx.Document(docx_path)
    print(f"Total paragraphs in body: {len(doc.paragraphs)}")
    parent_tags = set()
    for idx, p in enumerate(doc.paragraphs):
        parent = p._element.getparent()
        parent_tags.add(parent.tag if parent is not None else "None")
    print(f"Parent tags: {parent_tags}")
else:
    print(f"File not found: {docx_path}")
