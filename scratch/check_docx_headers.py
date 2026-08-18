import docx
from pathlib import Path

docx_path = "output/TestReport_DULAL NAHA (KOL2604070057)_2600128556_11d3f985-7e08-40be-91b4-48b35be363ee_report.docx"

if Path(docx_path).exists():
    doc = docx.Document(docx_path)
    print(f"Total sections: {len(doc.sections)}")
    for s_idx, section in enumerate(doc.sections):
        print(f"\nSection {s_idx+1}:")
        headers = [
            ("default", section.header),
            ("first_page", section.first_page_header),
            ("even_page", section.even_page_header)
        ]
        for h_name, header in headers:
            if header is not None:
                is_linked = header.is_linked_to_previous
                p_count = len(header.paragraphs)
                img_count = 0
                for p in header.paragraphs:
                    for run in p.runs:
                        if "pic:pic" in run._r.xml:
                            img_count += 1
                print(f"  {h_name} header: is_linked={is_linked}, paragraphs={p_count}, images={img_count}")
else:
    print(f"File not found: {docx_path}")
