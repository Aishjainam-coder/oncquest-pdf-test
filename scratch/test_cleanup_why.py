import docx
from pathlib import Path
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
sys.path.append(".")
import fitz
import os
import re

pdf_path = "outsourcing pdf/TestReport_DULAL NAHA (KOL2604070057)_2600128556_11d3f985-7e08-40be-91b4-48b35be363ee.pdf"
docx_path = "output/trace_pipeline_output.docx"

doc_word = docx.Document()

# Let's create a dummy doc and inject headers and footers to see if it happens there
logo_image_path = Path("assets/header_image2.jpeg")
sig_image_path = Path("assets/dr_vinay_signature.png")

for s in doc_word.sections:
    h = s.header
    p_h = h.paragraphs[0]
    p_h.add_run().add_picture(str(logo_image_path.absolute()), width=Inches(1.5))
    
    f = s.footer
    p_f = f.paragraphs[0]
    p_f.add_run().add_picture(str(sig_image_path.absolute()), width=Inches(1.25))

# Add some body paragraphs
p1 = doc_word.add_paragraph("Hello World")
p2 = doc_word.add_paragraph("")
p3 = doc_word.add_paragraph("End")

def get_h_f_images(doc_obj):
    h_imgs = sum(1 for s in doc_obj.sections for p in s.header.paragraphs for r in p.runs if "pic:pic" in r._element.xml)
    f_imgs = sum(1 for s in doc_obj.sections for p in s.footer.paragraphs for r in p.runs if "pic:pic" in r._element.xml)
    return h_imgs, f_imgs

print(f"Before cleanup: header images={get_h_f_images(doc_word)[0]}, footer images={get_h_f_images(doc_word)[1]}")

# Run cleanup
for idx, p in enumerate(list(doc_word.paragraphs)):
    if not p.text.strip() and not any(r._element.xpath('.//w:drawing') for r in p.runs):
        p_elem = p._element
        p_parent = p_elem.getparent()
        if p_parent is not None:
            p_parent.remove(p_elem)
            print(f"Removed paragraph {idx+1}")
            print(f"  After removal: header images={get_h_f_images(doc_word)[0]}, footer images={get_h_f_images(doc_word)[1]}")
            
doc_word.save("output/dummy_test_saved.docx")
doc_check = docx.Document("output/dummy_test_saved.docx")
print(f"After save: header images={get_h_f_images(doc_check)[0]}, footer images={get_h_f_images(doc_check)[1]}")
