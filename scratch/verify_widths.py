import docx
from pathlib import Path

docx_path = Path("output/TestReport_DULAL NAHA (KOL2604070057)_2600128556_11d3f985-7e08-40be-91b4-48b35be363ee.docx")
doc = docx.Document(docx_path)

print(f"Total tables in generated DOCX: {len(doc.tables)}")

for idx, tbl in enumerate(doc.tables):
    rows = tbl.rows
    if not rows:
        continue
    ncols = len(rows[0].cells)
    
    # Safely convert list of cells to text, encoding unknown characters
    header_text = [cell.text.encode('ascii', errors='backslashreplace').decode('ascii') for cell in rows[0].cells]
    
    if ncols > 1:
        print(f"\nTable {idx} (cols: {ncols}, rows: {len(rows)}):")
        print(f"  Headers: {header_text}")
        cell_widths_in = []
        for cell in rows[0].cells:
            w_in = cell.width.inches if cell.width else None
            cell_widths_in.append(f"{w_in:.2f} in" if w_in is not None else "None")
        print(f"  Cell widths: {', '.join(cell_widths_in)}")
        
        if len(rows) > 1:
            first_data_row = [cell.text.encode('ascii', errors='backslashreplace').decode('ascii') for cell in rows[1].cells]
            print(f"  First data row: {first_data_row[:5]}") # Print first few columns
