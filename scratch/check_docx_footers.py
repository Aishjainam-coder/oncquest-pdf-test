import docx
from pathlib import Path

docx_path = "output/TestReport_DULAL NAHA (KOL2604070057)_2600128556_11d3f985-7e08-40be-91b4-48b35be363ee_report.docx"

if Path(docx_path).exists():
    doc = docx.Document(docx_path)
    print(f"Total sections: {len(doc.sections)}")
    for s_idx, section in enumerate(doc.sections):
        print(f"\nSection {s_idx+1}:")
        footers = [
            ("default", section.footer),
            ("first_page", section.first_page_footer),
            ("even_page", section.even_page_footer)
        ]
        for f_name, footer in footers:
            if footer is not None:
                is_linked = footer.is_linked_to_previous
                p_count = len(footer.paragraphs)
                img_count = 0
                for p in footer.paragraphs:
                    # check for runs with pictures
                    for run in p.runs:
                        if "pic:pic" in run._r.xml:
                            img_count += 1
                print(f"  {f_name} footer: is_linked={is_linked}, paragraphs={p_count}, images={img_count}")
else:
    print(f"File not found: {docx_path}")
