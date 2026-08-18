import docx
from pathlib import Path

path = Path("output/debug_test_report_saved.docx")
print(f"File exists: {path.exists()}")
print(f"Size: {path.stat().st_size if path.exists() else 'N/A'} bytes")

if path.exists():
    doc = docx.Document(path)
    print(f"Sections: {len(doc.sections)}")
    for idx, sec in enumerate(doc.sections):
        print(f"Section {idx+1}:")
        print(f"  Header paragraphs: {len(sec.header.paragraphs)}")
        for p_idx, p in enumerate(sec.header.paragraphs):
            h_imgs = sum(1 for r in p.runs if "pic:pic" in r._element.xml)
            print(f"    p {p_idx+1}: text='{p.text}', images={h_imgs}")
            
        print(f"  Footer paragraphs: {len(sec.footer.paragraphs)}")
        for p_idx, p in enumerate(sec.footer.paragraphs):
            f_imgs = sum(1 for r in p.runs if "pic:pic" in r._element.xml)
            print(f"    p {p_idx+1}: text='{p.text}', images={f_imgs}")
