import docx
from pathlib import Path

docx_path = Path(r"c:\Users\aishwarya.jain\Downloads\Oncquest-main\Oncquest-main\output\TestReport_DULAL NAHA (KOL2604070057)_2600128556_11d3f985-7e08-40be-91b4-48b35be363ee_report.docx")

if not docx_path.exists():
    print(f"File not found: {docx_path}")
else:
    doc = docx.Document(docx_path)
    print("=== DOCX Paragraphs containing 'TEST' or 'NAME' ===")
    found = False
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            if "TEST" in text.upper() or "NAME" in text.upper():
                print(f"  Paragraph {i}: {repr(text)}")
                found = True
    
    # Also check tables
    print("\n=== DOCX Table Cells containing 'TEST' or 'NAME' ===")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if "TEST" in text.upper() or "NAME" in text.upper():
                    print(f"  Table {t_idx}, Row {r_idx}, Col {c_idx}: {repr(text)}")
                    found = True
    
    if not found:
        print("No matches found for TEST or NAME in docx!")
