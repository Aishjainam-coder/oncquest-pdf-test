import docx
from pathlib import Path
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
sys.path.append(".")
import fitz
import os
import re

pdf_path = "outsourcing pdf/TestReport_DULAL NAHA (KOL2604070057)_2600128556_11d3f985-7e08-40be-91b4-48b35be363ee.pdf"
docx_path = "output/trace_step_output.docx"

pdf_p = str(Path(pdf_path).absolute())
docx_p = str(Path(docx_path).absolute())

from pdf2docx import Converter as PDF2DocxConverter

# 1. Redact PDF
temp_pdf_path = docx_p + ".temp_redacted.pdf"
doc = fitz.open(pdf_p)
for page_num in range(len(doc)):
    page = doc[page_num]
    rect = fitz.Rect(450, 715, 570, 820)
    page.add_redact_annot(rect, fill=(1, 1, 1))
    page.apply_redactions()
doc.save(temp_pdf_path)
doc.close()

# 2. Convert to docx
cv_obj = PDF2DocxConverter(temp_pdf_path)
cv_obj.convert(docx_p)
cv_obj.close()
os.remove(temp_pdf_path)

doc_word = docx.Document(docx_p)

# Inject signature
sig_image_path = Path("assets/dr_vinay_signature.png")
if sig_image_path.exists():
    for s_idx, section in enumerate(doc_word.sections):
        footers = [section.footer, section.first_page_footer, section.even_page_footer]
        for footer in footers:
            if footer is not None:
                if s_idx == 0 or not footer.is_linked_to_previous:
                    p = footer.paragraphs[0] if len(footer.paragraphs) == 1 and footer.paragraphs[0].text == "" else footer.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.add_run().add_picture(str(sig_image_path.absolute()), width=Inches(1.25))

# Inject logo
logo_image_path = Path("assets/header_image2.jpeg")
if logo_image_path.exists():
    for s_idx, section in enumerate(doc_word.sections):
        headers = [section.header, section.first_page_header, section.even_page_header]
        for header in headers:
            if header is not None:
                if s_idx == 0 or not header.is_linked_to_previous:
                    p = header.paragraphs[0] if len(header.paragraphs) == 1 and header.paragraphs[0].text == "" else header.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.add_run().add_picture(str(logo_image_path.absolute()), width=Inches(1.5))

def get_h_f_images(doc_obj):
    h_imgs = sum(1 for s in doc_obj.sections for p in s.header.paragraphs for r in p.runs if "pic:pic" in r._element.xml)
    f_imgs = sum(1 for s in doc_obj.sections for p in s.footer.paragraphs for r in p.runs if "pic:pic" in r._element.xml)
    return h_imgs, f_imgs

h_init, f_init = get_h_f_images(doc_word)
print(f"Initial: header images={h_init}, footer images={f_init}")

# Run cleanup step-by-step
paragraphs = list(doc_word.paragraphs)
for idx, p in enumerate(paragraphs):
    if not p.text.strip() and not any(r._element.xpath('.//w:drawing') for r in p.runs):
        p_elem = p._element
        p_parent = p_elem.getparent()
        if p_parent is not None:
            p_parent.remove(p_elem)
            h_now, f_now = get_h_f_images(doc_word)
            if h_now < h_init or f_now < f_init:
                print(f"--- DELETION AFFECTED IMAGES at paragraph index {idx} ---")
                print(f"Paragraph text: '{p.text}'")
                print(f"Parent tag: {p_parent.tag}")
                print(f"Paragraph XML: {p_elem.xml[:200]}")
                print(f"New counts: header images={h_now}, footer images={f_now}")
                break
            h_init, f_init = h_now, f_now
