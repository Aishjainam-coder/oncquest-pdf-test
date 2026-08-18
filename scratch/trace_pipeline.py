import docx
from pathlib import Path
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
sys.path.append(".")
import fitz
import os
import re

pdf_path = "outsourcing pdf/TestReport_DULAL NAHA (KOL2604070057)_2600128556_11d3f985-7e08-40be-91b4-48b35be363ee.pdf"
docx_path = "output/trace_pipeline_output.docx"

pdf_p = str(Path(pdf_path).absolute())
docx_p = str(Path(docx_path).absolute())

from pdf2docx import Converter as PDF2DocxConverter

# 1. Redact PDF
temp_pdf_path = docx_p + ".temp_redacted.pdf"
doc = fitz.open(pdf_p)
for page_num in range(len(doc)):
    page = doc[page_num]
    rect = fitz.Rect(450, 715, 570, 820)
    page.add_redact_annot(rect, fill=(1, 1, 1))
    page.apply_redactions()
doc.save(temp_pdf_path)
doc.close()

# 2. Convert to docx
cv_obj = PDF2DocxConverter(temp_pdf_path)
cv_obj.convert(docx_p)
cv_obj.close()
os.remove(temp_pdf_path)

doc_word = docx.Document(docx_p)

def print_status(label):
    print(f"\n--- {label} ---")
    for s_idx, section in enumerate(doc_word.sections):
        headers = [("default", section.header), ("first_page", section.first_page_header), ("even_page", section.even_page_header)]
        footers = [("default", section.footer), ("first_page", section.first_page_footer), ("even_page", section.even_page_footer)]
        for h_name, header in headers:
            h_imgs = sum(1 for p in header.paragraphs for r in p.runs if "pic:pic" in r._element.xml)
            print(f"  Header {h_name}: paragraphs={len(header.paragraphs)}, images={h_imgs}")
        for f_name, footer in footers:
            f_imgs = sum(1 for p in footer.paragraphs for r in p.runs if "pic:pic" in r._element.xml)
            print(f"  Footer {f_name}: paragraphs={len(footer.paragraphs)}, images={f_imgs}")

print_status("After raw conversion")

# 3. Inject signature
sig_image_path = Path("assets/dr_vinay_signature.png")
if sig_image_path.exists():
    for s_idx, section in enumerate(doc_word.sections):
        footers = [section.footer, section.first_page_footer, section.even_page_footer]
        for footer in footers:
            if footer is not None:
                if s_idx == 0 or not footer.is_linked_to_previous:
                    if len(footer.paragraphs) == 1 and footer.paragraphs[0].text == "":
                        p = footer.paragraphs[0]
                    else:
                        p = footer.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run()
                    run.add_picture(str(sig_image_path.absolute()), width=Inches(1.25))

print_status("After signature injection")

# 4. Inject logo
logo_image_path = Path("assets/header_image2.jpeg")
if logo_image_path.exists():
    for s_idx, section in enumerate(doc_word.sections):
        headers = [section.header, section.first_page_header, section.even_page_header]
        for header in headers:
            if header is not None:
                if s_idx == 0 or not header.is_linked_to_previous:
                    if len(header.paragraphs) == 1 and header.paragraphs[0].text == "":
                        p = header.paragraphs[0]
                    else:
                        p = header.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run()
                    run.add_picture(str(logo_image_path.absolute()), width=Inches(1.5))

print_status("After logo injection")

# 5. SNG replace
from converter import replace_sng_in_docx_obj
replace_sng_in_docx_obj(doc_word)

print_status("After SNG replace")

# 6. Spacing cleanup
for p in list(doc_word.paragraphs):
    if not p.text.strip() and not any(r._element.xpath('.//w:drawing') for r in p.runs):
        p_elem = p._element
        p_parent = p_elem.getparent()
        if p_parent is not None:
            p_parent.remove(p_elem)

print_status("After spacing cleanup")

doc_word.save(docx_p)
print_status("After save")
