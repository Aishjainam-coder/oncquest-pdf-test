import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from pdf2docx import Converter
import re

pdf_path = Path("outsourcing pdf/TestReport_DULAL NAHA (KOL2604070057)_2600128556_11d3f985-7e08-40be-91b4-48b35be363ee.pdf").absolute()
out_docx_path = Path("scratch/test_final_output.docx").absolute()

def set_docx_table_black_borders(table, color="000000", sz="6"):
    tblPr = table._tbl.tblPr
    for child in list(tblPr):
        if child.tag.endswith("tblBorders"):
            tblPr.remove(child)
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        b.append(e)
    tblPr.append(b)
    
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for child in list(tcPr):
                if child.tag.endswith("tcBorders"):
                    tcPr.remove(child)
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                e = OxmlElement(f"w:{edge}")
                e.set(qn("w:val"), "single")
                e.set(qn("w:sz"), str(sz))
                e.set(qn("w:space"), "0")
                e.set(qn("w:color"), color)
                tcBorders.append(e)
            tcPr.append(tcBorders)

# 1. Convert PDF directly
cv = Converter(str(pdf_path))
cv.convert(str(out_docx_path))
cv.close()

# 2. Post process docx
doc = docx.Document(str(out_docx_path))

# Apply black border to EVERY table/box
for tbl in doc.tables:
    set_docx_table_black_borders(tbl, color="000000", sz="6")

# Replace test name and SNG
pattern_sng = re.compile(r"SNG\s+Gene?(?:['’‘]|&[a-zA-Z0-9#]+;)?s\s+Lab\s+pvt\.?\s*ltd", re.IGNORECASE)

def fix_paragraph(p):
    if "Liquidseq Actionable Genomic Profiling Panel" in p.text:
        p.text = re.sub(r'Liquidseq\s+Actionable\s+Genomic\s+Profiling\s+Panel', 'TEST NAME', p.text, flags=re.IGNORECASE)
    if "On Illumina Novaseq 6000 Platform" in p.text:
        p.text = re.sub(r'On\s+Illumina\s+Novaseq\s+6000\s+Platform', '', p.text, flags=re.IGNORECASE)
    if pattern_sng.search(p.text):
        p.text = pattern_sng.sub("Laboratory", p.text)

for p in doc.paragraphs:
    fix_paragraph(p)

for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                fix_paragraph(p)

doc.save(str(out_docx_path))
print("Saved post-processed docx. Tables count:", len(doc.tables))
