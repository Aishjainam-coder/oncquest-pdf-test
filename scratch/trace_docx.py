import docx
from pathlib import Path
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

docx_path = "output/debug_test_report.docx"
sig_image_path = Path("assets/dr_vinay_signature.png")

doc_word = docx.Document(docx_path)

def count_images(doc_obj, label):
    print(f"\n[{label}]")
    for s_idx, section in enumerate(doc_obj.sections):
        footers = [
            ("default", section.footer),
            ("first_page", section.first_page_footer),
            ("even_page", section.even_page_footer)
        ]
        for f_name, footer in footers:
            if footer is not None:
                img_count = 0
                for p in footer.paragraphs:
                    for run in p.runs:
                        if "pic:pic" in run._r.xml:
                            img_count += 1
                print(f"  {f_name} footer: paragraphs={len(footer.paragraphs)}, images={img_count}")

# 1. Border styling
for tbl in doc_word.tables:
    pass # Border styling doesn't touch footers

# 2. Perform test name replacement
for p in doc_word.paragraphs:
    if "Liquidseq Actionable Genomic Profiling Panel" in p.text:
        p.text = re.sub(r'Liquidseq\s+Actionable\s+Genomic\s+Profiling\s+Panel', 'TEST NAME', p.text, flags=re.IGNORECASE)

count_images(doc_word, "After test name replacement (body)")

# 3. Inject signature
if sig_image_path.exists():
    for s_idx, section in enumerate(doc_word.sections):
        footers = [section.footer, section.first_page_footer, section.even_page_footer]
        for footer in footers:
            if footer is not None:
                if s_idx == 0 or not footer.is_linked_to_previous:
                    if len(footer.paragraphs) == 1 and footer.paragraphs[0].text == "":
                        p = footer.paragraphs[0]
                    else:
                        p = footer.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run()
                    run.add_picture(str(sig_image_path.absolute()), width=Inches(1.25))

count_images(doc_word, "After signature injection")

# 4. Perform SNG replacement
import sys
sys.path.append(".")
from converter import replace_sng_in_docx_obj
replace_sng_in_docx_obj(doc_word)

count_images(doc_word, "After replace_sng_in_docx_obj")

# 5. Preserve exact PDF vertical layout & spacing
for p in list(doc_word.paragraphs):
    if not p.text.strip() and not any(r._element.xpath('.//w:drawing') for r in p.runs):
        p_elem = p._element
        p_parent = p_elem.getparent()
        if p_parent is not None:
            p_parent.remove(p_elem)

count_images(doc_word, "After layout spacing cleanup")
