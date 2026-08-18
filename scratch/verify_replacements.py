import json
import docx
from pathlib import Path

json_path = Path("extracted_jsons/sng_test_input.json")
html_path = Path("output/sng_test_input_template.html")
docx_path = Path("output/sng_test_input_report.docx")

print("=== VERIFYING JSON OUTPUT ===")
if json_path.exists():
    with open(json_path, "r", encoding="utf-8") as f:
        json_data = json.load(f)
    json_str = json.dumps(json_data)
    
    # Check for SNG
    sng_count = json_str.lower().count("sng")
    lab_count = json_str.lower().count("laboratory")
    print(f"Occurrences of 'SNG' in JSON: {sng_count}")
    print(f"Occurrences of 'Laboratory' in JSON: {lab_count}")
    
    # Let's inspect pages text
    for page in json_data.get("pages", []):
        page_text = page.get("page_text", "")
        if "Laboratory" in page_text:
            print(f"Found 'Laboratory' in page_text: '{page_text.strip()}'")
else:
    print("[ERROR] JSON file not found!")

print("\n=== VERIFYING HTML OUTPUT ===")
if html_path.exists():
    with open(html_path, "r", encoding="utf-8") as f:
        html_content = f.read()
    
    sng_count = html_content.lower().count("sng")
    lab_count = html_content.lower().count("laboratory")
    print(f"Occurrences of 'SNG' in HTML: {sng_count}")
    print(f"Occurrences of 'Laboratory' in HTML: {lab_count}")
    
    # Print lines containing Laboratory
    for line in html_content.split("\n"):
        if "Laboratory" in line:
            print(f"Found line in HTML: {line.strip()[:100]}...")
else:
    print("[ERROR] HTML file not found!")

print("\n=== VERIFYING DOCX OUTPUT ===")
if docx_path.exists():
    doc = docx.Document(docx_path)
    sng_count = 0
    lab_count = 0
    
    # Paragraphs
    for p in doc.paragraphs:
        sng_count += p.text.lower().count("sng")
        if "laboratory" in p.text.lower():
            lab_count += 1
            print(f"Found in DOCX paragraph: '{p.text}'")
            
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    sng_count += p.text.lower().count("sng")
                    if "laboratory" in p.text.lower():
                        lab_count += 1
                        print(f"Found in DOCX table cell: '{p.text}'")
                        
    print(f"Occurrences of 'SNG' in DOCX: {sng_count}")
    print(f"Occurrences of 'Laboratory' in DOCX (paragraph/cell count): {lab_count}")
else:
    print("[ERROR] DOCX file not found!")
