import json
from docx import Document

# Check output docx content
doc = Document('output/TestReport_DULAL NAHA (KOL2604070057)_2600128556_11d3f985-7e08-40be-91b4-48b35be363ee.docx')

print(f"Sections: {len(doc.sections)}")
for i, sec in enumerate(doc.sections):
    from docx.shared import Pt
    print(f"  Section {i+1}: top_margin={sec.top_margin}, bottom_margin={sec.bottom_margin}")

print(f"\nParagraphs: {len(doc.paragraphs)}")
for i, p in enumerate(doc.paragraphs[:30]):
    text = p.text.strip()
    if text:
        print(f"  [{i}] style={p.style.name}, align={p.alignment}, text='{text[:80]}'")

print(f"\nTables: {len(doc.tables)}")
for i, t in enumerate(doc.tables):
    rows = len(t.rows)
    cols = len(t.columns)
    cell0 = t.cell(0,0).text[:50] if rows > 0 and cols > 0 else ''
    print(f"  Table {i+1}: {rows}x{cols}, first_cell='{cell0}'")
