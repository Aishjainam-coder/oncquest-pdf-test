from pdf2docx import Converter
from docx import Document
import fitz

pdf_path = r"outsourcing pdf/TestReport_DULAL NAHA (KOL2604070057)_2600128556_11d3f985-7e08-40be-91b4-48b35be363ee.pdf"
orig_docx = r"scratch/orig_direct.docx"

cv = Converter(pdf_path)
cv.convert(orig_docx)
cv.close()

# Count lines in original docx
doc_orig = Document(orig_docx)
lines = []
for p in doc_orig.paragraphs:
    if p.text.strip():
        lines.append(p.text.strip())
for t in doc_orig.tables:
    for r in t.rows:
        for c in r.cells:
            if c.text.strip():
                lines.append(c.text.strip())

print(f"Direct conversion of Original PDF to DOCX line count: {len(lines)}")
