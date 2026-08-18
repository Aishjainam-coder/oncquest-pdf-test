import docx
from pathlib import Path
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

pdf_path = "outsourcing pdf/TestReport_DULAL NAHA (KOL2604070057)_2600128556_11d3f985-7e08-40be-91b4-48b35be363ee.pdf"
docx_path = "output/debug_test_report.docx"

# Let's inspect converter.py signature injection
import sys
sys.path.append(".")
from converter import convert_pdf_to_word

convert_pdf_to_word(pdf_path, docx_path)

# Check the results
if Path(docx_path).exists():
    doc = docx.Document(docx_path)
    print(f"Total sections: {len(doc.sections)}")
    for s_idx, section in enumerate(doc.sections):
        footers = [
            ("default", section.footer),
            ("first_page", section.first_page_footer),
            ("even_page", section.even_page_footer)
        ]
        for f_name, footer in footers:
            if footer is not None:
                p_texts = [p.text for p in footer.paragraphs]
                img_count = 0
                for p in footer.paragraphs:
                    for run in p.runs:
                        if "pic:pic" in run._r.xml:
                            img_count += 1
                print(f"  {f_name} footer: paragraphs={len(footer.paragraphs)}, p_texts={p_texts}, images={img_count}")
