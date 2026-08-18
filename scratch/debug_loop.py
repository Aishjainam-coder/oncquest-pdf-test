import docx
from pathlib import Path
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

docx_path = "output/TestReport_DULAL NAHA (KOL2604070057)_2600128556_11d3f985-7e08-40be-91b4-48b35be363ee_report.docx"
sig_image_path = Path("assets/dr_vinay_signature.png")

print(f"sig_image_path exists: {sig_image_path.exists()}")
print(f"sig_image_path absolute: {sig_image_path.absolute()}")

doc_word = docx.Document(docx_path)
print(f"Total sections: {len(doc_word.sections)}")

for s_idx, section in enumerate(doc_word.sections):
    print(f"\nSection {s_idx+1}:")
    footers = [
        ("default", section.footer),
        ("first_page", section.first_page_footer),
        ("even_page", section.even_page_footer)
    ]
    for f_name, footer in footers:
        if footer is not None:
            is_linked = footer.is_linked_to_previous
            p_count = len(footer.paragraphs)
            p_text = footer.paragraphs[0].text if p_count > 0 else "N/A"
            print(f"  {f_name} footer: is_linked={is_linked}, paragraphs={p_count}, p_text='{p_text}'")
            
            # Evaluate condition: s_idx == 0 or not footer.is_linked_to_previous
            cond = (s_idx == 0 or not is_linked)
            print(f"    Condition (s_idx == 0 or not is_linked): {cond}")
            if cond:
                if len(footer.paragraphs) == 1 and footer.paragraphs[0].text == "":
                    p = footer.paragraphs[0]
                    print(f"      Reusing paragraph 0")
                else:
                    p = footer.add_paragraph()
                    print(f"      Added new paragraph")
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run()
                print(f"      Added run")
                try:
                    run.add_picture(str(sig_image_path.absolute()), width=Inches(1.25))
                    print(f"      Successfully added picture to run")
                except Exception as e:
                    print(f"      Error adding picture: {e}")

doc_word.save("output/debug_test_report_saved.docx")
print("\nSaved output/debug_test_report_saved.docx")

# Re-read and check images
doc_check = docx.Document("output/debug_test_report_saved.docx")
for s_idx, section in enumerate(doc_check.sections):
    footers = [
        ("default", section.footer),
        ("first_page", section.first_page_footer),
        ("even_page", section.even_page_footer)
    ]
    for f_name, footer in footers:
        if footer is not None:
            img_count = 0
            for p in footer.paragraphs:
                for run in p.runs:
                    if "pic:pic" in run._r.xml:
                        img_count += 1
            print(f"After save - {f_name} footer: images={img_count}")
