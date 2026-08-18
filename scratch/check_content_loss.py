import fitz
import json
from docx import Document
from pathlib import Path

pdf_path = r"outsourcing pdf/TestReport_DULAL NAHA (KOL2604070057)_2600128556_11d3f985-7e08-40be-91b4-48b35be363ee.pdf"
json_path = r"extracted_jsons/TestReport_DULAL NAHA (KOL2604070057)_2600128556_11d3f985-7e08-40be-91b4-48b35be363ee.json"
docx_path = r"output/TestReport_DULAL NAHA (KOL2604070057)_2600128556_11d3f985-7e08-40be-91b4-48b35be363ee_report.docx"

# 1. Read PDF text
doc = fitz.open(pdf_path)
pdf_text_lines = []
for p in doc:
    pdf_text_lines.extend([line.strip() for line in p.get_text("text").splitlines() if line.strip()])

# 2. Read DOCX text
docx_doc = Document(docx_path)
docx_text_lines = []
for p in docx_doc.paragraphs:
    if p.text.strip():
        docx_text_lines.append(p.text.strip())
for t in docx_doc.tables:
    for row in t.rows:
        for cell in row.cells:
            if cell.text.strip():
                docx_text_lines.append(cell.text.strip())

# 3. Read JSON text
with open(json_path, "r", encoding="utf-8") as f:
    json_data = json.load(f)

json_text_str = json.dumps(json_data)

print(f"Total lines in PDF: {len(pdf_text_lines)}")
print(f"Total lines in DOCX: {len(docx_text_lines)}")

# Search for key phrases in PDF that might be missing in DOCX
missing_in_docx = []
for line in pdf_text_lines:
    if len(line) < 4:
        continue
    # check if line or normalized line is in docx
    found = any(line.lower() in d_line.lower() for d_line in docx_text_lines)
    if not found:
        # check if it's in json
        in_json = line.lower() in json_text_str.lower()
        missing_in_docx.append((line, in_json))

print(f"\nTotal lines in PDF missing in DOCX: {len(missing_in_docx)}")
print("\nSample missing lines in DOCX (first 30):")
for line, in_json in missing_in_docx[:30]:
    print(f"  [In JSON: {in_json}] '{line}'")
