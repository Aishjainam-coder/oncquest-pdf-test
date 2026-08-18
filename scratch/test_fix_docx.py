import docx
from pathlib import Path
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

docx_path = "output/TestReport_DULAL NAHA (KOL2604070057)_2600128556_11d3f985-7e08-40be-91b4-48b35be363ee_report.docx"
sig_image_path = "assets/dr_vinay_signature.png"

if Path(docx_path).exists():
    doc = docx.Document(docx_path)
    print(f"Original sections: {len(doc.sections)}")
    
    # Let's clear all existing footers first to reset them, or we can just load the redacted file and run the fix
    # Actually, we can just run the fix on a clean copy. Let's see if we have the redacted temp pdf.
    # Wait, the conversion pipeline creates a redacted temp pdf, and then converts it to docx.
    # Let's write a script that does the post-processing correctly.
    # To test this, let's run the post-processing on doc:
    # First, let's clear the footer paragraphs.
    for s_idx, section in enumerate(doc.sections):
        footers = [
            ("default", section.footer),
            ("first_page", section.first_page_footer),
            ("even_page", section.even_page_footer)
        ]
        for f_name, footer in footers:
            if footer is not None:
                # Clear all paragraphs in the footer
                # Note: to clear a footer, we can delete paragraphs, but docx requires at least one paragraph.
                # So we can delete all but one, and clear the text of that one.
                while len(footer.paragraphs) > 1:
                    p = footer.paragraphs[-1]
                    p._element.getparent().remove(p._element)
                if footer.paragraphs:
                    footer.paragraphs[0].text = ""
                    # Remove all runs
                    for run in footer.paragraphs[0].runs:
                        run._r.getparent().remove(run._r)
    
    # Now let's inject signatures using the fixed logic
    for s_idx, section in enumerate(doc.sections):
        footers = [
            ("default", section.footer),
            ("first_page", section.first_page_footer),
            ("even_page", section.even_page_footer)
        ]
        for f_name, footer in footers:
            if footer is not None:
                # Only add signature if it is the first section or NOT linked to previous
                if s_idx == 0 or not footer.is_linked_to_previous:
                    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                    p.text = ""
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run()
                    run.add_picture(str(sig_image_path), width=Inches(1.2))
                    print(f"Added signature to Section {s_idx+1} {f_name} footer")
                else:
                    print(f"Section {s_idx+1} {f_name} footer is linked, skipping")
                    
    fixed_path = "output/fixed_report.docx"
    doc.save(fixed_path)
    print(f"Saved fixed docx to: {fixed_path}")
    
    # Now verify the saved file
    doc2 = docx.Document(fixed_path)
    for s_idx, section in enumerate(doc2.sections):
        footers = [
            ("default", section.footer),
            ("first_page", section.first_page_footer),
            ("even_page", section.even_page_footer)
        ]
        for f_name, footer in footers:
            if footer is not None:
                img_count = 0
                for p in footer.paragraphs:
                    for run in p.runs:
                        if "pic:pic" in run._r.xml:
                            img_count += 1
                print(f"Section {s_idx+1} {f_name} footer: images={img_count}")
else:
    print(f"File not found: {docx_path}")
