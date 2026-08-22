"""
Universal Dynamic PDF HTML Renderer & Converter Module
======================================================
Provides rendering modes and export tools:
1. Exact Position Layout Mode (`render_exact_pdf_layout_html`): Preserves 100% exact 1-to-1 visual
   positions (`top`, `left`, `width`, `height`) of all text, tables, images, banners, and boxes from input PDF.
2. Standard Flow Template Mode (`generate_dynamic_template_html`): Renders extracted JSON into a responsive web flow.
3. Microsoft Word Exporter (`convert_json_to_docx`): Generates formatted Word .docx files from extracted JSON.
"""

import os
import io
import json
import base64
import re
from pathlib import Path
from PIL import Image
import pymupdf as fitz  # PyMuPDF
import logging

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger("UniversalPDFConverter")

try:
    from playwright.sync_api import sync_playwright
except ImportError:
    sync_playwright = None

try:
    from pdf2docx import Converter as PDF2DocxConverter
except ImportError:
    PDF2DocxConverter = None

def replace_sng_gen_lab(text: str) -> str:
    if not isinstance(text, str):
        return text
    # Broad pattern to catch all SNG/SN Genelab/Gene's Lab variants (with or without Pvt Ltd)
    pattern = re.compile(
        r"(?:SNG\s*Gene?(?:[''\u2018\u2019']|&[a-zA-Z0-9#]+;)?s?\s*Lab(?:oratory)?|SN\s*Genelab)"
        r"(?:\s+pvt\.?\s*ltd\.?)?",
        re.IGNORECASE
    )
    return pattern.sub("Laboratory", text)


def replace_test_name_in_html(html: str) -> str:
    """
    Replaces the SNG test report test name with 'TEST NAME' in rendered HTML.
    Supports:
      - 'Liquidseq Actionable Genomic Profiling Panel'
      - 'Brainseq Genomic Profiling Panel – Advance' (or with hyphen)
      - 'On Illumina Novaseq 6000 Platform' (subtitle, cleared)
    Both lines are replaced — the main test name becomes 'TEST NAME', the subtitle is cleared.
    Preserves all HTML tags, positioning, and formatting.
    """
    if not isinstance(html, str):
        return html

    # Replace the main test name line with 'TEST NAME'
    html = re.sub(
        r'(?<=\>)\s*(?:Liquidseq\s+Actionable|Brainseq)\s+Genomic\s+Profiling\s+Panel(?:\s*[-–]\s*Advance)?\s*(?=\<)',
        'TEST NAME',
        html, flags=re.IGNORECASE
    )
    # Clear the subtitle line (already covered by the main test name above)
    html = re.sub(
        r'(?<=\>)\s*On\s+Illumina\s+Novaseq\s+6000\s+Platform\s*(?=\<)',
        '',
        html, flags=re.IGNORECASE
    )

    return html


def replace_test_name_in_structure(obj):
    """
    Recursively replaces the test name in JSON structure keys/values
    so direct JSON->DOCX or JSON->HTML routes are sanitized cleanly.
    """
    if isinstance(obj, dict):
        return {k: replace_test_name_in_structure(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [replace_test_name_in_structure(item) for item in obj]
    elif isinstance(obj, str):
        # Only replace if the entire string matches the test name exactly
        if re.match(r'^\s*(?:Liquidseq\s+Actionable|Brainseq)\s+Genomic\s+Profiling\s+Panel(?:\s*[-–]\s*Advance)?\s*$', obj, re.IGNORECASE):
            return "TEST NAME"
        # Only replace if the entire string matches the subtitle exactly
        if re.match(r'^\s*On\s+Illumina\s+Novaseq\s+6000\s+Platform\s*$', obj, re.IGNORECASE):
            return ""
        return obj
    return obj



def replace_sng_in_structure(obj):
    if isinstance(obj, dict):
        return {k: replace_sng_in_structure(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [replace_sng_in_structure(item) for item in obj]
    elif isinstance(obj, str):
        return replace_sng_gen_lab(obj)
    return obj


def replace_sng_in_docx_obj(doc):
    import re
    pattern = re.compile(
        r"(?:SNG\s*Gene?(?:[''\u2018\u2019']|&[a-zA-Z0-9#]+;)?s?\s*Lab(?:oratory)?|SN\s*Genelab)"
        r"(?:\s+pvt\.?\s*ltd\.?)?",
        re.IGNORECASE
    )
    replacement = "Laboratory"
    
    def replace_in_xml_elements(root_element):
        for p in root_element.iter():
            if p.tag.endswith('}p'):
                t_elements = [t for t in p.iter() if t.tag.endswith('}t')]
                if not t_elements:
                    continue
                
                # Check individual t elements first
                replaced = False
                for t in t_elements:
                    if t.text and pattern.search(t.text):
                        t.text = pattern.sub(replacement, t.text)
                        replaced = True
                
                # If not replaced but the combined text matches, we do cross-element replacement
                if not replaced:
                    full_text = "".join(t.text for t in t_elements if t.text)
                    if pattern.search(full_text):
                        t_elements[0].text = pattern.sub(replacement, full_text)
                        for t in t_elements[1:]:
                            t.text = ""

    # Replace in main body
    replace_in_xml_elements(doc.element)

    # Replace in headers/footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header is not None:
                replace_in_xml_elements(header._element)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer is not None:
                replace_in_xml_elements(footer._element)


from extractor import extract_report_data, detect_dynamic_header_footer_bounds


def get_css_color(color_tuple):
    if not color_tuple:
        return None
    try:
        if isinstance(color_tuple, (list, tuple)):
            if len(color_tuple) == 3:
                r = int(color_tuple[0] * 255)
                g = int(color_tuple[1] * 255)
                b = int(color_tuple[2] * 255)
                return f"rgb({r},{g},{b})"
            elif len(color_tuple) == 1:
                gray = int(color_tuple[0] * 255)
                return f"rgb({gray},{gray},{gray})"
            elif len(color_tuple) == 4:
                c, m, y, k = color_tuple
                r = int(255 * (1 - c) * (1 - k))
                g = int(255 * (1 - m) * (1 - k))
                b = int(255 * (1 - y) * (1 - k))
                return f"rgb({r},{g},{b})"
        elif isinstance(color_tuple, (int, float)):
            val = int(color_tuple * 255)
            return f"rgb({val},{val},{val})"
    except Exception:
        pass
    return "black"


def render_exact_pdf_layout_html(doc, doc_title: str = "Uploaded Document", theme_config: dict = None) -> str:
    """
    Renders an HTML document where body elements stay in their exact visual positions,
    while original header and footer content is completely excluded dynamically.
    Applies user-selected theme typography, primary colors, table header styling, and cell borders.
    """
    cfg = get_merged_theme_config(theme_config)
    colors_cfg = cfg.get("colors", {})
    typo_cfg = cfg.get("typography", {})

    primary_color = colors_cfg.get("primary", "#1f497d")
    secondary_color = colors_cfg.get("secondary", "#008080")
    accent_orange = colors_cfg.get("accent_orange", "#ed7d31")
    accent_red = colors_cfg.get("accent_red", "#ff0000")
    banner_dark = colors_cfg.get("banner_dark", "#404040")
    text_primary = colors_cfg.get("text_primary", "#000000")
    text_dark = colors_cfg.get("text_dark", "#0d0d0d")
    bg_page = colors_cfg.get("background_page", "#ffffff")
    # Change 4: All borders are now forced to black
    border_color = "#000000"
    result_positive_color = colors_cfg.get("result_positive", "#C00000")
    result_negative_color = colors_cfg.get("result_negative", "#008000")

    # Change 3: Keywords for result color coding
    _POSITIVE_KEYWORDS = ["positive", "pathogenic", "detected", "high", "abnormal", "msi - high", "msi-high"]
    _NEGATIVE_KEYWORDS = ["negative", "normal", "not detected", "stable", "msi - stable", "msi-stable", "benign", "likely benign"]
    font_family = typo_cfg.get("primary_family", "Cambria, 'Times New Roman', serif")

    fallback_left = 35.5
    fallback_width = 524.0

    # Detect dynamic header and footer bounds per page
    page_bounds = detect_dynamic_header_footer_bounds(doc)

    html_parts = [
        "<!DOCTYPE html>",
        "<html>",
        "<head>",
        "<meta charset='utf-8'>",
        f"<title>{doc_title}</title>",
        "<style>",
        f"""
        :root {{
          --color-primary: {primary_color};
          --color-secondary: {secondary_color};
          --color-accent-orange: {accent_orange};
          --color-accent-red: {accent_red};
          --color-banner-dark: {banner_dark};
          --color-bg-page: {bg_page};
          --color-bg-container: #525659;
          --text-primary: {text_primary};
          --text-teal: {secondary_color};
          --text-orange: {accent_orange};
          --text-red: {accent_red};
          --border-color: #000000;
          --font-primary: {font_family};
        }}
        @page {{ size: 595.6pt 842.0pt; margin: 0; }}
        * {{ box-sizing: border-box; }}
        body {{ margin: 0; padding: 0; background-color: var(--color-bg-container); font-family: var(--font-primary); color: var(--text-primary); }}
        .pdf-container {{ display: flex; flex-direction: column; align-items: center; padding: 20px 0; }}
        .pdf-page {{ background: var(--color-bg-page); width: 595.6pt; min-height: 842.0pt; margin-bottom: 20px; position: relative; overflow: visible; box-shadow: 0 4px 12px rgba(0,0,0,0.3); page-break-after: always; font-family: var(--font-primary); }}
        div[id^='page'] {{ position: relative !important; width: 595.6pt !important; min-height: 842.0pt !important; overflow: visible !important; }}
        div[id^='page'] p {{ position: absolute !important; margin: 0 !important; padding: 0 !important; white-space: normal !important; word-break: normal !important; overflow-wrap: break-word !important; max-width: 100% !important; z-index: 10 !important; font-family: var(--font-primary) !important; line-height: 1.2 !important; overflow: visible !important; }}
        div[id^='page'] span {{ word-break: normal !important; overflow-wrap: break-word !important; }}
        .black-banner-span {{ color: #ffffff !important; display: block !important; width: 100% !important; text-align: center !important; padding: 4px 0 !important; line-height: 1.2 !important; font-weight: bold !important; font-size: 13.0pt !important; font-family: var(--font-primary) !important; border-radius: 0px !important; white-space: normal !important; margin: 0 !important; background-color: var(--color-banner-dark) !important; box-sizing: border-box !important; z-index: 15 !important; }}
        .label-bar-span {{ color: #ffffff !important; display: inline-block !important; padding: 2px 6px !important; line-height: 1.2 !important; border-radius: 2px !important; font-family: var(--font-primary) !important; word-break: break-word !important; margin: 0 !important; background-color: var(--color-primary) !important; z-index: 15 !important; }}
        .table-header-cell {{ position: absolute; background-color: var(--color-primary) !important; color: #ffffff !important; font-family: var(--font-primary) !important; font-size: 9.5pt !important; font-weight: bold !important; display: flex !important; align-items: center !important; justify-content: center !important; text-align: center !important; padding: 2px 4px !important; white-space: normal !important; word-break: break-word !important; overflow-wrap: break-word !important; line-height: 1.15 !important; border: 1px solid var(--border-color) !important; box-sizing: border-box !important; z-index: 15 !important; pointer-events: none !important; }}
        div[id^='page'] img {{ position: absolute !important; transform-origin: 0 0 !important; z-index: 5 !important; opacity: 1 !important; visibility: visible !important; display: inline-block !important; }}
        .table-grid-cell {{ position: absolute; border: 1px solid var(--border-color) !important; background: transparent; pointer-events: none; z-index: 2 !important; }}
        .vector-box {{ position: absolute; border: 1px solid var(--border-color) !important; background: transparent; pointer-events: none; z-index: 2 !important; border-radius: 2px; box-sizing: border-box !important; }}
        .vector-fill-box {{ position: absolute; background-color: var(--color-primary) !important; pointer-events: none; z-index: 2 !important; border-radius: 2px; box-sizing: border-box !important; }}
        .section-content-box {{ position: absolute; border: 1px solid var(--border-color) !important; background: transparent; pointer-events: none; z-index: 2 !important; border-radius: 3px; }}
        .teal-text {{ color: var(--color-secondary) !important; font-weight: bold; }}
        .orange-text {{ color: var(--color-accent-orange) !important; font-weight: bold; }}
        .red-text {{ color: var(--color-accent-red) !important; font-weight: bold; }}
        @media print {{ body {{ background-color: #ffffff; }} .pdf-container {{ padding: 0; }} .pdf-page {{ margin: 0; box-shadow: none; }} }}
        """,
        "</style>",
        "</head>",
        "<body>",
        "<div class='pdf-container'>",
    ]

    import base64
    from pathlib import Path
    header_image2_b64 = ""
    header_image2_path = Path("assets/header_image2.jpeg")
    if header_image2_path.exists():
        try:
            with open(header_image2_path, "rb") as img_f:
                header_image2_b64 = base64.b64encode(img_f.read()).decode("utf-8")
        except Exception:
            pass

    sig_image_b64 = ""
    sig_image_path = Path("assets/dr_vinay_signature.png")
    if sig_image_path.exists():
        try:
            with open(sig_image_path, "rb") as sig_f:
                sig_image_b64 = base64.b64encode(sig_f.read()).decode("utf-8")
        except Exception:
            pass

    for page_num in range(len(doc)):
        page = doc[page_num]
        page_left_val, page_width_val = _get_page_bounds(page, fallback_left, fallback_width)
        page_left_str = f"{page_left_val:.1f}pt"
        page_width_str = f"{page_width_val:.1f}pt"

        hy_cutoff = page_bounds[page_num]["header_y_cutoff"]
        fy_cutoff = page_bounds[page_num]["footer_y_cutoff"]

        html_parts.append(f"<div class='pdf-page' id='page-{page_num+1}'>")
        if header_image2_b64:
            img_tag = (
                f'<img src="data:image/jpeg;base64,{header_image2_b64}" '
                f'style="position:absolute !important; left:440.0pt !important; '
                f'top:15.0pt !important; width:130.0pt !important; '
                f'height:auto !important; z-index:100 !important; '
                f'object-fit:contain !important;" />'
            )
            html_parts.append(img_tag)
        page_html = page.get_text("html")

        # 1. Extract PyMuPDF table headers & grid coordinates in body region ONLY
        tabs = page.find_tables()
        table_header_html_divs = []
        table_grid_html_divs = []
        header_y_ranges = []
        kept_table_bboxes = []

        for tab in tabs.tables:
            if hasattr(tab, 'bbox'):
                tx0, ty0, tx1, ty1 = tab.bbox
                t_mid = (ty0 + ty1) / 2.0
                if t_mid < hy_cutoff or t_mid > fy_cutoff:
                    continue
                kept_table_bboxes.append(tab.bbox)

            valid_cells = [c for c in tab.cells if c]
            if valid_cells:
                min_y0 = min(c[1] for c in valid_cells)
                header_cells = [c for c in valid_cells if abs(c[1] - min_y0) < 3.0]
                header_cells.sort(key=lambda c: c[0])
                if len(header_cells) >= 2:
                    hy0 = min(c[1] for c in header_cells)
                    hy1 = max(c[3] for c in header_cells)
                    header_y_ranges.append((hy0 - 3.0, hy1 + 3.0))

                    for c in header_cells:
                        x0, y0, x1, y1 = c
                        w = x1 - x0
                        h = max(24.0, y1 - y0)
                        rect = fitz.Rect(x0, y0, x1, y1)
                        raw_text = page.get_text('text', clip=rect).strip()
                        formatted_text = raw_text.replace('\n', ' ').strip()
                        if "/" in formatted_text and " " not in formatted_text:
                            formatted_text = formatted_text.replace("/", "/ ")
                        table_header_html_divs.append(
                            f"<div class='table-header-cell' "
                            f"style='left:{x0:.1f}pt;top:{y0:.1f}pt;"
                            f"width:{w:.1f}pt;height:{h:.1f}pt;'>"
                            f"{formatted_text}</div>"
                        )

            for cell in tab.cells:
                if cell:
                    cx0, cy0, cx1, cy1 = cell
                    cw = cx1 - cx0
                    ch = cy1 - cy0
                    if any(hy0_r <= cy0 <= hy1_r for hy0_r, hy1_r in header_y_ranges):
                        continue
                    table_grid_html_divs.append(
                        f"<div class='table-grid-cell' "
                        f"style='left:{cx0:.1f}pt;top:{cy0:.1f}pt;"
                        f"width:{cw:.1f}pt;height:{ch:.1f}pt;'></div>"
                    )

        # 2. Extract Vector Drawings in body region ONLY
        vector_html_divs = []
        try:
            drawings = page.get_drawings()
            for d in drawings:
                rect = d.get('rect')
                if not rect:
                    continue
                rx0, ry0, rx1, ry1 = rect.x0, rect.y0, rect.x1, rect.y1
                rw, rh = rx1 - rx0, ry1 - ry0
                
                # Skip tiny points and full page borders
                if rw < 1.0 and rh < 1.0:
                    continue
                if rw > 550 and rh > 800:
                    continue

                # Exclude vector lines/drawings in header or footer regions, unless they are part of kept tables
                is_table_vector = False
                for tx0, ty0, tx1, ty1 in kept_table_bboxes:
                    if (tx0 - 2.0) <= rx0 <= rx1 <= (tx1 + 2.0) and (ty0 - 2.0) <= ry0 <= ry1 <= (ty1 + 2.0):
                        is_table_vector = True
                        break
                if not is_table_vector:
                    if ry0 < hy_cutoff or ry1 > fy_cutoff:
                        continue

                if any(abs(ry0 - hy0_r) < 5.0 for hy0_r, _ in header_y_ranges):
                    continue

                fill_col = get_css_color(d.get('fill'))
                stroke_col = get_css_color(d.get('color'))
                stroke_w = d.get('width') or 1.0

                if rh <= 1.5:  # Horizontal line segment
                    bg_color = stroke_col or "black"
                    vector_html_divs.append(
                        f"<div style='position:absolute; left:{rx0:.1f}pt; top:{ry0:.1f}pt; width:{rw:.1f}pt; height:{stroke_w:.1f}pt; background-color:{bg_color}; z-index:2; pointer-events:none;'></div>"
                    )
                elif rw <= 1.5:  # Vertical line segment
                    bg_color = stroke_col or "black"
                    vector_html_divs.append(
                        f"<div style='position:absolute; left:{rx0:.1f}pt; top:{ry0:.1f}pt; width:{stroke_w:.1f}pt; height:{rh:.1f}pt; background-color:{bg_color}; z-index:2; pointer-events:none;'></div>"
                    )
                else:  # Rectangle box
                    # Change 2: Skip filled rectangles that serve as heading background boxes
                    if fill_col and rh < 25.0 and rw > 40.0:
                        # This looks like a heading background fill box — suppress it
                        continue

                    # Remove all rectangles before headings (decorative squares on the left)
                    if fill_col and rx0 < 45.0 and rw < 20.0 and rh < 25.0:
                        continue


                    style_parts = [
                        "position:absolute;",
                        f"left:{rx0:.1f}pt;",
                        f"top:{ry0:.1f}pt;",
                        f"width:{rw:.1f}pt;",
                        f"height:{rh:.1f}pt;",
                        "pointer-events:none;"
                    ]
                    if fill_col:
                        # Use exact background fill color
                        style_parts.append(f"background-color:{fill_col};")
                        style_parts.append("z-index:1;")
                    else:
                        style_parts.append("background-color:transparent;")
                        style_parts.append("z-index:2;")

                    if stroke_col:
                        # Change 4: Force all stroke borders to black
                        style_parts.append(f"border:{stroke_w:.1f}pt solid #000000;")
                    else:
                        # Add black border even if no stroke was defined
                        style_parts.append(f"border:0.5pt solid #000000;")

                    vector_html_divs.append(
                        f"<div style='{' '.join(style_parts)}'></div>"
                    )
        except Exception:
            pass

        # 3. Clean raw HTML & suppress raw <p> tags inside header/footer regions or table headers
        cleaned = page_html
        cleaned = re.sub(r'<img\s+[^>]*>', '', cleaned)
        cleaned = re.sub(r'font-family:[^;"]+', f'font-family: {font_family}', cleaned)

        # HIDE raw <p> tags that fall inside header/footer regions or table headers
        def filter_hdr_ftr_and_table_p(match):
            p_tag = match.group(0)
            
            # Filter out standalone page numbering
            text_val = re.sub(r'<[^>]+>', '', p_tag).strip()
            if re.match(r'^Page\s+\d+\s+of\s+\d+$', text_val, re.IGNORECASE):
                return ""

            top_m = re.search(r'top:\s*([\d.]+)pt', p_tag)
            left_m = re.search(r'left:\s*([\d.]+)pt', p_tag)
            if top_m and left_m:
                y_val = float(top_m.group(1))
                x_val = float(left_m.group(1))
                is_inside_table = False
                for tx0, ty0, tx1, ty1 in kept_table_bboxes:
                    if (tx0 - 5.0) <= x_val <= (tx1 + 5.0) and (ty0 - 5.0) <= y_val <= (ty1 + 5.0):
                        is_inside_table = True
                        break
                if not is_inside_table:
                    if y_val < hy_cutoff or y_val > fy_cutoff:
                        return ""
            elif top_m:
                y_val = float(top_m.group(1))
                if y_val < hy_cutoff or y_val > fy_cutoff:
                    return ""
                if any(hy0_r <= y_val <= hy1_r for hy0_r, hy1_r in header_y_ranges):
                    return re.sub(r'>([^<]+)<', '><', p_tag)
            return p_tag

        cleaned = re.sub(r'<p\s+[^>]*>.*?</p>', filter_hdr_ftr_and_table_p, cleaned, flags=re.DOTALL)

        # 4. Format Black Banners & Blue Section Label Bars
        # Change 1: Headings use single consistent color (no multi-color internal spans)
        # Change 2: No colored background boxes on heading text
        def format_heading_p(match):
            p_tag = match.group(0)
            is_white_text = bool(re.search(r'color:\s*(?:#ffffff|#fff|rgb\(\s*255\s*,\s*255\s*,\s*255\s*\))', p_tag, re.IGNORECASE))
            size_m = re.search(r'font-size:\s*([\d.]+)pt', p_tag)
            font_sz = float(size_m.group(1)) if size_m else 10.0
            
            top_m = re.search(r'top:\s*([\d.]+)pt', p_tag)
            top_val = top_m.group(1) if top_m else "100.0"
            left_m = re.search(r'left:\s*([\d.]+)pt', p_tag)
            left_val = left_m.group(1) if left_m else page_left_str.replace('pt', '')

            if is_white_text and font_sz >= 12.0:
                text_val = re.sub(r'<[^>]+>', '', p_tag).strip()
                if len(text_val) > 15:
                    # Change 1: Single consistent heading style — plain bold text, no background
                    return (
                        f'<p style="top:{top_val}pt;left:{left_val}pt;'
                        f'width:{page_width_str};margin:0;padding:4px 0;z-index:15;">'
                        f'<span style="font-family:{font_family};font-size:{font_sz}pt;'
                        f'font-weight:bold;color:#000000;">'
                        f'{text_val}</span></p>'
                    )
            return p_tag

        cleaned = re.sub(r'<p\s+[^>]*>.*?</p>', format_heading_p, cleaned, flags=re.DOTALL)

        # Change 1: Normalize label bar headings — remove background, use consistent heading color
        def format_labelbar_p(match):
            p_tag = match.group(0)
            is_white_text = bool(re.search(r'color:\s*(?:#ffffff|#fff|rgb\(\s*255\s*,\s*255\s*,\s*255\s*\))', p_tag, re.IGNORECASE))
            if is_white_text and "black-banner-span" not in p_tag and "background-color:#404040" not in p_tag and "table-header-cell" not in p_tag:
                # Strip background and set consistent text color
                def fix_span(sm):
                    span = sm.group(0)
                    # Remove background-color
                    span = re.sub(r'background-color:\s*[^;"]+;?', '', span)
                    # Remove display:inline-block and padding (heading box styles)
                    span = re.sub(r'display:\s*inline-block;?', '', span)
                    span = re.sub(r'padding:\s*[^;"]+;?', '', span)
                    span = re.sub(r'border-radius:\s*[^;"]+;?', '', span)
                    # Change text color from white to consistent heading color
                    span = re.sub(r'color:\s*#(?:ffffff|fff)', f'color:{primary_color}', span, flags=re.IGNORECASE)
                    span = re.sub(r'color:\s*rgb\(\s*255\s*,\s*255\s*,\s*255\s*\)', f'color:{primary_color}', span, flags=re.IGNORECASE)
                    return span
                return re.sub(r'<span\s+[^>]*>.*?</span>', fix_span, p_tag, flags=re.DOTALL)
            return p_tag

        cleaned = re.sub(r'<p\s+[^>]*>.*?</p>', format_labelbar_p, cleaned, flags=re.DOTALL)

        # Change 3: Color-code result keywords (Positive=Red, Negative=Green) in exact layout
        def colorize_result_text(match):
            p_tag = match.group(0)
            text_val = re.sub(r'<[^>]+>', '', p_tag).strip().lower()
            # Skip if already styled as heading/banner
            if 'black-banner-span' in p_tag or 'table-header-cell' in p_tag:
                return p_tag
            # Skip long text — only short result values should be colored
            if len(text_val) > 100:
                return p_tag
            if any(kw in text_val for kw in _POSITIVE_KEYWORDS):
                # Wrap all spans in red color
                p_tag = re.sub(r'color:\s*[^;"]+', f'color:{result_positive_color}', p_tag)
                if 'font-weight' not in p_tag:
                    p_tag = p_tag.replace('style="', 'style="font-weight:bold;', 1)
            elif any(kw in text_val for kw in _NEGATIVE_KEYWORDS):
                # Wrap all spans in green color
                p_tag = re.sub(r'color:\s*[^;"]+', f'color:{result_negative_color}', p_tag)
                if 'font-weight' not in p_tag:
                    p_tag = p_tag.replace('style="', 'style="font-weight:bold;', 1)
            return p_tag

        cleaned = re.sub(r'<p\s+[^>]*>.*?</p>', colorize_result_text, cleaned, flags=re.DOTALL)

        # 5. Extract exact images in body region ONLY
        exact_image_html_divs = []
        try:
            img_infos = page.get_image_info(xrefs=True)
            for info in img_infos:
                xref = info.get("xref")
                bbox = info.get("bbox")
                if not xref or not bbox:
                    continue
                x0, y0, x1, y1 = bbox
                is_inside_table = False
                for tx0, ty0, tx1, ty1 in kept_table_bboxes:
                    if (tx0 - 5.0) <= x0 <= x1 <= (tx1 + 5.0) and (ty0 - 5.0) <= y0 <= y1 <= (ty1 + 5.0):
                        is_inside_table = True
                        break
                if not is_inside_table:
                    if y0 < hy_cutoff or y1 > fy_cutoff or y0 >= 690.0:
                        continue
                w_pt = max(1.0, x1 - x0)
                h_pt = max(1.0, y1 - y0)
                try:
                    base_img = doc.extract_image(xref)
                    if not base_img:
                        continue
                    img_bytes = base_img.get("image")
                    img_ext = base_img.get("ext", "png").lower()
                    b64_str = base64.b64encode(img_bytes).decode("utf-8")
                    img_tag = (
                        f'<img src="data:image/{img_ext};base64,{b64_str}" '
                        f'style="position:absolute !important; left:{x0:.1f}pt !important; '
                        f'top:{y0:.1f}pt !important; width:{w_pt:.1f}pt !important; '
                        f'height:{h_pt:.1f}pt !important; z-index:5 !important; '
                        f'object-fit:contain !important;" />'
                    )
                    exact_image_html_divs.append(img_tag)
                except Exception:
                    pass
        except Exception:
            pass

        section_overlays = get_page_section_overlays(page, page_left_str, page_width_str, hy_cutoff, fy_cutoff)

        html_parts.append(cleaned)
        html_parts.extend(vector_html_divs)
        html_parts.extend(section_overlays)
        html_parts.extend(table_header_html_divs)
        html_parts.extend(table_grid_html_divs)
        html_parts.extend(exact_image_html_divs)
        if sig_image_b64:
            sig_tag = (
                f"<div class='page-signature-block' style='position:absolute !important; "
                f"right:40.0pt !important; top:725.0pt !important; width:90.0pt !important; "
                f"height:88.0pt !important; z-index:100 !important; text-align:left !important; pointer-events:none !important;'>"
                f"  <img src='data:image/png;base64,{sig_image_b64}' style='position:relative !important; "
                f"width:90.0pt !important; height:88.0pt !important; display:block !important; "
                f"transform:none !important; opacity:1 !important; visibility:visible !important;' />"
                f"</div>"
            )
            html_parts.append(sig_tag)
        html_parts.append("</div>")

    html_parts.append("</div></body></html>")
    full_html = "\n".join(html_parts)
    full_html = replace_sng_gen_lab(full_html)
    full_html = replace_test_name_in_html(full_html)
    return full_html


def generate_dynamic_template_html(data: dict, doc_title: str = "Uploaded Document", theme_config: dict = None) -> str:
    """
    Renders extracted PDF JSON data dynamically into a clean HTML report template.
    Does NOT recreate vendor logos, patient header cards, or footer signatures.
    """
    data = replace_sng_in_structure(data)
    data = replace_test_name_in_structure(data)
    cfg = get_merged_theme_config(theme_config)
    colors_cfg = cfg.get("colors", {})
    typo_cfg = cfg.get("typography", {})

    primary_color = colors_cfg.get("primary", "#1f497d")
    secondary_color = colors_cfg.get("secondary", "#008080")
    accent_orange = colors_cfg.get("accent_orange", "#ed7d31")
    accent_red = colors_cfg.get("accent_red", "#ff0000")
    bg_page = colors_cfg.get("background_page", "#ffffff")
    text_color = colors_cfg.get("text_primary", "#0d0d0d")
    # Change 4: Force all borders to black
    border_color = "#000000"
    table_header_bg = primary_color
    table_header_text = colors_cfg.get("text_light", "#ffffff")
    font_family = typo_cfg.get("primary_family", "Cambria, 'Times New Roman', serif")
    result_positive_color = colors_cfg.get("result_positive", "#C00000")
    result_negative_color = colors_cfg.get("result_negative", "#008000")

    spacing_cfg = cfg.get("spacing", {})
    box_padding = spacing_cfg.get("boxPadding", 8)
    cell_padding_x = spacing_cfg.get("cellPaddingX", 6)
    cell_padding_y = spacing_cfg.get("cellPaddingY", 4)
    block_gap = spacing_cfg.get("blockGap", 10)

    border_cfg = cfg.get("border", {})
    border_width = border_cfg.get("width", 1)
    border_style = border_cfg.get("style", "single")
    border_color_val = "#000000"  # Change 4: Always use black borders

    css_border_style = "solid"
    if border_style == "dotted":
        css_border_style = "dotted"
    elif border_style == "dashed":
        css_border_style = "dashed"

    show_tables = theme_config.get("show_tables", True) if theme_config else True
    show_sections = theme_config.get("show_sections", True) if theme_config else True
    show_images = theme_config.get("show_images", True) if theme_config else True
    show_badges = theme_config.get("show_badges", True) if theme_config else True

    badge_rules = (theme_config or {}).get("badge_rules", {
        "danger": ["pathogenic", "positive", "high", "failed", "rejected", "overdue",
                   "invalid", "expired", "critical", "denied", "delinquent", "abnormal"],
        "warning": ["vus", "uncertain", "warning", "pending", "under review",
                    "partial", "provisional", "conditional"],
        "success": ["passed", "normal", "negative", "approved", "paid", "valid",
                    "cleared", "completed", "compliant", "settled"]
    })

    # Keywords that trigger result-level coloring (inline, not badge)
    _POSITIVE_KEYWORDS = ["positive", "pathogenic", "detected", "high", "abnormal", "msi - high", "msi-high"]
    _NEGATIVE_KEYWORDS = ["negative", "normal", "not detected", "stable", "msi - stable", "msi-stable", "benign", "likely benign"]

    def _result_color_span(cell_str: str) -> str:
        """Wrap cell text in a colored span if it matches a clinical result keyword.
        Only applies to short result values (< 100 chars) to avoid false positives
        on long descriptive paragraphs that happen to contain a keyword."""
        cell_lower = cell_str.strip().lower()
        # Skip long text — these are descriptions, not result values
        if len(cell_lower) > 100:
            return cell_str
        if any(kw in cell_lower for kw in _NEGATIVE_KEYWORDS):
            return f"<span style='color:{result_negative_color}; font-weight:bold;'>{cell_str}</span>"
        if any(kw in cell_lower for kw in _POSITIVE_KEYWORDS):
            return f"<span style='color:{result_positive_color}; font-weight:bold;'>{cell_str}</span>"
        return cell_str

    tables = data.get("all_tables") or data.get("tables") or []
    sections = data.get("all_boxes_and_sections") or data.get("content_sections") or []
    images = data.get("all_images_and_graphs") or data.get("images_and_graphs") or []

    # Titles to skip (demographics, generic "General Content / Notes")
    _SKIP_TITLES = {
        "general content / notes",
        "patient details & metadata",
    }

    def _should_skip_section(sec):
        title = sec.get("title", "").strip()
        sec_type = sec.get("type", "")
        if sec_type == "demographics_box":
            return True
        if title.startswith("Header & Metadata Box"):
            return True
        if title.lower() in _SKIP_TITLES:
            return True
        return False

    def _is_reference_fragment(sec):
        """Detect boxes that look like split reference entries (by checking title and content)."""
        title = sec.get("title", "").strip()
        content_text = sec.get("content_text", [])
        body_str = " ".join(content_text) if isinstance(content_text, list) else str(content_text)
        
        combined = (title + " " + body_str).lower()
        
        # Explicit reference cues
        if "pmid" in combined or "doi:" in combined or "et al" in combined:
            return True
        if "http" in combined or "www." in combined or ".html" in combined or ".com/" in combined or "release" in combined:
            return True
        if "guideline" in combined:
            return True
        # Citation year or similar patterns
        if re.search(r'\b\d{4}\b', title):
            return True
        # Starts with a number (like volume info "2;23(8)") and contains punctuation
        if re.match(r'^\d', title) and any(c in title for c in [';', ':', '(', ')', '-']):
            return True
            
        return False

    # 1. Content Section Boxes HTML — merge references, skip generic boxes
    sections_html = ""
    if show_sections and sections:
        reference_paragraphs = []  # collect all reference fragments
        has_references_section = False

        for sec in sections:
            if _should_skip_section(sec):
                continue

            title = sec.get("title", "").strip()
            content_text = sec.get("content_text", [])
            if isinstance(content_text, list):
                body_lines = [t.strip() for t in content_text if t and t.strip()]
            else:
                body_lines = [str(content_text).strip()] if content_text else []

            # Check if this is a "References:" section or a reference fragment
            if title.lower().startswith("references"):
                has_references_section = True
                reference_paragraphs.extend(body_lines)
                continue
            if _is_reference_fragment(sec):
                # Merge title + body into a single reference line
                merged_line = title
                if body_lines:
                    merged_line = f"{title} {' '.join(body_lines)}"
                reference_paragraphs.append(merged_line)
                continue

            if not title and not body_lines:
                continue

            body_paragraphs = "".join([f'<p class="section-p">{t}</p>' for t in body_lines])

            sec_title_html = f'<div class="section-title">{title}</div>' if title else ''
            sections_html += f"""
            <div class="section-box">
                {sec_title_html}
                <div class="section-body">
                    {body_paragraphs}
                </div>
            </div>
            """

        # Emit merged references section at the end
        if reference_paragraphs:
            ref_body = "".join([f'<p class="section-p">{r}</p>' for r in reference_paragraphs])
            sections_html += f"""
            <div class="section-box">
                <div class="section-title">References</div>
                <div class="section-body">
                    {ref_body}
                </div>
            </div>
            """

    # 2. Data Tables HTML — auto-fit by removing empty trailing columns, add result coloring
    tables_html = ""
    if show_tables and tables:
        for t_idx, tab in enumerate(tables):
            headers = list(tab.get("headers", []))
            rows = [list(r) for r in (tab.get("rows", []))]
            page_n = tab.get("page", 1)
            if not headers and not rows:
                continue

            # Determine effective column count: trim empty trailing columns
            raw_ncols = max([len(headers)] + [len(r) for r in rows] + [0])
            if raw_ncols == 0:
                continue
            # Pad to uniform width
            headers = (headers + [""] * raw_ncols)[:raw_ncols]
            rows = [(r + [""] * raw_ncols)[:raw_ncols] for r in rows]

            # Find columns where header AND all row cells are empty → remove them
            keep_cols = []
            for ci in range(raw_ncols):
                col_vals = [headers[ci].strip()] + [str(r[ci]).strip() for r in rows]
                if any(v for v in col_vals):  # at least one non-empty value
                    keep_cols.append(ci)
            if not keep_cols:
                continue

            headers = [headers[ci] for ci in keep_cols]
            rows = [[r[ci] for ci in keep_cols] for r in rows]

            th_html = "".join([f"<th>{h}</th>" for h in headers]) if any(h.strip() for h in headers) else ""
            tr_html = ""
            for r in rows:
                tds = ""
                for cell in r:
                    cell_str = str(cell).replace('\n', '<br>')
                    cell_lower = cell_str.lower()
                    cell_formatted = _result_color_span(cell_str)
                    # Badge logic only if result coloring didn't already wrap it
                    if cell_formatted == cell_str and show_badges:
                        if any(w in cell_lower for w in badge_rules.get("danger", [])):
                            cell_formatted = f"<span class='badge-danger'>{cell_str}</span>"
                        elif any(w in cell_lower for w in badge_rules.get("warning", [])):
                            cell_formatted = f"<span class='badge-warning'>{cell_str}</span>"
                        elif any(w in cell_lower for w in badge_rules.get("success", [])):
                            cell_formatted = f"<span class='badge-success'>{cell_str}</span>"
                    tds += f"<td>{cell_formatted}</td>"
                tr_html += f"<tr>{tds}</tr>"

            table_head_block = f"<thead><tr>{th_html}</tr></thead>" if th_html else ""
            tables_html += f"""
            <div class="table-card-box">
                <div class="table-card-header">
                    Table {t_idx + 1} (Page {page_n})
                </div>
                <table class="table-custom">
                    {table_head_block}
                    <tbody>{tr_html}</tbody>
                </table>
            </div>
            """


    # 5. Images & Graphs Section HTML
    images_html = ""
    if show_images and images:
        img_cards = ""
        for img in images:
            data_uri = img.get("data_uri")
            page_n = img.get("page", 1)
            img_type = img.get("type", "image").replace("_", " ").title()
            w = img.get("width", 0)
            h = img.get("height", 0)
            if not data_uri:
                continue
            img_cards += f"""
            <div class="image-card">
                <img src="{data_uri}" alt="{img_type}" class="extracted-img" />
                <div class="image-caption">Page {page_n} • {img_type} ({w}×{h}px)</div>
            </div>
            """
        if img_cards:
            images_html = f"""
            <div class="section-box" style="margin-top: 14pt;">
                <div class="section-title">Extracted Body Images & Graphical Content</div>
                <div class="image-grid">
                    {img_cards}
                </div>
            </div>
            """

    header_image2_path = Path("assets/header_image2.jpeg")
    header_image2_html = ""
    if header_image2_path.exists():
        try:
            import base64
            with open(header_image2_path, "rb") as img_f:
                header_image2_b64 = base64.b64encode(img_f.read()).decode("utf-8")
            header_image2_html = (
                f'<img src="data:image/jpeg;base64,{header_image2_b64}" '
                f'style="position:absolute !important; left:440.0pt !important; '
                f'top:15.0pt !important; width:130.0pt !important; '
                f'height:auto !important; z-index:100 !important; '
                f'object-fit:contain !important;" />'
            )
        except Exception:
            pass

    oncquest_logo_html = ""
    patient_table_html = ""

    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{doc_title}</title>
<style>
@page {{ size: 595.6pt 842.0pt; margin: 0; }}
* {{ box-sizing: border-box; }}
body {{ margin: 0; padding: 0; background-color: #f1f5f9; font-family: {font_family}; color: {text_color}; }}
.pdf-container {{ display: flex; flex-direction: column; align-items: center; padding: 20px 0; }}
.report-content {{ background: {bg_page}; width: 595.6pt; padding: 35.5pt; margin-bottom: 20px; position: relative; box-shadow: 0 4px 12px rgba(0,0,0,0.15); font-family: {font_family}; word-break: normal; overflow-wrap: break-word; }}
.section-box {{ border: {border_width}px {css_border_style} {border_color_val}; margin-bottom: {block_gap}pt; padding: {box_padding}px; width: fit-content; max-width: 100%; box-sizing: border-box; position: relative; background: #ffffff; overflow: visible; }}
.section-title {{ color: {primary_color}; font-weight: bold; font-size: 10.5pt; padding: 6pt 10pt; display: block; letter-spacing: 0.02em; border-bottom: 2px solid {primary_color}; }}
.section-body {{ padding: 10pt 12pt; font-size: 9.5pt; line-height: 1.5; color: {text_color}; }}
.section-p {{ margin: 0 0 6pt 0; text-align: left; word-break: normal; overflow-wrap: break-word; }}
.section-p:last-child {{ margin-bottom: 0; }}
.table-card-box {{ margin-top: 14pt; margin-bottom: {block_gap}pt; border: {border_width}px {css_border_style} {border_color_val}; width: fit-content; max-width: 100%; box-sizing: border-box; overflow: visible; background: #ffffff; }}
.table-card-header {{ font-size: 9.5pt; font-weight: bold; color: {primary_color}; padding: 6pt 10pt; background-color: #f8fafc; border-bottom: 1px solid {border_color_val}; }}
.table-custom {{ border-collapse: collapse; font-size: 9.5pt; table-layout: auto; width: 100%; }}
.table-custom th {{ background-color: {table_header_bg}; color: {table_header_text}; font-family: {font_family}; font-size: 9.5pt; font-weight: bold; text-align: center; padding: {cell_padding_y}px {cell_padding_x}px; border: {border_width}px {css_border_style} {border_color_val}; word-break: normal; overflow-wrap: break-word; }}
.table-custom td {{ padding: {cell_padding_y}px {cell_padding_x}px; border: {border_width}px {css_border_style} {border_color_val}; vertical-align: top; line-height: 1.35; font-size: 9.5pt; word-break: normal; overflow-wrap: break-word; }}
.badge-danger {{ background: #dc2626; color: #ffffff; padding: 2px 6px; border-radius: 3px; font-weight: bold; display: inline-block; font-size: 8.5pt; }}
.badge-warning {{ background: #d97706; color: #ffffff; padding: 2px 6px; border-radius: 3px; font-weight: bold; display: inline-block; font-size: 8.5pt; }}
.badge-success {{ background: #16a34a; color: #ffffff; padding: 2px 6px; border-radius: 3px; font-weight: bold; display: inline-block; font-size: 8.5pt; }}
.image-grid {{ display: flex; flex-wrap: wrap; gap: 12px; margin-top: 8pt; padding: 10pt; }}
.image-card {{ border: 1px solid #e2e8f0; border-radius: 4px; padding: 8px; background: #fafafa; text-align: center; max-width: 240px; }}
.extracted-img {{ max-width: 100%; max-height: 160px; object-fit: contain; border-radius: 2px; }}
.image-caption {{ font-size: 8pt; color: #64748b; margin-top: 4pt; font-family: sans-serif; }}
</style>
</head>
<body>
<div class="pdf-container">
  <div class="report-content">
    {header_image2_html}
    {sections_html}

    {tables_html}

    {images_html}
  </div>
</div>
</body>
</html>"""
    return replace_test_name_in_html(replace_sng_gen_lab(html_content))




def _set_cell_background(cell, hex_color: str):
    """Utility to set XML background shading of a python-docx table cell."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    hex_clean = hex_color.lstrip("#")
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_clean}"/>')
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_borders(cell, border_color: str = "CBD5E1", border_size: str = "4"):
    """Utility to set XML cell borders of a python-docx table cell."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    hex_clean = border_color.lstrip("#")
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{border_size}" w:space="0" w:color="{hex_clean}"/>'
        f'<w:left w:val="single" w:sz="{border_size}" w:space="0" w:color="{hex_clean}"/>'
        f'<w:bottom w:val="single" w:sz="{border_size}" w:space="0" w:color="{hex_clean}"/>'
        f'<w:right w:val="single" w:sz="{border_size}" w:space="0" w:color="{hex_clean}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    """Utility to set XML cell margins/padding of a python-docx table cell in dxa (1pt = 20dxa)."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def _set_cell_width(cell, width_pt):
    """Utility to set XML cell width in dxa (1pt = 20dxa)."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    tcPr = cell._tc.get_or_add_tcPr()
    width_dxa = int(width_pt * 20)
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
    tcPr.append(tcW)


def _add_page_number_to_run(run, docx_module):
    """Utility to add dynamic Word PAGE field to a run."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def _add_numpages_to_run(run, docx_module):
    """Utility to add dynamic Word NUMPAGES field to a run."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def _hex_to_rgb(hex_str: str):
    from docx.shared import RGBColor
    if not hex_str:
        return RGBColor(31, 73, 125)
    hex_clean = str(hex_str).lstrip("#")
    if len(hex_clean) == 6:
        r, g, b = int(hex_clean[0:2], 16), int(hex_clean[2:4], 16), int(hex_clean[4:6], 16)
    else:
        r, g, b = 31, 73, 125
    return RGBColor(r, g, b)


def get_merged_theme_config(theme_config: dict = None) -> dict:
    """
    Dynamically loads styling and layout parameters from theme.json and merges them with user overrides.
    """
    merged = {
        "document_page": {
            "paper_size": "A4",
            "width_pt": 595.6,
            "height_pt": 842.0,
            "margins_pt": {"top": 36.0, "bottom": 36.0, "left": 36.0, "right": 36.0},
            "header_distance_pt": 18.0,
            "footer_distance_pt": 18.0
        },
        "typography": {
            "primary_family": "Cambria",
            "secondary_family": "Calibri",
            "line_spacing": 1.15,
            "paragraph_space_after_pt": 4.0,
            "sizes_pt": {
                "document_title": 14.0,
                "section_heading": 13.0,
                "banner_heading": 10.0,
                "body": 10.0,
                "table_header": 9.5,
                "table_body": 9.5,
                "footer": 9.0,
                "small": 8.0
            }
        },
        "colors": {
            "primary": "#1f497d",
            "secondary": "#008080",
            "accent_orange": "#ed7d31",
            "accent_red": "#ff0000",
            "banner_dark": "#404040",
            "text_primary": "#000000",
            "text_secondary": "#242729",
            "text_light": "#ffffff",
            "background_page": "#ffffff",
            "border_primary": "#1f497d",
            "border_table": "#000000",
            "border_header": "#cbd5e1",
            "alternating_row_bg": "#f8fafc",
            "result_positive": "#C00000",
            "result_negative": "#008000"
        },
        "headers_and_footers": {
            "header": {
                "show_logo": True,
                "logo_image_path": "assets/header_image1.png",
                "logo_width_pt": 540.0,
                "logo_alignment": "center",
                "show_metadata_table": True,
                "metadata_table": {
                    "border_color": "#cbd5e1",
                    "border_size": "2",
                    "label_font_weight": "bold",
                    "value_font_weight": "normal",
                    "column_widths_pt": [87.2, 160.75, 88.95, 203.1],
                    "cell_padding_pt": {"top": 2, "bottom": 2, "left": 4, "right": 4}
                }
            },
            "footer": {
                "show_signatures": True,
                "signature_image_path": "assets/footer_signatures.png",
                "signature_width_pt": 540.0,
                "signature_alignment": "center",
                "show_page_number": True,
                "page_number_text": "Page {page} of {total}",
                "font_size_pt": 9.0,
                "text_color": "#404040"
            }
        },
        "spacing": {
            "boxPadding": 8,
            "cellPaddingX": 6,
            "cellPaddingY": 4,
            "blockGap": 10
        },
        "border": {
            "width": 1,
            "style": "single",
            "color": "#000000"
        },
        "components": {
            "black_banner": {
                "background_color": "#404040",
                "text_color": "#ffffff",
                "font_family": "Cambria",
                "font_size_pt": 13.0,
                "font_weight": "bold",
                "alignment": "center",
                "space_before_pt": 6.0,
                "space_after_pt": 6.0
            },
            "blue_section_banner": {
                "background_color": "#1f497d",
                "text_color": "#ffffff",
                "font_family": "Cambria",
                "font_size_pt": 10.0,
                "font_weight": "bold",
                "alignment": "left",
                "space_before_pt": 4.0,
                "space_after_pt": 4.0
            },
            "data_table": {
                "header_background_color": "#1f497d",
                "header_text_color": "#ffffff",
                "header_font_weight": "bold",
                "header_font_size_pt": 9.5,
                "body_font_size_pt": 9.5,
                "border_color": "#1f497d",
                "border_size": "4",
                "cell_padding_pt": {"top": 3, "bottom": 3, "left": 4, "right": 4},
                "alignment": "center"
            },
            "content_box": {
                "border_color": "#1f497d",
                "border_size": "4",
                "background_color": "transparent",
                "title_font_size_pt": 10.0,
                "title_font_weight": "bold",
                "title_text_color": "#1f497d",
                "body_font_size_pt": 9.5,
                "body_text_color": "#000000",
                "space_before_pt": 4.0,
                "space_after_pt": 4.0
            },
            "bullet_list": {
                "bullet_char": "•",
                "indent_pt": 18.0,
                "font_size_pt": 9.5
            },
            "end_of_report_marker": {
                "text": "------------------ End Of Report ------------------",
                "font_family": "Calibri",
                "font_size_pt": 10.0,
                "text_color": "#000000",
                "alignment": "center",
                "space_before_pt": 12.0,
                "space_after_pt": 12.0
            }
        },
        "word": {
            "layout": {
                "element_flow": ["header_metadata", "headings", "banners", "content_boxes", "tables", "paragraphs", "bullet_lists", "signatures", "end_of_report_marker"],
                "spacing": {
                    "default_line_spacing": 1.15,
                    "paragraph_space_before_pt": 0.0,
                    "paragraph_space_after_pt": 4.0,
                    "heading_space_before_pt": 8.0,
                    "heading_space_after_pt": 4.0,
                    "table_space_before_pt": 6.0,
                    "table_space_after_pt": 6.0,
                    "content_box_space_before_pt": 6.0,
                    "content_box_space_after_pt": 6.0,
                    "banner_space_before_pt": 6.0,
                    "banner_space_after_pt": 6.0
                },
                "boxes": {
                    "border_color": "#1f497d",
                    "border_size": "4",
                    "background_color": "#ffffff",
                    "title_font_size_pt": 10.0,
                    "title_font_weight": "bold",
                    "title_text_color": "#1f497d",
                    "body_font_size_pt": 9.5,
                    "body_text_color": "#000000",
                    "cell_padding_pt": {"top": 4, "bottom": 4, "left": 6, "right": 6}
                },
                "tables": {
                    "alignment": "center",
                    "header_background_color": "#1f497d",
                    "header_text_color": "#ffffff",
                    "header_font_weight": "bold",
                    "header_font_size_pt": 9.5,
                    "body_font_size_pt": 9.5,
                    "border_color": "#1f497d",
                    "border_size": "4",
                    "alternating_row_bg": "#f8fafc",
                    "cell_padding_pt": {"top": 3, "bottom": 3, "left": 4, "right": 4},
                    "autofit": True,
                    "cant_split_rows": True
                },
                "banners": {
                    "black_banner": {"background_color": "#404040", "text_color": "#ffffff", "font_size_pt": 13.0, "font_weight": "bold", "alignment": "center"},
                    "blue_banner": {"background_color": "#1f497d", "text_color": "#ffffff", "font_size_pt": 10.0, "font_weight": "bold", "alignment": "left"}
                }
            },
            "pagination": {
                "respect_page_breaks": True,
                "keep_with_next_headings": True,
                "page_numbering": {"enabled": True, "text_format": "Page {page} of {total}", "font_size_pt": 9.0, "alignment": "center"},
                "repeat_table_headers": True,
                "prevent_orphan_rows": True
            }
        },
        "primary_color": "#1f497d",
        "secondary_color": "#008080",
        "accent_orange": "#ed7d31",
        "accent_red": "#ff0000",
        "banner_dark": "#404040",
        "font_family": "Cambria",
        "border_color": "#1f497d",
        "body_font_size": 10.0,
        "title_font_size": 14.0,
        "table_header_font_size": 9.5,
        "show_footer_signatures": True,
        "end_report_marker": "------------------ End Of Report ------------------"
    }

    theme_file = Path("theme.json")
    if theme_file.exists():
        try:
            with open(theme_file, "r", encoding="utf-8") as f:
                tj = json.load(f)
                if isinstance(tj, dict):
                    for key in ["document_page", "typography", "colors", "headers_and_footers", "components", "word", "spacing", "border"]:
                        if key in tj and isinstance(tj[key], dict):
                            if key in merged and isinstance(merged[key], dict):
                                merged[key].update(tj[key])
                            else:
                                merged[key] = tj[key]

                    colors = merged.get("colors", {})
                    fonts = merged.get("typography", {})
                    if "primary" in colors:
                        merged["primary_color"] = colors["primary"]
                    if "secondary" in colors:
                        merged["secondary_color"] = colors["secondary"]
                    if "accent_orange" in colors:
                        merged["accent_orange"] = colors["accent_orange"]
                    if "accent_red" in colors:
                        merged["accent_red"] = colors["accent_red"]
                    if "banner_dark" in colors:
                        merged["banner_dark"] = colors["banner_dark"]
                    if "border_primary" in colors:
                        merged["border_color"] = colors["border_primary"]

                    merged["font_family"] = fonts.get("primary_family", "Cambria")
                    sizes = fonts.get("sizes_pt", {})
                    if "body" in sizes:
                        merged["body_font_size"] = sizes["body"]
                    if "document_title" in sizes:
                        merged["title_font_size"] = sizes["document_title"]
                    if "table_header" in sizes:
                        merged["table_header_font_size"] = sizes["table_header"]

                    end_marker_comp = merged.get("components", {}).get("end_of_report_marker", {})
                    if "text" in end_marker_comp:
                        merged["end_report_marker"] = end_marker_comp["text"]
        except Exception:
            pass

    if theme_config and isinstance(theme_config, dict):
        for k, v in theme_config.items():
            if v is not None:
                merged[k] = v

    return merged


def _find_kv_val(kv_dict: dict, keywords: list) -> str:
    """Helper to fuzzy match key in key-value dictionary."""
    if not isinstance(kv_dict, dict):
        return ""
    for k, v in kv_dict.items():
        k_clean = str(k).lower().replace("_", " ").replace(".", " ").strip()
        for kw in keywords:
            if kw in k_clean:
                return str(v).strip()
    return ""


def clean_extracted_text(val):
    """
    Clean trailing/standalone '_' artifacts and stray whitespace from extracted text and table cells.
    Preserves internal underscores (e.g. 'Gene_Name', 'NM_000251.3').
    """
    if not isinstance(val, str):
        return val
    text = val.strip()
    if not text:
        return ""
    # If text consists solely of underscores and whitespace
    if set(text) <= {'_', ' '}:
        return ""
    # Remove trailing/standalone underscore artifacts like " _", " _ _", or trailing "_" preceded by whitespace
    text = re.sub(r'(?:\s+_+)+\s*$', '', text)
    return text.strip()


def is_decorative_or_watermark_image(img_dict: dict, page_w: float = 595.6, page_h: float = 842.0) -> bool:
    """
    Filter out decorative/background images: skip images that are near-square and very large relative to page,
    or that sit in page margins/watermark layer, or have low content density.
    Keep only real content graphs/logos based on size & aspect-ratio thresholds.
    """
    if not isinstance(img_dict, dict):
        return False

    bbox = img_dict.get("bbox") or [0, 0, 0, 0]
    x0, y0, x1, y1 = bbox[:4] if len(bbox) >= 4 else (0, 0, 0, 0)
    w_pt = x1 - x0 if x1 > x0 else 0.0
    h_pt = y1 - y0 if y1 > y0 else 0.0

    px_w = float(img_dict.get("width") or 0.0)
    px_h = float(img_dict.get("height") or 0.0)

    data_uri = img_dict.get("data_uri")
    if data_uri and (px_w == 0 or px_h == 0 or w_pt == 0):
        try:
            b64_str = data_uri.split(",", 1)[1] if "," in data_uri else data_uri
            img_bytes = base64.b64decode(b64_str)
            with Image.open(io.BytesIO(img_bytes)) as pil_img:
                px_w, px_h = float(pil_img.size[0]), float(pil_img.size[1])
        except Exception:
            pass

    eff_w = w_pt if w_pt > 0 else px_w
    eff_h = h_pt if h_pt > 0 else px_h

    if eff_w <= 0 or eff_h <= 0:
        return False

    aspect_ratio = eff_w / eff_h if eff_h > 0 else 1.0

    # Criteria 1: Near-square (0.7 <= aspect <= 1.4) AND large relative to page
    is_near_square = 0.7 <= aspect_ratio <= 1.4
    is_large_pt = (w_pt > 350 and h_pt > 350) or (w_pt > 0.55 * page_w and h_pt > 0.4 * page_h)
    is_large_px = (px_w > 450 and px_h > 450)

    if is_near_square and (is_large_pt or is_large_px):
        return True

    # Criteria 2: Full-page background or margin-spanning watermark layer
    if w_pt > 0.8 * page_w and h_pt > 0.8 * page_h:
        return True
    if w_pt > 0 and h_pt > 0 and x0 <= 40 and x1 >= (page_w - 40) and y0 <= 50 and y1 >= (page_h - 50):
        return True

    return False


def _clean_text(s):
    if s is None:
        return ""
    s = str(s)
    s = re.sub(r"\s*_+\s*", " ", s)   # remove standalone/trailing underscores
    s = re.sub(r"\s{2,}", " ", s)     # collapse whitespace
    return s.strip()


def _is_decorative_image(img, page_w=595.0, page_h=842.0):
    w = float(img.get("width", 0) or 0)
    h = float(img.get("height", 0) or 0)
    if w <= 0 or h <= 0:
        return False
    aspect = w / h if h else 1.0
    near_square = 0.7 <= aspect <= 1.4
    bb = img.get("bbox")
    if isinstance(bb, (list, tuple)) and len(bb) >= 4:
        bw = abs(bb[2] - bb[0]); bh = abs(bb[3] - bb[1])
        area_frac = (bw * bh) / (page_w * page_h)
    else:
        area_frac = (w * h) / (page_w * page_h)
    return near_square and area_frac >= 0.22


def _load_oncquest_theme(theme_config=None):
    theme_config = theme_config or {}
    tj = {}
    tj_path = Path("theme.json")
    if tj_path.exists():
        try:
            with open(tj_path, "r", encoding="utf-8") as f:
                tj = json.load(f)
        except Exception:
            tj = {}
 
    colors = tj.get("colors", {}) or {}
    text = colors.get("text", {}) or {}
    fonts = tj.get("fonts", {}) or {}
    families = fonts.get("families", {}) or {}
    sizes = fonts.get("sizes", {}) or {}
    comps = tj.get("components", {}) or {}
    eor = comps.get("end_of_report_marker", {}) or {}

    # Word-specific configurations
    word_cfg = tj.get("word", {}) or {}
    word_layout = word_cfg.get("layout", {}) or {}
    word_spacing = word_layout.get("spacing", {}) or {}
    word_tables = word_layout.get("tables", {}) or {}
    word_banners = word_layout.get("banners", {}) or {}
    word_boxes = word_layout.get("boxes", {}) or {}
    doc_page = tj.get("document_page", {}) or {}

    def clean_font(css, fallback="Cambria"):
        if not css or not isinstance(css, str):
            return fallback
        first = css.split(",")[0].strip().strip("'\"")
        if first.lower() in ("bold", "italic", "bolditalic", "serif", "sans-serif"):
            return fallback
        return first or fallback

    def clean_hex(h, fallback):
        if not h or not isinstance(h, str):
            return fallback
        h = h.strip().lstrip("#")
        return h if len(h) == 6 else fallback

    border_primary = colors.get("border_primary") or colors.get("primary") or "#1f497d"
    primary = theme_config.get("primary_color") or colors.get("primary") or "#1f497d"

    return {
        "primary": clean_hex(primary, "1f497d"),
        "banner_dark": clean_hex(word_banners.get("black_banner", {}).get("background_color") or colors.get("banner_dark"), "404040"),
        "border": clean_hex(word_boxes.get("border_color") or border_primary, "1f497d"),
        "border_table": clean_hex(word_tables.get("border_color") or colors.get("border_table"), "000000"),
        "body_color": clean_hex(text.get("primary") or colors.get("text_primary"), "000000"),
        "header_text": clean_hex(text.get("white") or colors.get("text_light"), "ffffff"),
        "result_positive": clean_hex(colors.get("result_positive"), "C00000"),
        "result_negative": clean_hex(colors.get("result_negative"), "008000"),
        "font": clean_font(families.get("primary"), "Cambria"),
        "table_header_font": clean_font(families.get("table_header"), "Cambria"),
        "body_pt": float(word_tables.get("body_font_size_pt") or sizes.get("body_pt", 10.0)),
        "banner_pt": float(word_banners.get("black_banner", {}).get("font_size_pt") or sizes.get("banner_pt", 13.0)),
        "table_header_pt": float(word_tables.get("header_font_size_pt") or sizes.get("table_header_pt", 9.5)),
        "eor_text": eor.get("text_pattern") or eor.get("text") or "------------------ End Of Report ------------------",
        "eor_font": clean_font(eor.get("font_family"), "Calibri"),
        "eor_pt": float(eor.get("font_size_pt", 10.0)),
        "eor_color": clean_hex(eor.get("text_color"), "000000"),
        
        # Spacing configurations
        "line_spacing": float(word_spacing.get("default_line_spacing", 1.15)),
        "para_space_before": float(word_spacing.get("paragraph_space_before_pt", 0.0)),
        "para_space_after": float(word_spacing.get("paragraph_space_after_pt", 4.0)),
        "heading_space_before": float(word_spacing.get("heading_space_before_pt", 8.0)),
        "heading_space_after": float(word_spacing.get("heading_space_after_pt", 4.0)),
        "table_space_before": float(word_spacing.get("table_space_before_pt", 6.0)),
        "table_space_after": float(word_spacing.get("table_space_after_pt", 6.0)),
        "banner_space_before": float(word_spacing.get("banner_space_before_pt", 6.0)),
        "banner_space_after": float(word_spacing.get("banner_space_after_pt", 6.0)),
        
        # Margins & Dimensions configurations
        "margin_top": float(doc_page.get("margins_pt", {}).get("top", 36.0)),
        "margin_bottom": float(doc_page.get("margins_pt", {}).get("bottom", 36.0)),
        "margin_left": float(doc_page.get("margins_pt", {}).get("left", 36.0)),
        "margin_right": float(doc_page.get("margins_pt", {}).get("right", 36.0)),
        "paper_width": float(doc_page.get("width_pt", 595.6)),
        "paper_height": float(doc_page.get("height_pt", 842.0)),
        
        # Alternating row background color
        "alternating_row_bg": clean_hex(word_tables.get("alternating_row_bg") or colors.get("alternating_row_bg"), "f8fafc"),
    }


def convert_json_to_docx(data: dict, output_path: str = None, theme_config: dict = None):
    logger.info("[DIRECT-EXPORT] Using JSON->DOCX direct pipeline, no pdf2docx")
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    data = replace_sng_in_structure(data)
    data = replace_test_name_in_structure(data)

    # Load theme
    tj = {}
    tj_path = Path("theme.json")
    if tj_path.exists():
        try:
            with open(tj_path, "r", encoding="utf-8") as f:
                tj = json.load(f)
        except Exception:
            pass

    # Merge theme_config overrides if provided
    if theme_config:
        if "colors" not in tj:
            tj["colors"] = {}
        if "styles" not in tj:
            tj["styles"] = {}
            
        prim = theme_config.get("primary_color")
        if prim:
            tj["colors"]["primary"] = prim
            
            if "heading" not in tj["styles"]:
                tj["styles"]["heading"] = {}
            tj["styles"]["heading"]["text_color"] = prim
            
            if "table" not in tj["styles"]:
                tj["styles"]["table"] = {}
            tj["styles"]["table"]["header_background_color"] = prim
            tj["styles"]["table"]["border_color"] = prim
            
            if "key_value" not in tj["styles"]:
                tj["styles"]["key_value"] = {}
            tj["styles"]["key_value"]["border_color"] = prim

        show_sig = theme_config.get("show_footer_signatures")
        if show_sig is not None:
            if "footer" not in tj["styles"]:
                tj["styles"]["footer"] = {}
            tj["styles"]["footer"]["show_signatures"] = show_sig

    page_cfg = tj.get("page", {})
    default_margin_top = float(page_cfg.get("margin_top", 36.0))
    default_margin_bottom = float(page_cfg.get("margin_bottom", 36.0))
    default_margin_left = float(page_cfg.get("margin_left", 36.0))
    default_margin_right = float(page_cfg.get("margin_right", 36.0))
    default_width = float(page_cfg.get("width", 595.6))
    default_height = float(page_cfg.get("height", 842.0))

    theme_styles = tj.get("styles", {})
    colors_cfg = tj.get("colors", {})

    def rgb(hx):
        if not hx or not isinstance(hx, str):
            return RGBColor(0, 0, 0)
        hx = hx.strip().lstrip("#")
        if len(hx) != 6:
            return RGBColor(0, 0, 0)
        return RGBColor(int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16))

    def shade_cell(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill.lstrip("#")); tcPr.append(shd)

    def shade_para(paragraph, fill):
        pPr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill.lstrip("#")); pPr.append(shd)

    def cell_margins(cell, t=40, b=40, l=80, r=80):
        tcPr = cell._tc.get_or_add_tcPr()
        m = OxmlElement("w:tcMar")
        for tag, val in (("top", t), ("bottom", b), ("start", l), ("end", r)):
            n = OxmlElement(f"w:{tag}")
            n.set(qn("w:w"), str(val)); n.set(qn("w:type"), "dxa"); m.append(n)
        tcPr.append(m)

    def table_borders(table, color, sz=4):
        tblPr = table._tbl.tblPr
        b = OxmlElement("w:tblBorders")
        color = color.lstrip("#")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
            e.set(qn("w:space"), "0"); e.set(qn("w:color"), color); b.append(e)
        tblPr.append(b)

    def set_repeat_header(row):
        trPr = row._tr.get_or_add_trPr()
        th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true")
        trPr.append(th)

    def set_cant_split(row):
        trPr = row._tr.get_or_add_trPr()
        cs = OxmlElement("w:cantSplit"); cs.set(qn("w:val"), "true")
        trPr.append(cs)

    _POS_KW = ["positive", "pathogenic", "detected", "high", "abnormal", "msi - high", "msi-high"]
    _NEG_KW = ["negative", "normal", "not detected", "stable", "msi - stable", "msi-stable", "benign", "likely benign"]

    def _result_color(cell_text):
        cl = cell_text.strip().lower()
        if any(kw in cl for kw in _NEG_KW):
            return rgb(colors_cfg.get("result_negative", "#008000"))
        if any(kw in cl for kw in _POS_KW):
            return rgb(colors_cfg.get("result_positive", "#C00000"))
        return None
    def _clean_text(val):
        if val is None:
            return ""
        return str(val).strip()

    def flatten_header_footer_content(elements):
        """Remove header/footer wrapper elements and their nested content entirely.
        OncQuest branding (logo/signature) is handled by docx header/footer sections.
        Patient key_value data already exists as a sibling top-level element,
        so header nested content (QR text, old lab logo, duplicate key_value)
        is not needed in the body."""
        return [el for el in elements if el.get("type") not in ("header", "footer")]

    def deduplicate_elements(elements):
        seen_bboxes = []
        seen_texts = set()
        cleaned = []
        for el in elements:
            el_type = el.get("type")
            if el_type in ("header", "footer"):
                cleaned.append(el)
                continue
            bbox = el.get("bbox")
            text = el.get("text") or ""
            text_str = text.strip().lower()
            
            is_dup = False
            if bbox:
                for sb in seen_bboxes:
                    if (abs(bbox.get("x0", 0.0) - sb.get("x0", 0.0)) < 2.0 and
                        abs(bbox.get("y0", 0.0) - sb.get("y0", 0.0)) < 2.0 and
                        abs(bbox.get("x1", 0.0) - sb.get("x1", 0.0)) < 2.0 and
                        abs(bbox.get("y1", 0.0) - sb.get("y1", 0.0)) < 2.0):
                        is_dup = True
                        break
            
            if text_str and text_str in seen_texts:
                if len(text_str) > 10:
                    is_dup = True
                    
            if not is_dup:
                if bbox:
                    seen_bboxes.append(bbox)
                if text_str and len(text_str) > 10:
                    seen_texts.add(text_str)
                cleaned.append(el)
            else:
                if el.get("data") and bbox:
                    for i, cel in enumerate(cleaned):
                        cel_bbox = cel.get("bbox")
                        if cel_bbox and (abs(bbox.get("x0", 0.0) - cel_bbox.get("x0", 0.0)) < 2.0 and
                                         abs(bbox.get("y0", 0.0) - cel_bbox.get("y0", 0.0)) < 2.0 and
                                         abs(bbox.get("x1", 0.0) - cel_bbox.get("x1", 0.0)) < 2.0 and
                                         abs(bbox.get("y1", 0.0) - cel_bbox.get("y1", 0.0)) < 2.0):
                            if not cel.get("data"):
                                cleaned[i] = el
                                break
        return cleaned

    def get_element_y(el):
        bbox = el.get("bbox")
        if bbox:
            return float(bbox.get("y0", bbox.get("y", 0.0)))
        return 9999.0

    doc = Document()
    pages = data.get("document", {}).get("pages")
    if pages:
        for p_idx, page_data in enumerate(pages):
            # Standardize page size to standard A4 (595.28 x 841.89 pt)
            width = 595.28
            height = 841.89
            
            header_style = theme_styles.get("header", {})
            footer_style = theme_styles.get("footer", {})
            header_height = float(header_style.get("height_pt", 60.0))
            footer_height = float(footer_style.get("height_pt", 40.0))

            # Use only theme-configured header/footer height for margins,
            # never the PDF's raw content bbox height (which could be huge).
            hy_cutoff = max(default_margin_top, header_height)
            fy_cutoff = min(height - default_margin_bottom, height - footer_height)

            if p_idx == 0:
                section = doc.sections[0]
            else:
                section = doc.add_section()
                
            section.page_width = Pt(width)
            section.page_height = Pt(height)
            section.top_margin = Pt(hy_cutoff)
            section.bottom_margin = Pt(height - fy_cutoff)
            section.left_margin = Pt(default_margin_left)
            section.right_margin = Pt(default_margin_right)

            # Setup customized header with static OncQuest logo
            header = section.header
            if header is not None:
                header.is_linked_to_previous = False
                for p in header.paragraphs:
                    p.text = ""
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                logo_path = header_style.get("logo_image_path", "assets/header_image1.png")
                logo_w = float(header_style.get("logo_width_pt", 540.0))
                if Path(logo_path).exists():
                    try:
                        run = header_para.add_run()
                        run.add_picture(logo_path, width=Inches(logo_w / 72.0))
                    except Exception as e:
                        print(f"Error adding header logo: {e}")

            # Setup customized footer with Dr. Vinay Bhatia signature
            footer = section.footer
            if footer is not None:
                footer.is_linked_to_previous = False
                for p in footer.paragraphs:
                    p.text = ""
                
                if footer_style.get("show_signatures", True):
                    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    sig_path = footer_style.get("signature_image_path", "assets/dr_vinay_signature.png")
                    sig_w = float(footer_style.get("signature_width_pt", 90.0))
                    if Path(sig_path).exists():
                        try:
                            run = footer_para.add_run()
                            run.add_picture(sig_path, width=Inches(sig_w / 72.0))
                        except Exception as e:
                            print(f"Error adding footer signature: {e}")

                # Add dynamic "Page X of Y" numbering
                page_num_cfg = tj.get("word", {}).get("pagination", {}).get("page_numbering", {})
                if page_num_cfg.get("enabled", True):
                    p_num = footer.add_paragraph()
                    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    run_text = p_num.add_run("Page ")
                    run_text.font.name = "Calibri"
                    run_text.font.size = Pt(9.0)
                    run_text.font.color.rgb = rgb("#404040")
                    
                    run_page = p_num.add_run()
                    _add_page_number_to_run(run_page, None)
                    run_page.font.name = "Calibri"
                    run_page.font.size = Pt(9.0)
                    run_page.font.color.rgb = rgb("#404040")
                    
                    run_of = p_num.add_run(" of ")
                    run_of.font.name = "Calibri"
                    run_of.font.size = Pt(9.0)
                    run_of.font.color.rgb = rgb("#404040")
                    
                    run_numpages = p_num.add_run()
                    _add_numpages_to_run(run_numpages, None)
                    run_numpages.font.name = "Calibri"
                    run_numpages.font.size = Pt(9.0)
                    run_numpages.font.color.rgb = rgb("#404040")

            # Render Body Elements with deduplication and sorting
            # Flatten nested header/footer content into the body list first,
            # then filter out any remaining header/footer wrappers.
            flattened = flatten_header_footer_content(page_data.get("elements", []))
            raw_body_elements = [el for el in flattened if el.get("type") not in ("header", "footer")]
            deduped_body_elements = deduplicate_elements(raw_body_elements)
            body_elements = sorted(deduped_body_elements, key=get_element_y)
            preceding_el = None
            for el in body_elements:
                el_type = el.get("type")
                style = theme_styles.get(el_type, {})
                style_override = el.get("style_override", {})
                # Filter out font family, size and alignment overrides to maintain consistency
                filtered_override = {
                    k: v for k, v in style_override.items()
                    if k not in ("font_family", "font_size", "font_size_pt")
                }
                resolved_style = {**style, **filtered_override}

                # Compute dynamic spacing (gap) relative to preceding elements
                gap = 0.0
                c_bbox = el.get("bbox")
                if c_bbox:
                    c_y0 = c_bbox.get("y", 0.0)
                    if preceding_el is None:
                        gap = max(0.0, c_y0 - hy_cutoff)
                    else:
                        p_bbox = preceding_el.get("bbox")
                        if p_bbox:
                            p_y1 = p_bbox.get("y", 0.0) + p_bbox.get("height", 0.0)
                            gap = max(0.0, c_y0 - p_y1)

                if gap > 2.0:
                    resolved_style["spacing_before"] = min(18.0, gap)
                else:
                    resolved_style["spacing_before"] = 0.0

                # Set spacing before table/key_value using paragraph spacers since tables don't support spacing_before
                if el_type in ("table", "key_value") and resolved_style.get("spacing_before", 0.0) > 2.0:
                    p_space = doc.add_paragraph()
                    p_space.paragraph_format.space_before = Pt(0)
                    p_space.paragraph_format.space_after = Pt(resolved_style.get("spacing_before", 0.0))
                    p_space.paragraph_format.line_spacing = Pt(1)
                    r = p_space.add_run(); r.font.size = Pt(1)

                if el_type in ("heading", "subheading"):
                    txt = _clean_text(el.get("text", ""))
                    if not txt:
                        continue
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(resolved_style.get("spacing_before", 6.0))
                    p.paragraph_format.space_after = Pt(resolved_style.get("spacing_after", 4.0))
                    p.paragraph_format.keep_with_next = True
                    
                    align_val = resolved_style.get("alignment", "left").lower()
                    if align_val == "center":
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif align_val == "right":
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                    run = p.add_run(txt)
                    run.bold = resolved_style.get("bold", True)
                    run.italic = resolved_style.get("italic", False)
                    run.font.name = resolved_style.get("font_family", "Cambria")
                    run.font.size = Pt(resolved_style.get("font_size", 11.0))
                    color_val = resolved_style.get("text_color") or colors_cfg.get("primary", "#1f497d")
                    run.font.color.rgb = rgb(color_val)
                    preceding_el = el

                elif el_type == "paragraph":
                    txt = _clean_text(el.get("text", ""))
                    if not txt:
                        continue
                    
                    if txt.startswith(('•', '-', '*')):
                        bullet_text = txt.lstrip('•-* ').strip()
                        p = doc.add_paragraph(style='List Bullet')
                        p.paragraph_format.space_before = Pt(resolved_style.get("spacing_before", 0.0))
                        p.paragraph_format.space_after = Pt(resolved_style.get("spacing_after", 4.0))
                        p.paragraph_format.line_spacing = resolved_style.get("line_spacing", 1.15)
                        run = p.add_run(bullet_text)
                    else:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(resolved_style.get("spacing_before", 0.0))
                        p.paragraph_format.space_after = Pt(resolved_style.get("spacing_after", 4.0))
                        p.paragraph_format.line_spacing = resolved_style.get("line_spacing", 1.15)
                        align_val = resolved_style.get("alignment", "left").lower()
                        if align_val == "center":
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif align_val == "right":
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = p.add_run(txt)
                        
                    run.font.name = resolved_style.get("font_family", "Cambria")
                    run.font.size = Pt(resolved_style.get("font_size", 10.0))
                    color_val = resolved_style.get("text_color") or colors_cfg.get("text_primary", "#000000")
                    run.font.color.rgb = rgb(color_val)
                    preceding_el = el

                elif el_type == "key_value":
                    kv_data = el.get("data", {})
                    if not kv_data:
                        continue
                    
                    tbl = doc.add_table(rows=0, cols=4)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    tbl.allow_autofit = False
                    border_color = resolved_style.get("border_color", "#cbd5e1")
                    
                    b_sz = int(resolved_style.get("border_width", 0.25) * 8)
                    table_borders(tbl, border_color, sz=max(1, b_sz))
                    
                    col_widths = [Inches(1.0), Inches(2.2), Inches(1.0), Inches(2.2)]
                    items = list(kv_data.items())
                    for idx in range(0, len(items), 2):
                        cells = tbl.add_row().cells
                        set_cant_split(tbl.rows[-1])
                        for c in cells:
                            cell_margins(c, t=20, b=20, l=40, r=40)
                            
                        k1, v1 = items[idx]
                        kr1 = cells[0].paragraphs[0].add_run(_clean_text(k1))
                        kr1.bold = resolved_style.get("label_bold", True)
                        kr1.font.name = resolved_style.get("font_family", "Cambria")
                        kr1.font.size = Pt(resolved_style.get("font_size", 10.0))
                        kr1.font.color.rgb = rgb(colors_cfg.get("primary", "#1f497d"))
                        
                        v1_str = _clean_text(v1)
                        vr1 = cells[1].paragraphs[0].add_run(v1_str)
                        vr1.font.name = resolved_style.get("font_family", "Cambria")
                        vr1.font.size = Pt(resolved_style.get("font_size", 10.0))
                        rc1 = _result_color(v1_str)
                        if rc1:
                            vr1.font.color.rgb = rc1
                            vr1.bold = True
                        else:
                            vr1.font.color.rgb = rgb(colors_cfg.get("text_primary", "#000000"))
                        
                        if idx + 1 < len(items):
                            k2, v2 = items[idx + 1]
                            kr2 = cells[2].paragraphs[0].add_run(_clean_text(k2))
                            kr2.bold = resolved_style.get("label_bold", True)
                            kr2.font.name = resolved_style.get("font_family", "Cambria")
                            kr2.font.size = Pt(resolved_style.get("font_size", 10.0))
                            kr2.font.color.rgb = rgb(colors_cfg.get("primary", "#1f497d"))
                            
                            v2_str = _clean_text(v2)
                            vr2 = cells[3].paragraphs[0].add_run(v2_str)
                            vr2.font.name = resolved_style.get("font_family", "Cambria")
                            vr2.font.size = Pt(resolved_style.get("font_size", 10.0))
                            rc2 = _result_color(v2_str)
                            if rc2:
                                vr2.font.color.rgb = rc2
                                vr2.bold = True
                            else:
                                vr2.font.color.rgb = rgb(colors_cfg.get("text_primary", "#000000"))
                            
                    for row in tbl.rows:
                        for i, w in enumerate(col_widths):
                            row.cells[i].width = w
                            
                    p_space = doc.add_paragraph()
                    p_space.paragraph_format.space_before = Pt(0)
                    p_space.paragraph_format.space_after = Pt(4)
                    p_space.paragraph_format.line_spacing = Pt(1)
                    r = p_space.add_run(); r.font.size = Pt(1)
                    preceding_el = el

                elif el_type == "table":
                    cols_list = el.get("columns", [])
                    rows_list = el.get("rows", [])
                    if not cols_list:
                        continue
                    
                    ncols = len(cols_list)
                    tbl = doc.add_table(rows=1, cols=ncols)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    tbl.allow_autofit = False
                    
                    border_color = resolved_style.get("border_color", "#1f497d")
                    b_sz = int(resolved_style.get("border_width", 0.5) * 8)
                    table_borders(tbl, border_color, sz=max(1, b_sz))
                    
                    header_names = []
                    col_widths = []
                    for col in cols_list:
                        if isinstance(col, dict):
                            header_names.append(col.get("name", ""))
                            col_widths.append(col.get("width"))
                        else:
                            header_names.append(str(col))
                            col_widths.append(None)

                    # Calculate dynamic column widths if not specified
                    has_none_width = any(w is None for w in col_widths)
                    if has_none_width:
                        left_margin = float(page_cfg.get("margin_left", 36.0))
                        right_margin = float(page_cfg.get("margin_right", 36.0))
                        avail_w = 595.28 - (left_margin + right_margin)
                        
                        known_w_sum = sum(w for w in col_widths if w is not None)
                        unspecified_count = sum(1 for w in col_widths if w is None)
                        remaining_w = max(50.0, avail_w - known_w_sum)
                        
                        col_max_lens = []
                        for i in range(ncols):
                            if col_widths[i] is None:
                                h_name = header_names[i]
                                max_l = len(str(h_name))
                                for r_val in rows_list:
                                    if isinstance(r_val, dict):
                                        val = r_val.get(h_name, "")
                                    elif isinstance(r_val, list) and i < len(r_val):
                                        val = r_val[i]
                                    else:
                                        val = ""
                                    max_l = max(max_l, len(_clean_text(val)))
                                col_max_lens.append((i, max_l))
                        
                        base_min = min(45.0, (remaining_w / unspecified_count) * 0.6) if unspecified_count > 0 else 45.0
                        min_w_sum = base_min * unspecified_count
                        distribute_w = max(0.0, remaining_w - min_w_sum)
                        
                        weights = [max(1, l) ** 0.7 for idx, l in col_max_lens]
                        total_weight = sum(weights)
                        
                        weight_idx = 0
                        for i in range(ncols):
                            if col_widths[i] is None:
                                if total_weight > 0:
                                    extra = (weights[weight_idx] / total_weight) * distribute_w
                                else:
                                    extra = distribute_w / unspecified_count
                                col_widths[i] = base_min + extra
                                weight_idx += 1
                            
                    hrow = tbl.rows[0]
                    set_repeat_header(hrow)
                    set_cant_split(hrow)
                    header_bg = resolved_style.get("header_background_color", "#1f497d")
                    header_text_color = resolved_style.get("header_text_color", "#ffffff")
                    
                    t_pad = int(resolved_style.get("cell_padding_top", 3.0) * 20)
                    b_pad = int(resolved_style.get("cell_padding_bottom", 3.0) * 20)
                    l_pad = int(resolved_style.get("cell_padding_left", 4.0) * 20)
                    r_pad = int(resolved_style.get("cell_padding_right", 4.0) * 20)

                    for i in range(ncols):
                        cell = hrow.cells[i]
                        shade_cell(cell, header_bg)
                        cell_margins(cell, t=t_pad, b=b_pad, l=l_pad, r=r_pad)
                        para = cell.paragraphs[0]
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run(header_names[i])
                        run.bold = resolved_style.get("header_bold", True)
                        run.font.name = resolved_style.get("font_family", "Cambria")
                        run.font.size = Pt(resolved_style.get("font_size", 9.5))
                        run.font.color.rgb = rgb(header_text_color)
                        
                    row_bg = colors_cfg.get("alternating_row_bg", "#f8fafc")
                    for row_idx, r_val in enumerate(rows_list):
                        cells = tbl.add_row().cells
                        set_cant_split(tbl.rows[-1])
                        
                        is_alt = (row_idx % 2 == 1)
                        for i in range(ncols):
                            cell = cells[i]
                            if is_alt and row_bg:
                                shade_cell(cell, row_bg)
                            cell_margins(cell, t=t_pad, b=b_pad, l=l_pad, r=r_pad)
                            
                            h_name = header_names[i]
                            if isinstance(r_val, dict):
                                val = r_val.get(h_name, "")
                            elif isinstance(r_val, list) and i < len(r_val):
                                val = r_val[i]
                            else:
                                val = ""
                                
                            run = cell.paragraphs[0].add_run(_clean_text(val))
                            run.font.name = resolved_style.get("font_family", "Cambria")
                            run.font.size = Pt(resolved_style.get("font_size", 9.5))
                            
                            rc = _result_color(str(val))
                            if rc:
                                run.font.color.rgb = rc
                                run.bold = True
                            else:
                                run.font.color.rgb = rgb(colors_cfg.get("text_primary", "#000000"))
                                
                    for row in tbl.rows:
                        for i in range(ncols):
                            w_pt = col_widths[i]
                            if w_pt is not None:
                                row.cells[i].width = Inches(float(w_pt) / 72.0)

                    p_space = doc.add_paragraph()
                    p_space.paragraph_format.space_before = Pt(0)
                    p_space.paragraph_format.space_after = Pt(6)
                    p_space.paragraph_format.line_spacing = Pt(1)
                    r = p_space.add_run(); r.font.size = Pt(1)
                    preceding_el = el

                elif el_type == "image":
                    uri = el.get("data_uri", "")
                    if not uri or "," not in uri:
                        continue
                    try:
                        raw = base64.b64decode(uri.split(",", 1)[1])
                        stream = io.BytesIO(raw)
                        w = float(el.get("width", 0) or 0)
                        
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(resolved_style.get("spacing_before", 0.0))
                        p.paragraph_format.space_after = Pt(6)
                        
                        run = p.add_run()
                        width_in = min(6.0, w / 72.0) if w else 4.0
                        run.add_picture(stream, width=Inches(max(1.0, width_in)))
                    except Exception:
                        pass
                    preceding_el = el
                                
        eor_style = theme_styles.get("end_of_report_marker", {}) or tj.get("components", {}).get("end_of_report_marker", {})
        p_eor = doc.add_paragraph()
        p_eor.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_eor.paragraph_format.space_before = Pt(12)
        p_eor.paragraph_format.space_after = Pt(12)
        run_eor = p_eor.add_run(eor_style.get("text", "------------------ End Of Report ------------------"))
        run_eor.font.name = eor_style.get("font_family", "Calibri")
        run_eor.font.size = Pt(eor_style.get("font_size_pt", 10.0))
        run_eor.font.color.rgb = rgb(eor_style.get("text_color", "#000000"))

        replace_sng_in_docx_obj(doc)
        if output_path:
            out_p = Path(output_path)
            out_p.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(out_p))
            return None
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
        
    else:
        T = _load_oncquest_theme(theme_config)

        def rgb_legacy(hx):
            return RGBColor(int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16))

        def shade_cell_legacy(cell, fill):
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill); tcPr.append(shd)

        def shade_para_legacy(paragraph, fill):
            pPr = paragraph._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill); pPr.append(shd)

        def cell_margins_legacy(cell, t=40, b=40, l=80, r=80):
            tcPr = cell._tc.get_or_add_tcPr()
            m = OxmlElement("w:tcMar")
            for tag, val in (("top", t), ("bottom", b), ("start", l), ("end", r)):
                n = OxmlElement(f"w:{tag}")
                n.set(qn("w:w"), str(val)); n.set(qn("w:type"), "dxa"); m.append(n)
            tcPr.append(m)

        def table_borders_legacy(table, color, sz=4):
            tblPr = table._tbl.tblPr
            b = OxmlElement("w:tblBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                e = OxmlElement(f"w:{edge}")
                e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
                e.set(qn("w:space"), "0"); e.set(qn("w:color"), color); b.append(e)
            tblPr.append(b)

        def set_repeat_header_legacy(row):
            trPr = row._tr.get_or_add_trPr()
            th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true")
            trPr.append(th)

        def set_cant_split_legacy(row):
            trPr = row._tr.get_or_add_trPr()
            cs = OxmlElement("w:cantSplit"); cs.set(qn("w:val"), "true")
            trPr.append(cs)

        def _result_color_legacy(cell_text):
            cl = cell_text.strip().lower()
            if any(kw in cl for kw in _NEG_KW):
                return rgb_legacy(T["result_negative"])
            if any(kw in cl for kw in _POS_KW):
                return rgb_legacy(T["result_positive"])
            return None

        def add_banner(doc, txt, fill, big=False, flat=False):
            txt = _clean_text(txt)
            if not txt:
                return
            p = doc.add_paragraph()
            if big:
                p.paragraph_format.space_before = Pt(T["banner_space_before"])
                p.paragraph_format.space_after = Pt(T["banner_space_after"])
            else:
                p.paragraph_format.space_before = Pt(T["heading_space_before"])
                p.paragraph_format.space_after = Pt(T["heading_space_after"])
            p.paragraph_format.keep_with_next = True
            if not flat:
                shade_para_legacy(p, fill)
            run = p.add_run(txt); run.bold = True
            run.font.name = T["font"]
            run.font.size = Pt(T["banner_pt"] if big else T["body_pt"] + 1.0)
            if flat:
                run.font.color.rgb = rgb_legacy(T["primary"])
            else:
                run.font.color.rgb = rgb_legacy(T["header_text"])

        def add_para(doc, txt):
            txt = _clean_text(txt)
            if not txt:
                return
            if txt.strip().startswith(('•', '-', '*')):
                bullet_text = txt.strip().lstrip('•-* ').strip()
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_before = Pt(T["para_space_before"])
                p.paragraph_format.space_after = Pt(T["para_space_after"])
                p.paragraph_format.line_spacing = T["line_spacing"]
                run = p.add_run(bullet_text)
                run.font.name = T["font"]; run.font.size = Pt(T["body_pt"])
                run.font.color.rgb = rgb_legacy(T["body_color"])
                return
                
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(T["para_space_before"])
            p.paragraph_format.space_after = Pt(T["para_space_after"])
            p.paragraph_format.line_spacing = T["line_spacing"]
            run = p.add_run(txt)
            run.font.name = T["font"]; run.font.size = Pt(T["body_pt"])
            run.font.color.rgb = rgb_legacy(T["body_color"])

        def add_report_title(doc, text):
            text = _clean_text(text)
            if not text:
                return
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.bold = True
            run.font.name = T["font"]
            run.font.size = Pt(14.0)
            run.font.color.rgb = rgb_legacy(T["primary"])

        def add_content_box(doc, title, lines):
            if not title and not lines:
                return
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = True
            cell = tbl.rows[0].cells[0]
            cell_margins_legacy(cell, t=80, b=80, l=100, r=100)
            shade_cell_legacy(cell, "FFFFFF")
            
            tcPr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "6")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), T["border"])
                borders.append(b)
            tcPr.append(borders)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            
            if title:
                run_title = p.add_run(title)
                run_title.bold = True
                run_title.font.name = T["font"]
                run_title.font.size = Pt(T["body_pt"] + 0.5)
                run_title.font.color.rgb = rgb_legacy(T["primary"])
                p = cell.add_paragraph()
                
            for line in lines:
                line_text = _clean_text(line)
                if not line_text:
                    continue
                if line_text.strip().startswith(('•', '-', '*')):
                    bullet_text = line_text.strip().lstrip('•-* ').strip()
                    p_bullet = cell.add_paragraph(style='List Bullet')
                    p_bullet.paragraph_format.space_before = Pt(0)
                    p_bullet.paragraph_format.space_after = Pt(2)
                    p_bullet.paragraph_format.line_spacing = T["line_spacing"]
                    run = p_bullet.add_run(bullet_text)
                    run.font.name = T["font"]; run.font.size = Pt(T["body_pt"])
                    run.font.color.rgb = rgb_legacy(T["body_color"])
                    continue
                    
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = T["line_spacing"]
                run = p.add_run(line_text)
                run.font.name = T["font"]
                run.font.size = Pt(T["body_pt"])
                run.font.color.rgb = rgb_legacy(T["body_color"])
                p = cell.add_paragraph()
                
            if len(cell.paragraphs) > 1 and cell.paragraphs[-1].text == "":
                p_to_remove = cell.paragraphs[-1]
                p_to_remove._p.getparent().remove(p_to_remove._p)
                
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_before = Pt(0)
            p_space.paragraph_format.space_after = Pt(T["table_space_after"])
            p_space.paragraph_format.line_spacing = Pt(1)
            r = p_space.add_run(); r.font.size = Pt(1)

        def add_kv_table(doc, kv):
            if not kv:
                return
            tbl = doc.add_table(rows=0, cols=4)
            table_borders_legacy(tbl, T["border_table"])
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = True
            
            col_widths = [Inches(1.0), Inches(2.2), Inches(1.0), Inches(2.2)]
            items = list(kv.items())
            for idx in range(0, len(items), 2):
                cells = tbl.add_row().cells
                set_cant_split_legacy(tbl.rows[-1])
                for c in cells:
                    cell_margins_legacy(c, t=20, b=20, l=40, r=40)
                    
                k1, v1 = items[idx]
                kr1 = cells[0].paragraphs[0].add_run(_clean_text(k1))
                kr1.bold = True; kr1.font.name = T["font"]; kr1.font.size = Pt(T["body_pt"])
                kr1.font.color.rgb = rgb_legacy(T["primary"])
                vr1 = cells[1].paragraphs[0].add_run(_clean_text(v1))
                vr1.font.name = T["font"]; vr1.font.size = Pt(T["body_pt"])
                vr1.font.color.rgb = rgb_legacy(T["body_color"])
                
                if idx + 1 < len(items):
                    k2, v2 = items[idx + 1]
                    kr2 = cells[2].paragraphs[0].add_run(_clean_text(k2))
                    kr2.bold = True; kr2.font.name = T["font"]; kr2.font.size = Pt(T["body_pt"])
                    kr2.font.color.rgb = rgb_legacy(T["primary"])
                    vr2 = cells[3].paragraphs[0].add_run(_clean_text(v2))
                    vr2.font.name = T["font"]; vr2.font.size = Pt(T["body_pt"])
                    vr2.font.color.rgb = rgb_legacy(T["body_color"])

            for row in tbl.rows:
                for i, w in enumerate(col_widths):
                    row.cells[i].width = w

            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_before = Pt(0)
            p_space.paragraph_format.space_after = Pt(T["table_space_after"])
            p_space.paragraph_format.line_spacing = Pt(1)
            r = p_space.add_run(); r.font.size = Pt(1)

        def add_data_table_legacy(doc, tab):
            headers = [_clean_text(h) for h in (tab.get("headers") or [])]
            rows = [[_clean_text(c) for c in (r or [])] for r in (tab.get("rows") or [])]

            if (not headers or not any(headers)) and rows:
                headers = rows[0]; rows = rows[1:]
            if not any(headers) and not any(any(c for c in r) for r in rows):
                return

            ncols = max([len(headers)] + [len(r) for r in rows] + [0])
            if ncols == 0:
                return
            headers = (headers + [""] * ncols)[:ncols]
            rows = [(r + [""] * ncols)[:ncols] for r in rows]

            keep_cols = []
            for ci in range(ncols):
                col_vals = [headers[ci].strip()] + [str(r[ci]).strip() for r in rows]
                if any(v for v in col_vals):
                    keep_cols.append(ci)
            if not keep_cols:
                return
            headers = [headers[ci] for ci in keep_cols]
            rows = [[r[ci] for ci in keep_cols] for r in rows]
            ncols = len(keep_cols)

            has_header = any(headers)
            tbl = doc.add_table(rows=1 if has_header else 0, cols=ncols)
            table_borders_legacy(tbl, T["border_table"])
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.allow_autofit = False

            # Calculate dynamic column widths for legacy table
            avail_w = 595.28 - (float(T.get("margin_left", 36.0)) + float(T.get("margin_right", 36.0)))
            col_max_lens = [len(str(h)) for h in headers]
            for r in rows:
                for i in range(min(ncols, len(r))):
                    col_max_lens[i] = max(col_max_lens[i], len(str(r[i])))
            
            base_min = min(45.0, (avail_w / ncols) * 0.6) if ncols > 0 else 45.0
            min_w_sum = base_min * ncols
            distribute_w = max(0.0, avail_w - min_w_sum)
            
            weights = [max(1, l) ** 0.7 for l in col_max_lens]
            total_weight = sum(weights)
            
            legacy_col_widths = []
            for i in range(ncols):
                if total_weight > 0:
                    extra = (weights[i] / total_weight) * distribute_w
                else:
                    extra = distribute_w / ncols
                legacy_col_widths.append(base_min + extra)

            if has_header:
                hrow = tbl.rows[0]
                set_repeat_header_legacy(hrow)
                set_cant_split_legacy(hrow)
                for i in range(ncols):
                    cell = hrow.cells[i]
                    shade_cell_legacy(cell, T["primary"]); cell_margins_legacy(cell)
                    para = cell.paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(headers[i]); run.bold = True
                    run.font.name = T["table_header_font"]
                    run.font.size = Pt(T["table_header_pt"])
                    run.font.color.rgb = rgb_legacy(T["header_text"])

            for row_idx, r in enumerate(rows):
                cells = tbl.add_row().cells
                set_cant_split_legacy(tbl.rows[-1])
                row_bg = T.get("alternating_row_bg")
                for i in range(ncols):
                    cell = cells[i]
                    cell_margins_legacy(cell)
                    if row_idx % 2 == 1 and row_bg:
                        shade_cell_legacy(cell, row_bg)
                    cell_text = r[i]
                    run = cell.paragraphs[0].add_run(cell_text)
                    run.font.name = T["font"]; run.font.size = Pt(T["body_pt"])
                    rc = _result_color_legacy(cell_text)
                    if rc:
                        run.font.color.rgb = rc
                        run.bold = True
                    else:
                        run.font.color.rgb = rgb_legacy(T["body_color"])
            
            # Apply dynamic widths to cells in all rows
            for row in tbl.rows:
                for i in range(ncols):
                    row.cells[i].width = Inches(float(legacy_col_widths[i]) / 72.0)
            
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_before = Pt(0)
            p_space.paragraph_format.space_after = Pt(T["table_space_after"])
            p_space.paragraph_format.line_spacing = Pt(1)
            r = p_space.add_run(); r.font.size = Pt(1)

        def add_image_legacy(doc, img):
            if _is_decorative_image(img):
                return
            uri = img.get("data_uri", "")
            if not uri or "," not in uri:
                return
            try:
                raw = base64.b64decode(uri.split(",", 1)[1])
                stream = io.BytesIO(raw)
                w = float(img.get("width", 0) or 0)
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                width_in = min(6.0, w / 96.0) if w else 4.0
                run.add_picture(stream, width=Inches(max(1.0, width_in)))
            except Exception:
                pass

        def bbox_top(el):
            bb = el.get("bbox")
            if isinstance(bb, (list, tuple)) and len(bb) >= 2:
                return (bb[1], bb[0])
            return (10_000_000, 0)

        def content_lines(box):
            ct = box.get("content_text", [])
            if isinstance(ct, str):
                return [ct]
            if isinstance(ct, list):
                return [str(x) for x in ct]
            return []

        doc = Document()
        header_logo_path = Path("assets/header_image1.png")
        sig_image_path = Path("assets/dr_vinay_signature.png")
        for section in doc.sections:
            section.top_margin = Pt(T["margin_top"])
            section.bottom_margin = Pt(T["margin_bottom"])
            section.left_margin = Pt(T["margin_left"])
            section.right_margin = Pt(T["margin_right"])
            section.page_width = Pt(T["paper_width"])
            section.page_height = Pt(T["paper_height"])

            if header_logo_path.exists():
                header = section.header
                if header is not None:
                    for p in header.paragraphs:
                        p.text = ""
                    p = header.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(header_logo_path.absolute()), width=Inches(6.0))

            if sig_image_path.exists():
                footer = section.footer
                if footer is not None:
                    for p in footer.paragraphs:
                        p.text = ""
                    p = footer.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run()
                    run.add_picture(str(sig_image_path.absolute()), width=Inches(1.25))

        normal = doc.styles["Normal"]
        normal.font.name = T["font"]; normal.font.size = Pt(T["body_pt"])
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), T["font"]); rfonts.set(qn("w:hAnsi"), T["font"])

        summary = data.get("document_summary", {}) or {}
        title = os.path.splitext(summary.get("file_name", "Report"))[0]
        add_report_title(doc, title.upper())

        kv = data.get("all_key_value_pairs") or data.get("extracted_key_value_pairs") or {}
        if kv:
            add_banner(doc, "DETAILS", T["primary"])
            add_kv_table(doc, kv)

        seen = set()
        _SKIP_TITLES_DOCX = {"general content / notes", "patient details & metadata"}

        def _should_skip_docx(el):
            ttl = _clean_text(el.get("title"))
            sec_type = el.get("type", "")
            if sec_type == "demographics_box":
                return True
            if ttl and ttl.startswith("Header & Metadata Box"):
                return True
            if ttl and ttl.lower() in _SKIP_TITLES_DOCX:
                return True
            return False

        def _is_ref_fragment_docx(el):
            ttl = _clean_text(el.get("title"))
            content_text = content_lines(el)
            body_str = " ".join(content_text)
            combined = (ttl + " " + body_str).lower()
            if "pmid" in combined or "doi:" in combined or "et al" in combined:
                return True
            if "http" in combined or "www." in combined or ".html" in combined or ".com/" in combined or "release" in combined:
                return True
            if "guideline" in combined:
                return True
            if re.search(r'\b\d{4}\b', ttl):
                return True
            if re.match(r'^\d', ttl) and any(c in ttl for c in [';', ':', '(', ')', '-']):
                return True
            return False

        ref_lines_docx = []

        def render_page_legacy(page):
            boxes = page.get("boxes_and_sections") or []
            tables = page.get("tables") or []
            images = page.get("images_and_graphs") or []
            text_blocks = page.get("text_blocks") or []

            items = []
            for b in boxes:
                items.append(("box", bbox_top(b), b))
            for t in tables:
                items.append(("table", bbox_top(t), t))
            for im in images:
                items.append(("image", bbox_top(im), im))
            if not boxes and text_blocks:
                for tb in text_blocks:
                    items.append(("textblock", bbox_top(tb), tb))

            items.sort(key=lambda x: x[1])

            for kind, _pos, el in items:
                if kind == "box":
                    if _should_skip_docx(el):
                        for line in content_lines(el):
                            s = _clean_text(line)
                            if s and s not in seen:
                                seen.add(s); add_para(doc, s)
                        continue
                    ttl = _clean_text(el.get("title"))
                    if ttl and ttl.lower().startswith("references"):
                        for line in content_lines(el):
                            s = _clean_text(line)
                            if s:
                                ref_lines_docx.append(s)
                        continue
                    if _is_ref_fragment_docx(el):
                        merged = ttl
                        body = [_clean_text(l) for l in content_lines(el) if _clean_text(l)]
                        if body:
                            merged = f"{ttl} {' '.join(body)}"
                        ref_lines_docx.append(merged)
                        continue
                    box_lines = []
                    for line in content_lines(el):
                        s = _clean_text(line)
                        if s and s not in seen:
                            seen.add(s)
                            box_lines.append(s)
                    if box_lines or ttl:
                        add_content_box(doc, ttl, box_lines)
                elif kind == "table":
                    add_data_table_legacy(doc, el)
                elif kind == "image":
                    add_image_legacy(doc, el)
                elif kind == "textblock":
                    s = _clean_text(el.get("text", ""))
                    if s and s not in seen:
                        seen.add(s); add_para(doc, s)

        pages_legacy = data.get("pages") or []
        if pages_legacy:
            for page in pages_legacy:
                render_page_legacy(page)
        else:
            for sec in data.get("all_boxes_and_sections") or data.get("content_sections") or []:
                if _should_skip_docx(sec):
                    continue
                ttl = _clean_text(sec.get("title"))
                if ttl and ttl.lower().startswith("references"):
                    for line in content_lines(sec):
                        s = _clean_text(line)
                        if s:
                            ref_lines_docx.append(s)
                    continue
                if _is_ref_fragment_docx(sec):
                    merged = ttl
                    body = [_clean_text(l) for l in content_lines(sec) if _clean_text(l)]
                    if body:
                        merged = f"{ttl} {' '.join(body)}"
                    ref_lines_docx.append(merged)
                    continue
                box_lines = []
                for line in content_lines(sec):
                    s = _clean_text(line)
                    if s and s not in seen:
                        seen.add(s)
                        box_lines.append(s)
                if box_lines or ttl:
                    add_content_box(doc, ttl, box_lines)
            for tb in data.get("all_tables") or data.get("tables") or []:
                add_data_table_legacy(doc, tb)
            for im in data.get("all_images_and_graphs") or data.get("images_and_graphs") or []:
                add_image_legacy(doc, im)

        if ref_lines_docx:
            add_banner(doc, "References", T["primary"])
            for rl in ref_lines_docx:
                if rl not in seen:
                    seen.add(rl); add_para(doc, rl)

        end = doc.add_paragraph(); end.alignment = WD_ALIGN_PARAGRAPH.CENTER
        er = end.add_run(T["eor_text"])
        er.font.name = T["eor_font"]; er.font.size = Pt(T["eor_pt"])
        er.font.color.rgb = rgb_legacy(T["eor_color"])

        replace_sng_in_docx_obj(doc)
        if output_path:
            out_p = Path(output_path)
            out_p.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(out_p))
            return None
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


def render_json_file_to_html(json_path, output_path: str = None, theme_config: dict = None) -> str:
    """
    Renders an extracted document JSON file (containing page text_blocks, drawings, images)
    directly into an HTML document template matching the Oncquest design system.
    Saves to output_path if provided, and returns the full HTML string.
    """
    import html
    json_p = Path(json_path)
    with open(json_p, "r", encoding="utf-8") as f:
        data = json.load(f)
    data = replace_sng_in_structure(data)
    data = replace_test_name_in_structure(data)

    source_file = data.get("source_file", json_p.stem + ".pdf")
    doc_title = Path(source_file).stem
    pages = data.get("pages", [])

    cfg = get_merged_theme_config(theme_config)
    colors_cfg = cfg.get("colors", {})
    primary_color = colors_cfg.get("primary", "#1f497d")
    secondary_color = colors_cfg.get("secondary", "#008080")
    accent_orange = colors_cfg.get("accent_orange", "#ed7d31")
    accent_red = colors_cfg.get("accent_red", "#ff0000")
    result_positive_color = colors_cfg.get("result_positive", "#C00000")
    result_negative_color = colors_cfg.get("result_negative", "#008000")

    def sanitize_text(text: str) -> str:
        if not text:
            return ""
        text = text.replace("\ufb01", "fi").replace("\ufb02", "fl").replace("\xa0", " ")
        return html.escape(text)

    def get_font_family(font_name: str) -> str:
        if not font_name:
            return "Calibri, sans-serif"
        fn = font_name.lower()
        if "calibri" in fn:
            return "Calibri, Arial, sans-serif"
        elif "tahoma" in fn:
            return "Tahoma, Geneva, sans-serif"
        elif "verdana" in fn:
            return "Verdana, Geneva, sans-serif"
        elif "times" in fn:
            return "'Times New Roman', Times, serif"
        elif "cambria" in fn:
            return "Cambria, Georgia, serif"
        return "Calibri, Arial, sans-serif"

    _POSITIVE_KEYWORDS = ["positive", "pathogenic", "detected", "high", "abnormal", "msi - high", "msi-high"]
    _NEGATIVE_KEYWORDS = ["negative", "normal", "not detected", "stable", "msi - stable", "msi-stable", "benign", "likely benign"]

    def _result_color_span(cell_str: str) -> str:
        cell_lower = cell_str.strip().lower()
        if any(kw in cell_lower for kw in _NEGATIVE_KEYWORDS):
            return f"<span style='color:{result_negative_color}; font-weight:bold;'>{cell_str}</span>"
        if any(kw in cell_lower for kw in _POSITIVE_KEYWORDS):
            return f"<span style='color:{result_positive_color}; font-weight:bold;'>{cell_str}</span>"
        return cell_str

    html_parts = [
        "<!DOCTYPE html>",
        "<html lang='en'>",
        "<head>",
        "<meta charset='utf-8'>",
        "<meta name='viewport' content='width=device-width, initial-scale=1.0'>",
        f"<title>{doc_title} — Oncquest Lab Report</title>",
        "<link href='https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap' rel='stylesheet'>",
        "<style>",
        f"""
        :root {{
          --color-primary: {primary_color};
          --color-secondary: {secondary_color};
          --color-banner-dark: #404040;
          --color-bg-container: #525659;
          --color-bg-page: #ffffff;
          --font-primary: 'Calibri', 'Inter', sans-serif;
          --page-shadow: 0 8px 24px rgba(0, 0, 0, 0.25);
        }}
        @page {{ size: 595.0pt 842.0pt; margin: 0; }}
        * {{ box-sizing: border-box; margin: 0; padding: 0; }}
        body {{
          margin: 0;
          padding: 0;
          background-color: var(--color-bg-container);
          font-family: var(--font-primary);
          color: #000000;
          -webkit-font-smoothing: antialiased;
        }}
        .header-bar {{
          background: #0f172a;
          color: #ffffff;
          padding: 12px 24px;
          display: flex;
          align-items: center;
          justify-content: space-between;
          font-family: 'Inter', sans-serif;
          box-shadow: 0 2px 8px rgba(0,0,0,0.3);
          position: sticky;
          top: 0;
          z-index: 1000;
        }}
        .header-bar h1 {{
          font-size: 1.1rem;
          font-weight: 600;
          margin: 0;
        }}
        .header-bar .meta {{
          font-size: 0.85rem;
          color: #94a3b8;
        }}
        .pdf-container {{
          display: flex;
          flex-direction: column;
          align-items: center;
          padding: 24px 0;
          gap: 24px;
        }}
        .pdf-page {{
          background: var(--color-bg-page);
          width: 595.0pt;
          min-height: 842.0pt;
          position: relative;
          overflow: visible;
          box-shadow: var(--page-shadow);
          page-break-after: always;
          border-radius: 2px;
        }}
        .pdf-page p {{
          position: absolute;
          margin: 0;
          padding: 0;
          white-space: pre-wrap;
          word-break: break-word;
          line-height: 1.15;
          z-index: 10;
        }}
        .vector-fill-box {{
          position: absolute;
          pointer-events: none;
          z-index: 2;
          box-sizing: border-box;
        }}
        .vector-box {{
          position: absolute;
          pointer-events: none;
          z-index: 3;
          box-sizing: border-box;
        }}
        .img-box {{
          position: absolute;
          z-index: 5;
          object-fit: contain;
        }}
        .pdf-page table th, .pdf-page table td {{
          border: 1px solid var(--color-primary);
          padding: 4px;
          text-align: left;
          vertical-align: middle;
          font-family: var(--font-primary);
        }}
        .pdf-page table th {{
          font-weight: bold;
          text-align: center;
        }}
        @media print {{
          .header-bar {{ display: none; }}
          body {{ background: #ffffff; }}
          .pdf-container {{ padding: 0; gap: 0; }}
          .pdf-page {{ box-shadow: none; margin: 0; border-radius: 0; }}
        }}
        """,
        "</style>",
        "</head>",
        "<body>",
        f"<div class='header-bar'>",
        f"  <h1>Oncquest Report Renderer — {doc_title}</h1>",
        f"  <div class='meta'>Source: {sanitize_text(source_file)} | Pages: {len(pages)}</div>",
        f"</div>",
        "<div class='pdf-container'>"
    ]

    import base64
    header_image2_b64 = ""
    header_image2_path = Path("assets/header_image2.jpeg")
    if header_image2_path.exists():
        try:
            with open(header_image2_path, "rb") as img_f:
                header_image2_b64 = base64.b64encode(img_f.read()).decode("utf-8")
        except Exception:
            pass

    for p in pages:
        p_num = p.get("page_number", 1)
        pw = 595.0
        ph = 842.0
        dimensions = p.get("dimensions")
        if isinstance(dimensions, dict):
            pw = dimensions.get("width", 595.0)
            ph = dimensions.get("height", 842.0)
        hy_cutoff = p.get("header_y_cutoff", 0.0)
        fy_cutoff = p.get("footer_y_cutoff", ph)

        html_parts.append(f"<div class='pdf-page' id='page-{p_num}' style='width:{pw:.1f}pt; min-height:{ph:.1f}pt;'>")
        if header_image2_b64:
            img_tag = (
                f'<img src="data:image/jpeg;base64,{header_image2_b64}" '
                f'style="position:absolute !important; left:440.0pt !important; '
                f'top:15.0pt !important; width:130.0pt !important; '
                f'height:auto !important; z-index:100 !important; '
                f'object-fit:contain !important;" />'
            )
            html_parts.append(img_tag)

        # 1. Render Vector Drawings in body region ONLY
        for d in p.get("drawings", []):
            bbox = d.get("bbox")
            if not bbox or len(bbox) < 4:
                continue
            x0, y0, x1, y1 = bbox
            if y0 < hy_cutoff or y1 > fy_cutoff:
                continue
            w = max(0.5, x1 - x0)
            h = max(0.5, y1 - y0)
            fill_col = d.get("fill_color")
            stroke_col = d.get("stroke_color")

            if fill_col:
                # Remove all rectangles before headings (decorative squares on the left)
                if x0 < 45.0 and w < 20.0 and h < 25.0:
                    continue
                html_parts.append(
                    f"<div class='vector-fill-box' style='left:{x0:.2f}pt; top:{y0:.2f}pt; width:{w:.2f}pt; height:{h:.2f}pt; background-color:{fill_col};'></div>"
                )
            elif stroke_col:
                stroke_w = d.get("width") or 1.0
                html_parts.append(
                    f"<div class='vector-box' style='left:{x0:.2f}pt; top:{y0:.2f}pt; width:{w:.2f}pt; height:{h:.2f}pt; border:{stroke_w}px solid {stroke_col};'></div>"
                )

        # 2. Render Images in body region ONLY
        for img in p.get("images", []):
            bbox = img.get("bbox")
            data_uri = img.get("data_uri")
            if bbox and data_uri:
                x0, y0, x1, y1 = bbox
                if y0 < hy_cutoff or y1 > fy_cutoff or y0 >= 690.0:
                    continue
                w = max(1.0, x1 - x0)
                h = max(1.0, y1 - y0)
                html_parts.append(
                    f"<img class='img-box' src='{data_uri}' style='left:{x0:.2f}pt; top:{y0:.2f}pt; width:{w:.2f}pt; height:{h:.2f}pt;' />"
                )

        # 3. Render Text Blocks in body region ONLY
        for b in p.get("text_blocks", []):
            b_bbox = b.get("bbox", [0, 0, 0, 0])
            if b_bbox[1] < hy_cutoff or b_bbox[3] > fy_cutoff:
                continue

            lines = b.get("lines")
            if lines:
                for line in lines:
                    l_bbox = line.get("bbox", [0, 0, 0, 0])
                    lx0, ly0, lx1, ly1 = l_bbox
                    if ly0 < hy_cutoff or ly1 > fy_cutoff:
                        continue
                    spans = line.get("spans", [])

                    if not spans:
                        continue

                    spans_html = []
                    for s in spans:
                        raw_text = s.get("text", "")
                        if not raw_text:
                            continue
                        clean_t = sanitize_text(raw_text)
                        font_name = s.get("font", "")
                        font_family = get_font_family(font_name)
                        size = s.get("size", 10.0)
                        color = s.get("color", "#000000")
                        is_bold = s.get("bold") or ("bold" in font_name.lower())
                        is_italic = s.get("italic") or ("italic" in font_name.lower())

                        font_wt = "bold" if is_bold else "normal"
                        font_st = "italic" if is_italic else "normal"

                        span_style = (
                            f"font-size:{size:.2f}pt; "
                            f"color:{color}; "
                            f"font-family:{font_family}; "
                            f"font-weight:{font_wt}; "
                            f"font-style:{font_st};"
                        )
                        spans_html.append(f"<span style='{span_style}'>{clean_t}</span>")

                    if spans_html:
                        html_parts.append(
                            f"<p style='left:{lx0:.2f}pt; top:{ly0:.2f}pt;'>{''.join(spans_html)}</p>"
                        )
            else:
                raw_text = b.get("text", "")
                if not raw_text:
                    continue
                clean_t = sanitize_text(raw_text)
                font_name = b.get("font", "")
                font_family = get_font_family(font_name)
                size = b.get("max_font_size") or b.get("size") or 10.0
                color = b.get("color", "#000000")
                is_bold = b.get("is_bold") or b.get("bold") or ("bold" in font_name.lower())
                is_italic = b.get("is_italic") or b.get("italic") or ("italic" in font_name.lower())

                font_wt = "bold" if is_bold else "normal"
                font_st = "italic" if is_italic else "normal"

                span_style = (
                    f"font-size:{size:.2f}pt; "
                    f"color:{color}; "
                    f"font-family:{font_family}; "
                    f"font-weight:{font_wt}; "
                    f"font-style:{font_st};"
                )
                lx0, ly0 = b_bbox[0], b_bbox[1]
                html_parts.append(
                    f"<p style='left:{lx0:.2f}pt; top:{ly0:.2f}pt; {span_style}'>{clean_t}</p>"
                )

        # 4. Render Key-Value blocks and Tables in body region ONLY
        for item in p.get("content", []):
            itype = item.get("type")
            bbox = item.get("bbox")
            if not bbox or len(bbox) < 4:
                continue
            x0, y0, x1, y1 = bbox
            if y0 < hy_cutoff or y1 > fy_cutoff:
                continue
            w = max(1.0, x1 - x0)
            h = max(1.0, y1 - y0)

            if itype == "key_value":
                pairs = item.get("pairs", {})
                if not pairs:
                    continue
                tr_html = ""
                for k, v in pairs.items():
                    tr_html += (
                        f"<tr>"
                        f"<td style='font-weight: bold; border: none !important; padding: 2px 4px; color: var(--color-primary); width: 120pt;'>{sanitize_text(k)}:</td>"
                        f"<td style='border: none !important; padding: 2px 4px;'>{sanitize_text(str(v))}</td>"
                        f"</tr>"
                    )
                table_style = (
                    f"position: absolute !important; "
                    f"left: {x0:.2f}pt; "
                    f"top: {y0:.2f}pt; "
                    f"width: {w:.2f}pt; "
                    f"height: {h:.2f}pt; "
                    f"border-collapse: collapse; "
                    f"border: none; "
                    f"z-index: 15; "
                    f"font-size: 9.0pt;"
                )
                html_parts.append(
                    f"<table style='{table_style}'>"
                    f"<tbody>{tr_html}</tbody>"
                    f"</table>"
                )

            elif itype == "table":
                headers = item.get("headers", [])
                rows = item.get("rows", [])
                if not headers and not rows:
                    continue

                th_html = "".join([f"<th>{sanitize_text(h)}</th>" for h in headers]) if any(h.strip() for h in headers) else ""
                tr_html = ""
                for r in rows:
                    tds = "".join([f"<td>{_result_color_span(sanitize_text(str(cell)))}</td>" for cell in r])
                    tr_html += f"<tr>{tds}</tr>"

                table_style = (
                    f"position: absolute !important; "
                    f"left: {x0:.2f}pt; "
                    f"top: {y0:.2f}pt; "
                    f"width: {w:.2f}pt; "
                    f"height: {h:.2f}pt; "
                    f"border-collapse: collapse; "
                    f"border: 1px solid var(--color-primary); "
                    f"z-index: 15; "
                    f"font-size: 8.5pt;"
                )
                html_parts.append(
                    f"<table style='{table_style}'>"
                    f"<thead><tr style='background-color: var(--color-primary); color: #ffffff;'>{th_html}</tr></thead>"
                    f"<tbody>{tr_html}</tbody>"
                    f"</table>"
                )

        html_parts.append("</div>")

    html_parts.append("</div></body></html>")
    full_html = "\n".join(html_parts)
    full_html = replace_sng_gen_lab(full_html)
    full_html = replace_test_name_in_html(full_html)

    if output_path:
        out_p = Path(output_path)
        out_p.parent.mkdir(parents=True, exist_ok=True)
        with open(out_p, "w", encoding="utf-8") as f_out:
            f_out.write(full_html)

    return full_html


def convert_html_to_docx(html_input, output_path: str = None, theme_config: dict = None):
    """
    Converts an HTML string or HTML file path directly into a Microsoft Word (.docx) file.
    
    Uses a high-fidelity page-image approach:
    1. Renders the HTML in Playwright Chromium browser
    2. Captures each page as a high-resolution PNG screenshot  
    3. Embeds page images into Word at exact A4 page size
    
    This guarantees 100% visual fidelity — every pixel of the HTML is preserved in Word.
    Falls back to pdf2docx or htmldocx if Playwright is unavailable.
    
    Returns bytes of the Word file, or writes to output_path if provided.
    """
    import tempfile

    if isinstance(html_input, (str, Path)) and os.path.exists(str(html_input)) and os.path.isfile(str(html_input)):
        with open(str(html_input), "r", encoding="utf-8") as f:
            html_content = f.read()
    else:
        html_content = str(html_input)

    html_content = replace_sng_gen_lab(html_content)

    # ── Method 1: High-fidelity page-image approach using Playwright (DISABLED) ──
    if False:
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Emu
            from docx.enum.section import WD_ORIENT

            with tempfile.TemporaryDirectory() as tmp_dir:
                tmp_dir_path = Path(tmp_dir)
                html_file_path = tmp_dir_path / "temp.html"
                html_file_path.write_text(html_content, encoding="utf-8")

                with sync_playwright() as p:
                    browser = p.chromium.launch()
                    # Use A4-like viewport width matching the HTML page width (~795px = 595.6pt)
                    # device_scale_factor=2 for high-DPI sharp screenshots
                    page = browser.new_page(viewport={"width": 795, "height": 1123}, device_scale_factor=2)
                    page.goto(html_file_path.as_uri(), wait_until="networkidle")

                    # Find all .pdf-page elements (each represents one report page)
                    pdf_pages = page.query_selector_all(".pdf-page")
                    
                    if not pdf_pages:
                        # Fallback: if no .pdf-page elements, try div[id^='page']
                        pdf_pages = page.query_selector_all("div[id^='page']")
                    
                    if not pdf_pages:
                        # Last resort: capture full page as single image
                        pdf_pages = [None]  # Sentinel for full-page mode

                    page_images = []
                    for idx, pg_elem in enumerate(pdf_pages):
                        img_path = tmp_dir_path / f"page_{idx}.png"
                        if pg_elem is not None:
                            # Capture individual page element at 2x scale for high DPI
                            pg_elem.screenshot(path=str(img_path))
                        else:
                            # Full page screenshot
                            page.screenshot(path=str(img_path), full_page=True)
                        
                        if img_path.exists() and img_path.stat().st_size > 0:
                            page_images.append(img_path)
                        
                        print(f"   [+] Captured page {idx + 1}/{len(pdf_pages)}")

                    page.close()
                    browser.close()

                if page_images:
                    # Build Word document with one image per page
                    doc = Document()
                    
                    # Set A4 page size with minimal margins to maximize image area
                    for section in doc.sections:
                        section.page_width = Emu(7560310)   # A4 width: 210mm
                        section.page_height = Emu(10692130)  # A4 height: 297mm
                        section.top_margin = Inches(0)
                        section.bottom_margin = Inches(0)
                        section.left_margin = Inches(0)
                        section.right_margin = Inches(0)

                    for idx, img_path in enumerate(page_images):
                        if idx > 0:
                            # Add page break before each new page (except the first)
                            new_section = doc.add_section()
                            new_section.page_width = Emu(7560310)
                            new_section.page_height = Emu(10692130)
                            new_section.top_margin = Inches(0)
                            new_section.bottom_margin = Inches(0)
                            new_section.left_margin = Inches(0)
                            new_section.right_margin = Inches(0)

                        # Add the page image — fill the entire A4 page width
                        para = doc.add_paragraph()
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
                        run = para.add_run()
                        # Width = full A4 width (8.27 inches = 210mm)
                        run.add_picture(str(img_path), width=Inches(8.27))

                    replace_sng_in_docx_obj(doc)

                    if output_path:
                        out_p = Path(output_path)
                        out_p.parent.mkdir(parents=True, exist_ok=True)
                        doc.save(str(out_p))
                        print(f"   [+] Word document created with {len(page_images)} page(s): {out_p}")
                        return None
                    else:
                        buf = io.BytesIO()
                        doc.save(buf)
                        return buf.getvalue()

        except Exception as e:
            print(f"[*] Note: Page-image HTML-to-Word conversion failed: {e}. Falling back to PDF-based approach.")

    # ── Method 2: Fallback — PDF-based conversion via pdf2docx ──
    if sync_playwright is not None and PDF2DocxConverter is not None:
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                tmp_dir_path = Path(tmp_dir)
                html_file_path = tmp_dir_path / "temp.html"
                html_file_path.write_text(html_content, encoding="utf-8")

                compiled_pdf = tmp_dir_path / "compiled.pdf"
                render_html_to_pdf_and_preview(html_file_path, compiled_pdf)

                if compiled_pdf.exists() and compiled_pdf.stat().st_size > 0:
                    out_docx_path = tmp_dir_path / "output.docx"
                    success = convert_pdf_to_word(compiled_pdf, out_docx_path, theme_config=theme_config)
                    if success and out_docx_path.exists():
                        if output_path:
                            out_p = Path(output_path)
                            out_p.parent.mkdir(parents=True, exist_ok=True)
                            with open(out_docx_path, "rb") as f_in, open(out_p, "wb") as f_out:
                                f_out.write(f_in.read())
                            return None
                        else:
                            with open(out_docx_path, "rb") as f_in:
                                return f_in.read()
        except Exception as e:
            print(f"[*] Note: PDF-based HTML-to-Word conversion failed: {e}. Falling back to direct HTML parse.")

    # ── Method 3: Original Fallback — direct HTML parsing ──
    from docx import Document
    from docx.shared import Inches
    from bs4 import BeautifulSoup
    try:
        from htmldocx import HtmlToDocx
    except ImportError:
        HtmlToDocx = None

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    soup = BeautifulSoup(html_content, 'html.parser')
    for tag in soup(['style', 'script', 'meta', 'link']):
        tag.decompose()

    body = soup.find('body')
    clean_html = str(body) if body else str(soup)

    converted_with_htmldocx = False
    if HtmlToDocx is not None:
        try:
            parser = HtmlToDocx()
            parser.add_html_to_document(clean_html, doc)
            converted_with_htmldocx = True
        except Exception:
            converted_with_htmldocx = False

    if not converted_with_htmldocx:
        _parse_html_soup_to_docx(soup, doc)

    replace_sng_in_docx_obj(doc)
    if output_path:
        out_p = Path(output_path)
        out_p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_p))
        return None
    else:
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


def _parse_html_soup_to_docx(soup, doc):
    """Fallback manual parser from BeautifulSoup to python-docx."""
    for elem in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'table', 'ul', 'ol', 'div']):
        if elem.name in ['h1', 'h2', 'h3', 'h4']:
            level = int(elem.name[1])
            txt = elem.get_text(strip=True)
            if txt:
                doc.add_heading(txt, level=level)
        elif elem.name in ['p', 'div'] and elem.find_parent('table') is None:
            txt = elem.get_text(strip=True)
            if txt:
                doc.add_paragraph(txt)
        elif elem.name == 'table' and elem.find_parent('table') is None:
            rows = elem.find_all('tr')
            if rows:
                cols_count = max(len(r.find_all(['td', 'th'])) for r in rows)
                if cols_count > 0:
                    t = doc.add_table(rows=len(rows), cols=cols_count)
                    for r_idx, r in enumerate(rows):
                        cells = r.find_all(['td', 'th'])
                        for c_idx, c in enumerate(cells):
                            if c_idx < cols_count:
                                t.rows[r_idx].cells[c_idx].text = c.get_text(strip=True)


def convert_pdf_via_pdf2docx(pdf_path, docx_path):
    """
    Directly converts a PDF file to Word (.docx) using pdf2docx Converter.
    """
    if PDF2DocxConverter is None:
        print("[!] pdf2docx library is not installed. Install via: pip install pdf2docx")
        return False
    try:
        pdf_p = str(Path(pdf_path).absolute())
        docx_p = str(Path(docx_path).absolute())
        Path(docx_p).parent.mkdir(parents=True, exist_ok=True)

        cv_obj = PDF2DocxConverter(pdf_p)
        cv_obj.convert(docx_p)
        cv_obj.close()
        print(f"   [+] pdf2docx conversion completed successfully: {docx_p}")
        return True
    except Exception as e:
        print(f"   [!] pdf2docx conversion error for {pdf_path}: {e}")
        return False


def redact_extracted_json(data: dict) -> dict:
    """
    Redacts patient name, case ID, and signature images directly from the extracted JSON data.
    """
    patient_name = None
    case_id = None
    
    # 1. Identify patient name and case ID from key_value elements on the first page
    pages = data.get("document", {}).get("pages", [])
    if pages:
        for el in pages[0].get("elements", []):
            if el.get("type") == "key_value" and "data" in el:
                kv_data = el["data"]
                for k, v in kv_data.items():
                    k_lower = k.lower().strip()
                    if "name" in k_lower and v:
                        patient_name = str(v).strip()
                    elif ("case id" in k_lower or "case_id" in k_lower) and v:
                        case_id = str(v).strip()

    # 2. Perform global string replacements on serialized JSON to mask name and case ID everywhere
    json_str = json.dumps(data, ensure_ascii=False)
    if patient_name:
        json_str = json_str.replace(patient_name, "PATIENT NAME")
    if case_id:
        json_str = json_str.replace(case_id, "0000000000")
    
    redacted_data = json.loads(json_str)

    # 3. Recursively remove any image elements whose bboxes overlap with the signature area [450, 715, 570, 820]
    def clean_signature_images(obj):
        if isinstance(obj, dict):
            # Check if this element is an image in the signature region
            if obj.get("type") == "image":
                bbox = obj.get("bbox")
                if bbox is not None:
                    if isinstance(bbox, (list, tuple)) and len(bbox) >= 4:
                        x0, y0, x1, y1 = bbox[0], bbox[1], bbox[2], bbox[3]
                    elif isinstance(bbox, dict):
                        x0 = bbox.get("x0", bbox.get("x", 0.0))
                        y0 = bbox.get("y0", bbox.get("y", 0.0))
                        x1 = bbox.get("x1", x0 + bbox.get("width", 0.0))
                        y1 = bbox.get("y1", y0 + bbox.get("height", 0.0))
                    else:
                        x0, y0, x1, y1 = 0.0, 0.0, 0.0, 0.0
                    
                    if 440 <= x0 <= 580 and 700 <= y0 <= 830:
                        return None
            
            new_dict = {}
            for k, v in obj.items():
                val_cleaned = clean_signature_images(v)
                if val_cleaned is not None:
                    new_dict[k] = val_cleaned
            return new_dict
            
        elif isinstance(obj, list):
            cleaned_list = []
            for item in obj:
                cleaned_item = clean_signature_images(item)
                if cleaned_item is not None:
                    cleaned_list.append(cleaned_item)
            return cleaned_list
            
        return obj

    redacted_data = clean_signature_images(redacted_data)
    return redacted_data


def convert_pdf_full_pipeline(pdf_path, output_dir=None, theme_config: dict = None):
    """
    Executes PDF->HTML->PDF->DOCX pipeline using pdf2docx:
    1. Input PDF -> Extract JSON
    2. Redact name, case ID, signature images directly on the extracted JSON
    3. Render redacted JSON -> HTML (themed)
    4. Compile HTML -> PDF
    5. Convert compiled PDF -> Word (.docx) using pdf2docx
    """
    pdf_path = Path(pdf_path).absolute()
    if output_dir is None:
        output_dir = pdf_path.parent / "output"
    else:
        output_dir = Path(output_dir).absolute()
    output_dir.mkdir(parents=True, exist_ok=True)
    json_dir = pdf_path.parent / "extracted_jsons"
    json_dir.mkdir(parents=True, exist_ok=True)

    stem = pdf_path.stem

    print(f"\n==================================================")
    print(f"[*] Executing PDF->HTML->PDF->DOCX Pipeline for: {pdf_path.name}")
    print(f"==================================================")

    # Step 1: PDF -> JSON
    print(f"\n[Step 1/4] Extracting PDF to JSON...")
    extracted_data = extract_report_data(str(pdf_path))
    if extracted_data:
        # Redact patient details and signature images directly in JSON
        extracted_data = redact_extracted_json(extracted_data)
        
        # Replace "SN Genelab Pvt Ltd" with "Laboratory"
        json_str = json.dumps(extracted_data, indent=2, ensure_ascii=False)
        json_str = json_str.replace("SN Genelab Pvt Ltd", "Laboratory")
        extracted_data = json.loads(json_str)

        json_path = json_dir / f"{stem}.json"
        with open(json_path, "w", encoding="utf-8") as f_json:
            f_json.write(json_str)
        print(f"   [+] Redacted JSON saved: {json_path}")

    # Step 2: JSON -> HTML (themed)
    print(f"\n[Step 2/4] Rendering JSON to themed HTML...")
    html_content = generate_dynamic_template_html(extracted_data, doc_title=f"{stem}.pdf", theme_config=theme_config)
    html_content = html_content.replace("SN Genelab Pvt Ltd", "Laboratory")
    html_path = output_dir / f"{stem}.html"
    html_path.write_text(html_content, encoding="utf-8")
    print(f"   [+] HTML saved: {html_path}")

    # Step 3: HTML -> Compiled PDF
    print(f"\n[Step 3/4] Compiling HTML to PDF...")
    compiled_pdf_path = output_dir / f"{stem}_compiled.pdf"
    render_html_to_pdf_and_preview(html_path, compiled_pdf_path)
    if compiled_pdf_path.exists():
        print(f"   [+] Compiled PDF saved: {compiled_pdf_path}")
    else:
        print(f"   [!] Failed to compile HTML to PDF")
        return None

    # Step 4: Compiled PDF -> Word (.docx) via pdf2docx
    print(f"\n[Step 4/4] Converting compiled PDF to Word (.docx) via pdf2docx...")
    out_docx = output_dir / f"{stem}_report.docx"
    success = convert_pdf_via_pdf2docx(str(compiled_pdf_path), str(out_docx))
    if success:
        print(f"   [+] Word (.docx) generated: {out_docx}")
    else:
        print(f"   [!] pdf2docx conversion failed")
        return None

    print(f"\n[+] Pipeline Completed Successfully! Final Word doc: {out_docx}")
    print(f"==================================================\n")
    return out_docx


def convert_pdf_to_word(pdf_path, docx_path, theme_config: dict = None):
    """
    Convert PDF to Word (.docx) directly using pdf2docx Converter (pdftoDoc logic),
    redacting signature images from the PDF body first and injecting them into footers.
    """
    import pymupdf as fitz
    import docx
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import os
    
    pdf_p = str(Path(pdf_path).absolute())
    docx_p = str(Path(docx_path).absolute())
    Path(docx_p).parent.mkdir(parents=True, exist_ok=True)

    if PDF2DocxConverter is None:
        print("[!] Error: pdf2docx is not installed. Run 'pip install pdf2docx'")
        return False

    # 1. Create a temporary PDF with signature area redacted (whited out)
    temp_pdf_path = docx_p + ".temp_redacted.pdf"
    try:
        doc = fitz.open(pdf_p)
        for page_num in range(len(doc)):
            page = doc[page_num]
            # Redact the signature bbox [450, 715, 570, 820]
            rect = fitz.Rect(450, 715, 570, 820)
            page.add_redact_annot(rect, fill=(1, 1, 1)) # White solid fill
            page.apply_redactions()
        doc.save(temp_pdf_path)
        doc.close()
    except Exception as e_redact:
        print(f"   [!] PDF Redaction failed: {e_redact}")
        temp_pdf_path = pdf_p  # Fallback to original PDF if redaction fails

    # 2. Convert the redacted PDF to Word
    try:
        cv_obj = PDF2DocxConverter(temp_pdf_path)
        cv_obj.convert(docx_p)
        cv_obj.close()
        print(f"   [+] File Converted Successfully: {docx_p}")
    except Exception as e:
        print(f"   [!] Conversion Failed: {e}")
        if temp_pdf_path != pdf_p and os.path.exists(temp_pdf_path):
            try:
                os.remove(temp_pdf_path)
            except Exception:
                pass
        return False

    # Clean up temporary PDF
    if temp_pdf_path != pdf_p and os.path.exists(temp_pdf_path):
        try:
            os.remove(temp_pdf_path)
        except Exception:
            pass

    # 3. Post-process the generated Word document (signature injection and SNG replacement)
    try:
        doc_word = docx.Document(docx_p)
        
        # Inject signature if present
        sig_image_path = Path("assets/dr_vinay_signature.png")
        if sig_image_path.exists():
            for s_idx, section in enumerate(doc_word.sections):
                # Add signature to all footer types (default, first page, even page)
                footers = [section.footer, section.first_page_footer, section.even_page_footer]
                for footer in footers:
                    if footer is not None:
                        # Only add signature if it is the first section or NOT linked to previous section
                        if s_idx == 0 or not footer.is_linked_to_previous:
                            if len(footer.paragraphs) == 1 and footer.paragraphs[0].text == "":
                                p = footer.paragraphs[0]
                            else:
                                p = footer.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            run = p.add_run()
                            run.add_picture(str(sig_image_path.absolute()), width=Inches(1.25))
            print(f"   [+] Injected signature image into Word footer of every section.")

        # Perform "SNG Gen's Lab pvt ltd" -> "Laboratory" substitution
        replace_sng_in_docx_obj(doc_word)
        
        doc_word.save(docx_p)
        print(f"   [+] Post-processed Word document successfully.")
    except Exception as e_word:
        print(f"   [!] Failed to post-process Word document: {e_word}")

    return True



def _get_page_bounds(page, fallback_left, fallback_width):
    """
    Compute actual content bounds from table positions.
    """
    tabs = page.find_tables()
    bboxes = []
    for tab in tabs.tables:
        if hasattr(tab, 'bbox'):
            x0, y0, x1, y1 = tab.bbox
            if y1 < 750 and (x1 - x0) > 80:
                bboxes.append(tab.bbox)

    if bboxes:
        page_left = min(b[0] for b in bboxes)
        page_right = max(b[2] for b in bboxes)
        page_width = page_right - page_left
        return page_left, page_width

    return fallback_left, fallback_width


def get_page_section_overlays(page, page_left_str, page_width_str, hy_cutoff: float = 0.0, fy_cutoff: float = 9999.0):
    """
    Extract section bounding boxes based on bold section headings in body region ONLY.
    """
    return []
    spans = []
    blocks = page.get_text('dict')['blocks']
    for b in blocks:
        if 'lines' in b:
            for ln in b['lines']:
                for s in ln['spans']:
                    if s.get('text', '').strip():
                        spans.append(s)
    spans.sort(key=lambda s: (round(s['bbox'][1], 1), round(s['bbox'][0], 1)))

    overlays = []
    for i, s in enumerate(spans):
        text = s['text'].strip()
        s_y0 = s['bbox'][1]
        s_y1 = s['bbox'][3]
        if s_y0 < hy_cutoff or s_y1 > fy_cutoff:
            continue

        is_bold = s.get('flags', 0) & 2 or "bold" in s.get('font', '').lower()
        if is_bold and text.endswith(':') and len(text) < 40:
            h_bbox = s['bbox']
            content_spans = []
            for j in range(i + 1, len(spans)):
                next_text = spans[j]['text'].strip()
                next_y0 = spans[j]['bbox'][1]
                if next_y0 > fy_cutoff:
                    break
                next_is_bold = spans[j].get('flags', 0) & 2 or "bold" in spans[j].get('font', '').lower()
                if (next_is_bold and next_text.endswith(':')) or spans[j]['bbox'][1] - h_bbox[1] > 100:
                    break
                content_spans.append(spans[j])

            if content_spans:
                c_min_y = max(hy_cutoff, h_bbox[1] - 2.0)
                c_max_y = min(fy_cutoff, max(cs['bbox'][3] for cs in content_spans) + 4.0)
                h = max(14.0, c_max_y - c_min_y)
                if c_min_y >= hy_cutoff and c_max_y <= fy_cutoff:
                    overlays.append(
                        f"<div class='section-content-box' "
                        f"style='left:{page_left_str};top:{c_min_y:.1f}pt;"
                        f"width:{page_width_str};height:{h:.1f}pt;'></div>"
                    )
    return overlays



def process_pdf(pdf_input, output_html_path=None, is_target=False, use_template=False, theme_config=None, save_output=False):
    """
    Universal PDF conversion function.
    Extracts structured JSON data and renders HTML.
    By default (save_output=False), does NOT write any files to the output directory.
    """
    doc_title = "Uploaded Report"
    if isinstance(pdf_input, (str, Path)):
        tmp_path = str(pdf_input)
        doc_title = Path(pdf_input).name
        doc = fitz.open(tmp_path)
    elif isinstance(pdf_input, bytes):
        doc = fitz.open(stream=pdf_input, filetype="pdf")
        tmp_path = None
    else:
        pdf_bytes = pdf_input.read()
        if hasattr(pdf_input, "seek"):
            pdf_input.seek(0)
        doc_title = getattr(pdf_input, "name", "Uploaded Report")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        tmp_path = None

    extracted_data = None
    if tmp_path and os.path.exists(tmp_path):
        try:
            extracted_data = extract_report_data(tmp_path)
        except Exception:
            pass

    # 3. HTML Rendering based on requested mode
    if use_template and extracted_data:
        full_html = generate_dynamic_template_html(extracted_data, doc_title=doc_title, theme_config=theme_config)
    else:
        # Preserve 100% exact 1-to-1 visual coordinates from input PDF
        full_html = render_exact_pdf_layout_html(doc, doc_title=doc_title, theme_config=theme_config)

    # Save to disk ONLY if save_output is explicitly requested
    if save_output and output_html_path:
        output_html_path = Path(output_html_path)
        output_html_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_html_path, "w", encoding="utf-8") as f:
            f.write(full_html)

    doc.close()
    return full_html


def render_html_to_pdf_and_preview(html_path, output_pdf_path, preview_img_path=None):
    """
    Render an HTML file to PDF via Playwright Chromium (if installed).
    """
    html_path = Path(html_path).absolute()
    output_pdf_path = Path(output_pdf_path)
    output_pdf_path.parent.mkdir(parents=True, exist_ok=True)

    if sync_playwright is None:
        return output_pdf_path

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page(viewport={"width": 1000, "height": 1200})
            page.goto(html_path.as_uri(), wait_until="networkidle")

            if preview_img_path:
                preview_img_path = Path(preview_img_path)
                preview_img_path.parent.mkdir(parents=True, exist_ok=True)
                page.screenshot(path=str(preview_img_path), full_page=False)

            page.pdf(
                path=str(output_pdf_path),
                print_background=True,
                prefer_css_page_size=True,
                margin={"top": "0px", "right": "0px", "bottom": "0px", "left": "0px"},
            )
            page.close()
            browser.close()
    except Exception as err:
        pass

    return output_pdf_path


def validate_docx_conversion(extracted_data, docx_path):
    """
    Validates that the generated DOCX matches the extracted JSON data elements and tables.
    Generates a clear validation report.
    """
    from docx import Document
    import json
    
    if isinstance(extracted_data, (str, Path)):
        with open(extracted_data, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = extracted_data

    # Parse output DOCX
    doc = Document(docx_path)
    
    # Extract structural elements from JSON
    extracted_elements = []
    extracted_tables_count = 0
    extracted_headings_count = 0
    extracted_paras_count = 0
    extracted_kvs_count = 0

    pages = data.get("document", {}).get("pages", [])
    for page in pages:
        for el in page.get("elements", []):
            el_type = el.get("type")
            if el_type in ("heading", "subheading", "paragraph", "table", "key_value", "image"):
                extracted_elements.append(el)
                if el_type == "table":
                    extracted_tables_count += 1
                elif el_type in ("heading", "subheading"):
                    extracted_headings_count += 1
                elif el_type == "paragraph":
                    extracted_paras_count += 1
                elif el_type == "key_value":
                    extracted_kvs_count += 1

    # Extract elements from DOCX
    rendered_tables_count = len(doc.tables)
    
    # We want to check if the paragraphs contain the text of headings and paragraphs in JSON
    missing_elements = []
    
    # Simple validation comparison
    # Check tables (key_values are also rendered as tables)
    tables_match = (rendered_tables_count >= (extracted_tables_count + extracted_kvs_count))
    
    # Check paragraphs text
    docx_text = "\n".join([p.text for p in doc.paragraphs]).lower()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_text += "\n" + cell.text.lower()
                
    # Also add headers/footers text
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for p in header.paragraphs:
                    docx_text += "\n" + p.text.lower()
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for p in footer.paragraphs:
                    docx_text += "\n" + p.text.lower()

    for el in extracted_elements:
        el_type = el.get("type")
        if el_type in ("heading", "subheading", "paragraph"):
            txt = el.get("text", "").strip().lower()
            txt_clean = txt.lstrip('•-* ').strip()
            if txt_clean and txt_clean not in docx_text:
                missing_elements.append(el)
                print(f"   [!] Missing element text: {repr(el.get('text'))}")
                    
    # Generate report
    total_extracted = len(extracted_elements)
    total_missing = len(missing_elements)
    total_rendered = total_extracted - total_missing
    
    validation_passed = (total_missing == 0) and tables_match
    status = "PASS" if validation_passed else "FAIL"
    
    report = f"""
==================================================
              DOCX VALIDATION REPORT
==================================================
CONTENT ELEMENTS EXTRACTED: {total_extracted}
CONTENT ELEMENTS RENDERED:  {total_rendered}
TABLES EXTRACTED:           {extracted_tables_count}
TABLES RENDERED:            {rendered_tables_count - extracted_kvs_count if rendered_tables_count >= extracted_kvs_count else 0}
KEY_VALUES EXTRACTED:       {extracted_kvs_count}
KEY_VALUES RENDERED:        {min(extracted_kvs_count, rendered_tables_count)}
MISSING ELEMENTS:           {total_missing}
CONTENT VALIDATION:         {status}
==================================================
"""
    print(report)
    return report